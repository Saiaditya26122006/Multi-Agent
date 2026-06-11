"""
Expand all existing agent explanation documents to match the detailed style
of the final 3 agents (Summary, Devil's Advocate, Council).

Reads each existing document, expands the content to be much more detailed,
and overwrites with the expanded version.
"""

import logging
from pathlib import Path

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)


def expand_opportunity_analyst(doc_path: Path) -> None:
    """Expand Opportunity Analyst explanation with much more detail."""
    doc = Document()

    title_para = doc.add_heading("Opportunity Analyst Agent - Business Idea Validation", level=1)
    title_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    sections = {
        "What This Agent Does": """The Opportunity Analyst is the first agent in the entire business planning pipeline. It is the gatekeeper that decides whether a business idea is worth analyzing further. When the CEO submits a new business idea, the Opportunity Analyst evaluates it across multiple dimensions: market size, problem severity, solution feasibility, competitive landscape, and timing.

Think of it as the initial sanity check before the system invests computational resources (and the CEO's time) into a full multi-section business plan. The agent asks: Is this a real problem? Is the market big enough? Is the proposed solution technically feasible? Are there insurmountable competitive barriers? Is the timing right?

The Opportunity Analyst does not do deep market research or financial modeling. That comes later. Its job is to quickly assess whether the idea passes the plausibility test. If the CEO proposes 'a social network for pet rocks', the Opportunity Analyst flags that the market is too small and the problem is not real. If the CEO proposes 'AI-powered cancer drug discovery', the Opportunity Analyst flags that the solution requires 10+ years of R&D and regulatory approval, which may not fit a startup timeline.

The agent produces a structured assessment with: problem definition (is this a real problem worth solving?), market size estimate (TAM/SAM/SOM), solution feasibility (can this be built with available technology?), competitive landscape overview (who else is solving this?), market timing assessment (why now?), key assumptions (what are we assuming about the market, customer, or solution?), critical uncertainties (what do we not know that could invalidate this idea?), go/no-go recommendation (should we proceed with full analysis?), and confidence score.

The output is designed to be read in 5 minutes and give a clear signal: green light (proceed to full analysis), yellow light (promising but needs clarification on specific points), or red light (fundamental flaws, do not proceed).

The Opportunity Analyst also sets the foundation for all downstream agents. The market size estimate it produces becomes input for the Financial Agent. The competitive landscape it maps becomes input for the Marketing Agent. The problem definition it validates becomes the anchor for the SWOT Analysis. If the Opportunity Analyst gets the fundamentals wrong, every downstream agent inherits that error.""",

        "Why We Built This Agent": """We built the Opportunity Analyst because not every business idea deserves a full business plan. Generating a complete 11-section business plan takes 15-30 minutes of computational time, costs $5-15 in LLM API calls, and produces 50+ pages of analysis. That investment only makes sense if the underlying idea is plausible.

In the early design of the system, we had no initial filter. The CEO could submit any idea, and the system would dutifully analyze it. We quickly found that CEOs would submit half-baked ideas just to see what the system would say. The system would produce a polished-looking business plan for a fundamentally flawed idea, which wasted everyone's time and eroded trust in the system.

The Opportunity Analyst solves the garbage-in problem. It catches fundamentally flawed ideas early before they consume resources. If the market is too small, if the problem is not real, if the solution is not feasible, the system tells the CEO immediately instead of pretending to build a business plan around a non-viable idea.

We also built it to educate the CEO about what makes a good business idea. By reading Opportunity Analyst assessments on multiple ideas, the CEO learns what 'market size' means, what 'competitive moat' means, what 'timing' means. Over time, the CEO submits higher-quality ideas because they internalized the evaluation criteria.

The agent also protects the system's credibility. If the system produces a detailed business plan for an obviously bad idea (like 'Uber for dog walking in rural Montana'), external stakeholders will question the system's judgment. But if the system says upfront 'this idea has a fatal flaw: market too small', it demonstrates good judgment and earns trust.

Finally, we built it because business planning should start with the market, not the solution. Many first-time founders fall in love with their solution (a cool technology, a clever product feature) without validating that the market needs it. The Opportunity Analyst forces the CEO to articulate the problem first, the customer pain, and the market opportunity before jumping to solution details.""",

        "Key Features": """First, it evaluates ideas across 5 critical dimensions: problem, market, solution, competition, and timing. This prevents the agent from focusing too narrowly on one aspect. An idea might have a big market but no competitive moat. Another idea might have a great solution but bad timing (too early or too late). The agent forces all dimensions to be considered.

Second, it uses a structured problem definition framework. The agent asks: What is the problem? Who has this problem? How painful is it? How do they solve it today? Why is the current solution inadequate? This prevents vague problem statements like 'universities need better software' and forces specificity like 'PhD students at research-intensive EU universities need faster epistemic validation of manuscripts to reduce desk rejection rates, currently taking 2-4 weeks of manual work per submission.'

Third, it produces a TAM/SAM/SOM market sizing estimate. TAM (Total Addressable Market) is the total revenue opportunity if the product achieved 100% market share. SAM (Serviceable Addressable Market) is the portion of TAM the business can realistically target given geographic, product, or customer constraints. SOM (Serviceable Obtainable Market) is the realistic short-term market share the business can capture. The agent does not just give one number. It shows the full funnel so the CEO understands what 'market size' actually means.

Fourth, it maps the competitive landscape with a positioning grid. The agent identifies 3-5 existing solutions (direct competitors, substitutes, and the status quo) and plots them on a 2x2 grid (usually cost vs features, or incumbents vs startups, or horizontal vs vertical). This shows the CEO where their idea fits and whether there is white space or if the market is saturated.

Fifth, it assesses timing using a 'why now?' framework. Good business ideas are not just good in theory. They are good now. The agent asks: What changed recently that makes this idea viable today but not 5 years ago? Is it a technology unlock (LLMs became cheap and good), a regulatory change (GDPR created compliance demand), a market shift (remote work increased demand for async tools), or a demographic trend (Gen Z users expect X)? If the agent cannot articulate a compelling 'why now?', that is a red flag.

Sixth, it produces explicit assumptions and uncertainties. The agent does not pretend to know things it does not know. If it estimates market size at $5B but that estimate is based on one industry report with no validation, the agent flags that: 'Market size assumed at $5B based on Gartner 2024 report, no primary data validation, uncertainty: high.' This transparency helps the CEO understand what is fact versus guess.

Seventh, it uses a confidence-calibrated go/no-go recommendation. The agent does not just say 'proceed' or 'do not proceed'. It says 'proceed with medium confidence' or 'do not proceed, high confidence this idea has fatal flaws' or 'proceed conditionally, clarify customer segment first'. The recommendation matches the strength of the evidence.

Eighth, it identifies critical questions that need answering before full analysis. If the idea is promising but has gaps, the agent lists 3-5 specific questions the CEO should answer: 'Who is the target customer - PhD students, professors, or university admin?' 'What is the willingness to pay - is this a $100/year tool or a $10K/year tool?' These questions become the input for the clarification loop (Phase 1 L1 agent).""",

        "How It Works Technically": """The Opportunity Analyst is a Claude Sonnet-powered agent that runs first in the pipeline. It uses the Intelligence Engine (4-step reasoning: Decompose, Produce, Challenge, Revise) to evaluate the business idea.

The agent is triggered when the CEO submits a new business idea via Telegram or the Streamlit UI. The input is the CEO's raw description (typically 1-3 paragraphs) plus any structured data the CEO provided (target market, customer type, pricing ballpark, geography).

The input schema includes: business_idea_description, target_market (B2B/B2C/B2B2C), customer_segment (who buys this), problem_statement (what problem does this solve), proposed_solution (high-level approach), geography (target region), any_constraints (budget, timeline, regulatory, team limitations).

The agent starts by reading the CEO's description and extracting the core elements: problem, customer, solution, market. If the CEO's description is vague ('I want to build a SaaS for universities'), the agent flags that it needs clarification before proceeding. If the description is specific, the agent proceeds with evaluation.

The agent uses Claude Sonnet (not Haiku) because opportunity evaluation requires strategic thinking, market sense, and the ability to reason about competition and timing. Haiku is too literal and often misses nuance in competitive dynamics.

The system prompt tells the agent: 'You are evaluating whether a business idea is worth pursuing. Be skeptical but not dismissive. Flag fatal flaws clearly. Highlight promising aspects. Quantify market size with realistic assumptions. Compare to existing solutions honestly. Ask critical questions the CEO needs to answer. Output a clear go/no-go recommendation with reasoning.'

The agent does not access external APIs in Phase 1 (no live market data, no competitive intelligence tools). It relies on Claude's training data (knowledge cutoff January 2025) plus any CEO-provided data. In Phase 2, the agent is enhanced to call web search APIs for live market data, but Phase 1 is closed-loop.

The output schema is a Pydantic model with: section_number ('1'), problem_definition (what problem, who has it, how painful, current solution, why inadequate), market_size_estimate (TAM, SAM, SOM with assumptions), solution_feasibility (technical feasibility, resource requirements, timeline estimate), competitive_landscape (3-5 competitors or substitutes, positioning grid, white space analysis), timing_assessment (why now, recent changes that enable this idea), key_assumptions (list of assumptions with confidence levels), critical_uncertainties (list of unknowns that could invalidate the idea), critical_questions (3-5 questions the CEO must answer), recommendation (proceed/do not proceed/clarify), recommendation_rationale (why this recommendation), confidence_score (high/medium/low), assumptions_used, uncertainties, input_tokens, output_tokens.

The agent writes its output to the agent_outputs table in Supabase with section_number='1'. It also updates the session state to AWAITING_CLARIFICATION (if the recommendation is 'clarify') or RESEARCH_RUNNING (if the recommendation is 'proceed') or PAUSED (if the recommendation is 'do not proceed').

Error handling: if the agent fails (API timeout, invalid output, Pydantic validation error), the system logs the failure and sends a Telegram message to the CEO: 'Opportunity analysis failed due to technical error. Please resubmit your idea or contact support.' The session state is set to PAUSED and the CEO is notified.""",

        "How It Connects to Other Agents": """The Opportunity Analyst is the root node in the agent dependency graph. No agent runs before it. All downstream agents depend on it either directly or indirectly.

The execution order is: (1) CEO submits business idea. (2) Opportunity Analyst evaluates it. (3) If recommendation is 'do not proceed', the pipeline stops and the CEO is notified. (4) If recommendation is 'clarify', the L1 Clarification Agent (Phase 1) is triggered to ask follow-up questions. (5) If recommendation is 'proceed', the Mother Agent (Phase 2) triggers the next section agents (Environment Research, Organisation Designer, etc.).

The Opportunity Analyst output feeds directly into multiple downstream agents. The Environment Research Agent (Section 3) uses the competitive landscape and market size estimates as input. The Marketing Agent (Section 8) uses the customer segment definition and problem statement. The Financial Agent (Section 12) uses the market size estimate (TAM/SAM/SOM) to sanity-check revenue projections. The Exit Strategy Agent (Section 14) uses the competitive landscape to identify potential acquirers.

The Opportunity Analyst also sets the baseline confidence level for the entire plan. If the Opportunity Analyst returns low confidence (because market size is uncertain or competitive landscape is unclear), that low confidence propagates downstream. The Financial Agent cannot have high confidence in revenue projections if the market size itself is uncertain.

The agent also feeds into the assumption registry (Phase 2 feature). Every assumption the Opportunity Analyst makes ('market size is $5B', 'customer acquisition cost will be $500') is logged in the assumptions table with source='opportunity_analyst', confidence_level, and validation_status='unvalidated'. As the CEO provides data or pilots validate assumptions, the assumption registry is updated, and downstream agents inherit the validated assumptions.

In Phase 3 (future), the Opportunity Analyst will also interact with the RAG knowledge base. If the CEO previously analyzed a similar business idea, the Opportunity Analyst will retrieve that analysis and use it as a reference point: 'You previously evaluated a similar idea (academic software for EU universities) and flagged low willingness-to-pay as a concern. How is this idea different?'

The Opportunity Analyst does not interact with the Devil's Advocate or Council agents directly. Those agents operate at the end of the pipeline, reviewing the complete plan. But the Opportunity Analyst's output is read by those agents as context. If the Council is deliberating whether to approve a plan, the Investor persona will review the Opportunity Analyst's market size estimate and ask 'how confident are we in this $5B TAM claim?'""",

        "Why This Matters (Real-World Impact)": """The Opportunity Analyst is the reason the multi-agent system can be trusted to give honest feedback. Without it, the system would be a 'yes-man' that produces a business plan for every idea the CEO submits, no matter how flawed.

In real-world usage, the Opportunity Analyst saves the CEO from wasting time on bad ideas. If the CEO is considering 5 different business ideas, they can run all 5 through the Opportunity Analyst in 30 minutes and immediately see which ones are worth pursuing. The agent might flag that Idea A has too much competition, Idea B has a tiny market, Idea C has great potential but needs clarification on pricing, and Ideas D and E are strong. The CEO can then focus their attention on C, D, and E instead of spreading effort across all 5.

The agent also prevents the sunk cost fallacy. If the CEO has already spent weeks thinking about an idea, they are emotionally invested and less likely to abandon it even when evidence suggests they should. The Opportunity Analyst provides an external, unbiased assessment that gives the CEO permission to kill a bad idea early instead of investing months and thousands of dollars into something fundamentally flawed.

The market sizing discipline is particularly impactful. Many first-time founders massively overestimate their market size ('the education market is $1 trillion!') without realizing that they can only realistically target a tiny fraction of that. The Opportunity Analyst forces the TAM/SAM/SOM breakdown, which gives the CEO a realistic view of the addressable market. This prevents situations where the Financial Agent projects $100M Year 5 revenue but the realistic SOM is only $50M total.

The competitive landscape mapping is also critical. First-time founders often believe their idea is unique when it is not. The Opportunity Analyst does the competitive research and shows the CEO: 'Here are 5 companies already doing something similar. Here is how they are positioned. Here is where you might fit.' This prevents the CEO from building something that already exists or entering a saturated market without a differentiation strategy.

The 'why now?' assessment helps the CEO articulate their timing thesis. Investors always ask 'why now?' If the CEO cannot answer that, the idea is probably not ready. The Opportunity Analyst forces that question early and helps the CEO develop a compelling timing narrative.

Finally, the Opportunity Analyst sets the tone for the entire system. By being honest, skeptical, and evidence-based in its assessment, it teaches the CEO that this system will not sugarcoat reality. That builds trust. The CEO learns that when the system says 'high confidence, proceed', that actually means something, because the system is willing to say 'do not proceed' when the idea is flawed.""",
    }

    for section_title, content in sections.items():
        doc.add_heading(section_title, level=2)
        paragraphs = content.strip().split("\n\n")
        for para_text in paragraphs:
            if para_text.strip():
                para = doc.add_paragraph(para_text.strip())
                para.style.font.size = Pt(11)

    doc.save(doc_path)
    logger.info(f"Expanded {doc_path.name} to {doc_path.stat().st_size} bytes")


def expand_environment_research(doc_path: Path) -> None:
    """Expand Environment Research explanation with much more detail."""
    doc = Document()

    title_para = doc.add_heading("Environment Research Agent - Market Landscape Analysis", level=1)
    title_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    sections = {
        "What This Agent Does": """The Environment Research Agent (Section 3) is the market intelligence specialist in the pipeline. After the Opportunity Analyst validates that the business idea is plausible, the Environment Research Agent goes deeper to map the external environment the business will operate in: regulatory landscape, macroeconomic trends, technology trends, customer behavior patterns, and ecosystem dynamics.

This agent answers questions like: What regulations will affect this business (GDPR, AI Act, industry-specific rules)? What macroeconomic trends are tailwinds or headwinds (recession, interest rates, VC funding environment)? What technology trends enable or threaten this business (new AI capabilities, platform shifts)? How are customer preferences changing? Who are the ecosystem players (suppliers, distribution partners, influencers)?

The Environment Research Agent does not focus on direct competitors. That was covered by the Opportunity Analyst in Section 1. This agent looks at the broader external forces that shape the opportunity space. It applies PESTLE analysis (Political, Economic, Social, Technological, Legal, Environmental) to identify factors the business cannot control but must navigate.

The agent produces a structured landscape report with: regulatory environment assessment (what laws, compliance requirements, or regulatory risks apply), macroeconomic trends (GDP growth, unemployment, interest rates, VC funding climate in target geography), technology trends (emerging technologies that enable or disrupt the business), customer behavior trends (how is the target customer changing their buying behavior, adoption patterns, or priorities), ecosystem analysis (who are the key partners, suppliers, distribution channels, influencers), geopolitical risks (trade restrictions, political instability, sanctions that could affect operations), key assumptions about the external environment, critical uncertainties (what external factors are most unpredictable), and confidence score.

The output is designed to give the CEO a realistic view of the external environment. If the agent identifies a regulatory headwind (like strict AI regulation in EU), that becomes input for the Tech Stack Agent to ensure compliance. If the agent identifies a macroeconomic tailwind (like increased university R&D budgets post-pandemic), that becomes input for the Marketing Agent to time the go-to-market strategy.

The Environment Research Agent also flags blind spots. If the CEO is focused on product features and ignoring the regulatory environment, this agent makes sure regulatory risks are surfaced early before they become blockers.""",

        "Why We Built This Agent": """We built the Environment Research Agent because great business ideas can fail due to external forces the founder did not anticipate. A startup might have a great product, strong team, and validated customer demand, but still fail because they did not account for regulatory changes, macroeconomic shifts, or technology platform changes.

In the early design of the system, we assumed the Opportunity Analyst's competitive landscape analysis was sufficient. We quickly realized that was too narrow. Competition is only one external force. Regulations, macroeconomic trends, and technology shifts can be just as consequential, and they often blindside founders because they are less visible than competitors.

The Environment Research Agent solves the blind spot problem. It forces the system to look beyond the immediate competitive landscape and consider the broader context. It asks: What is happening in the world that could affect this business, positively or negatively?

We also built it because first-time founders tend to be product-focused and externally naive. They think about features, pricing, and customers. They do not think about GDPR, interest rate environments, or the VC funding winter. The Environment Research Agent brings that external awareness into the planning process.

The agent also helps the CEO anticipate and prepare for external risks. If the agent flags that AI regulation is increasing in the EU, the CEO can plan for compliance costs and timelines upfront instead of being surprised later. If the agent flags that VC funding has dried up in a certain sector, the CEO can adjust their fundraising strategy to target alternative funding sources (bootstrapping, grants, angels) instead of assuming Series A will be easy.

We built it to make the business plan robust to external shocks. A plan that only works in ideal external conditions is fragile. A plan that accounts for regulatory headwinds, economic downturns, and technology disruptions is resilient. The Environment Research Agent is the resilience builder.

Finally, we built it because external environment analysis is tedious and time-consuming for humans. A founder could spend weeks researching regulations, reading economic reports, and tracking technology trends. The Environment Research Agent does that research in 2-3 minutes using Claude's training data plus (in Phase 2) live web search to get the latest regulatory updates and economic data.""",

        "Key Features": """First, it applies PESTLE framework systematically. PESTLE (Political, Economic, Social, Technological, Legal, Environmental) is a standard strategic analysis tool used by consultants and corporations. The agent evaluates each dimension and identifies the 2-3 most critical factors in each category. This prevents the agent from focusing too narrowly on one aspect (like just technology trends) and missing others (like regulatory changes).

Second, it tailors the analysis to the specific business context. The agent does not produce a generic PESTLE report. It customizes the analysis based on the business type (B2B vs B2C, SaaS vs marketplace, product vs service), geography (EU vs US vs Asia, each with different regulatory and economic contexts), and industry (education, healthcare, fintech, each with different regulatory intensity).

Third, it identifies tailwinds and headwinds explicitly. A tailwind is an external trend that helps the business (like increased demand for remote collaboration tools during COVID, or increased VC funding for AI startups in 2023-2024). A headwind is an external trend that hurts the business (like declining university budgets post-financial crisis, or GDPR compliance costs for small startups). The agent labels each factor as tailwind, headwind, or neutral so the CEO understands the net external environment.

Fourth, it assesses regulatory risk with geography-specific detail. If the business operates in the EU, the agent evaluates GDPR, AI Act, DPDP (if targeting India), sector-specific regulations (like medical device regulations for healthtech). If the business operates in the US, the agent evaluates HIPAA, COPPA, state-level privacy laws. The agent does not just say 'there are regulations.' It says 'GDPR Article 9 restricts processing of academic manuscript content without explicit consent, which affects your content analysis feature. Estimated compliance cost: $20K-50K for legal review and technical implementation.'

Fifth, it tracks macroeconomic indicators relevant to the business. For B2B SaaS, the agent tracks: VC funding trends (is capital available for startups in this sector?), corporate budget cycles (when do enterprises allocate budget for new tools?), unemployment rate (does the customer have resources to buy?). For B2C, the agent tracks: consumer spending trends, household savings rates, inflation. The agent uses recent data (from its training cutoff or, in Phase 2, live web search) to assess whether the macro environment is favorable or not.

Sixth, it identifies technology enablers and disruptors. An enabler is a new technology that makes the business possible or easier (like LLMs becoming cheap and capable in 2023, which enabled AI-powered SaaS products). A disruptor is a new technology that threatens the business (like AI tools that automate what the startup was planning to do manually). The agent flags both so the CEO can leverage enablers and defend against disruptors.

Seventh, it maps the ecosystem beyond direct competitors. The agent identifies: (1) Suppliers (who provides critical inputs? are they reliable? is there vendor lock-in risk?), (2) Distribution partners (who can help reach customers? what are their incentives?), (3) Complementary products (what tools do customers use alongside this product?), (4) Influencers (who shapes customer opinions in this space? academic thought leaders, industry analysts, bloggers?). This ecosystem map helps the CEO think about partnerships and go-to-market strategy.

Eighth, it produces confidence-calibrated uncertainty flags. The agent does not pretend to predict the future. If it identifies a regulatory risk but the timeline is unclear ('EU AI Act enforcement may begin in 2025 or 2026'), the agent flags that uncertainty: 'Regulatory timeline uncertain, range: 6-18 months, impact if earlier than expected: high.' This helps the CEO plan for multiple scenarios.""",

        "How It Works Technically": """The Environment Research Agent is a Claude Haiku-powered agent (cost optimization) that runs early in the pipeline, after the Opportunity Analyst (Section 1) completes.

The agent is triggered by the Mother Agent (Phase 2 orchestrator) once Section 1 output is available. The input is the Opportunity Analyst's output (business idea, market size, competitive landscape) plus any CEO-provided data about geography, target market, and industry.

The input schema includes: business_idea_summary, target_geography (EU/US/Asia/Global), industry_sector (education, healthcare, fintech, etc.), customer_type (B2B/B2C), stage (pre-launch, MVP, scaling), and the full Section 1 output (used for context).

The agent uses Claude Haiku instead of Sonnet because environment research is more about information retrieval and structured analysis than deep reasoning. Haiku is fast and cheap, and for a task like 'list relevant GDPR articles' or 'summarize macroeconomic trends in EU education sector', Haiku performs nearly as well as Sonnet.

The system prompt tells the agent: 'You are conducting external environment research for a business plan. Use PESTLE framework. Identify tailwinds and headwinds. Be specific about regulations (cite specific laws and articles). Be specific about macroeconomic trends (cite recent data with dates). Identify technology trends that enable or disrupt this business. Flag critical uncertainties where the external environment is unpredictable.'

In Phase 1, the agent relies on Claude's training data (knowledge cutoff January 2025). In Phase 2, the agent is enhanced with web search capability. The Mother Agent provides the Environment Research Agent with access to a web search API (via grounded_eval.py SEARCH_QUERIES), and the agent fetches live data on: recent regulatory updates (new laws, enforcement actions), latest macroeconomic data (GDP, unemployment, VC funding reports), recent technology news (new LLM releases, platform changes).

The output schema is a Pydantic model with: section_number ('3'), regulatory_environment (list of regulations that apply, compliance requirements, estimated compliance cost, timeline), macroeconomic_trends (GDP growth, interest rates, VC funding environment, consumer/business spending trends in target market), technology_trends (list of enabling technologies, list of disruptive threats, timeline for each), customer_behavior_trends (how is target customer changing buying behavior, adoption patterns, priorities), ecosystem_analysis (suppliers, partners, distributors, influencers, each with description and strategic importance), geopolitical_risks (trade restrictions, political instability, sanctions, each with probability and impact), external_tailwinds (list of positive external factors), external_headwinds (list of negative external factors), key_assumptions, critical_uncertainties, confidence_score, assumptions_used, uncertainties, input_tokens, output_tokens.

The agent writes its output to the agent_outputs table in Supabase with section_number='3'. It updates the session state to indicate Section 3 is complete, and the Mother Agent proceeds to the next section (Organisation Designer, Section 4).

Error handling: if web search fails (Phase 2), the agent falls back to training data and flags 'External data unavailable, using knowledge cutoff January 2025' in uncertainties. If the agent cannot determine regulatory requirements (because the industry or geography is niche), it flags 'Regulatory analysis incomplete, manual legal review required' and proceeds with low confidence.""",

        "How It Connects to Other Agents": """The Environment Research Agent (Section 3) runs early in the pipeline, after Section 1 (Opportunity Analyst) but before most other sections. It is a foundational section that provides context for downstream agents.

The execution order is: (1) Opportunity Analyst completes. (2) Environment Research Agent runs. (3) Organisation Designer, SWOT, Tech Stack, Marketing, Operations, Financial, Launch, and Exit Strategy agents run afterward (some in parallel, some sequentially based on dependencies).

The Environment Research output feeds into multiple downstream agents. The Tech Stack Agent (Section 6.5) uses the regulatory environment assessment to ensure compliance (if GDPR applies, the tech stack must include data residency in EU, encryption, consent management). The Marketing Agent (Section 8) uses the customer behavior trends to inform channel strategy (if the trend is toward lower ad engagement and higher community-led growth, the marketing plan shifts accordingly). The Financial Agent (Section 12) uses the macroeconomic trends to calibrate revenue assumptions (if VC funding is scarce, the plan may need to assume bootstrapping instead of seed funding). The Exit Strategy Agent (Section 14) uses the ecosystem analysis and technology trends to identify potential acquirers (if the tech trend is toward consolidation, that affects who might acquire the business).

The Environment Research Agent also feeds into the SWOT Synthesizer (Section 5). External tailwinds identified by the Environment Research Agent become Opportunities in the SWOT. External headwinds become Threats. The SWOT Synthesizer reads the Section 3 output and integrates it with internal strengths/weaknesses to produce the full SWOT matrix.

The agent does not depend on any other agent except the Opportunity Analyst (Section 1). It runs early because its output is needed by almost everyone else. It is a foundational research section.

The Environment Research Agent does not interact with the Devil's Advocate directly, but the Devil's Advocate may challenge its assumptions when reviewing downstream sections. For example, if the Financial Agent assumes 20% YoY revenue growth and justifies it by citing 'strong macroeconomic tailwinds from Section 3', the Devil's Advocate will check Section 3's confidence in that claim and may challenge it if the evidence is weak.

In Phase 3 (future), the Environment Research Agent will be continuously updated. Instead of running once at plan creation, it will run periodically (monthly or quarterly) to refresh the external environment assessment. If a new regulation is passed, if the macroeconomic environment shifts, or if a major technology disruption occurs, the agent re-runs and updates the business plan sections that depend on external environment assumptions.""",

        "Why This Matters (Real-World Impact)": """The Environment Research Agent is the reason the multi-agent system produces business plans that survive contact with reality. Many startup failures are not due to bad products or bad teams. They are due to external forces the founder did not anticipate: regulatory blockers, macroeconomic shifts, technology disruptions, customer behavior changes.

In real-world usage, the Environment Research Agent acts as the early warning system for external risks. If the agent flags that the EU AI Act will impose strict transparency requirements on AI-powered products, the CEO can plan for compliance upfront instead of being blindsided when they try to launch and realize they need 6 months of legal and technical work to comply.

The regulatory analysis is particularly impactful for first-time founders who are often naive about compliance. A founder building a healthtech product might not realize HIPAA compliance costs $50K-100K and takes 6-12 months. The Environment Research Agent surfaces that early, so the CEO can factor it into timelines and budgets.

The macroeconomic trend analysis helps the CEO make better timing decisions. If the agent identifies that VC funding for edtech startups has dried up, the CEO knows not to build a plan that assumes easy Series A fundraising. They can pivot to a bootstrap-first or grant-first strategy instead.

The technology trend analysis prevents the CEO from building something that is about to be disrupted. If the agent identifies that OpenAI just released a feature that directly competes with the startup's planned product, the CEO can pivot before investing months of development time.

The ecosystem mapping helps the CEO think strategically about partnerships and distribution. If the agent identifies that the target customer heavily relies on a specific platform (like universities use Canvas LMS), the CEO can explore a partnership or integration strategy to piggyback on that distribution channel.

The tailwinds/headwinds framing gives the CEO a net assessment of the external environment. If there are 5 tailwinds and 2 headwinds, the external environment is favorable. If there are 2 tailwinds and 5 headwinds, the CEO knows they are fighting uphill and may need to reconsider the idea or adjust the strategy to mitigate headwinds.

Finally, the Environment Research Agent makes the business plan defensible to external stakeholders. If the CEO pitches to an investor and the investor asks 'how do you plan to handle GDPR compliance?', the CEO can point to Section 3 of the business plan, which already analyzed the regulatory requirement, estimated the cost, and factored it into the tech stack and financials. That level of thoroughness builds investor confidence.""",
    }

    for section_title, content in sections.items():
        doc.add_heading(section_title, level=2)
        paragraphs = content.strip().split("\n\n")
        for para_text in paragraphs:
            if para_text.strip():
                para = doc.add_paragraph(para_text.strip())
                para.style.font.size = Pt(11)

    doc.save(doc_path)
    logger.info(f"Expanded {doc_path.name} to {doc_path.stat().st_size} bytes")


def main() -> None:
    """Expand all 10 agent explanation documents with much more detail."""
    output_dir = Path("/home/saiaditya26122006/multi-agent-system/explaination")

    logger.info("Expanding agent explanation documents to match detailed style...")

    # Expand Opportunity Analyst and Environment Research first
    expand_opportunity_analyst(output_dir / "opportunity_analyst_explanation.docx")
    expand_environment_research(output_dir / "environment_research_explanation.docx")

    logger.info("\n✅ First 2 documents expanded! Continuing with remaining 8...")

    # Note: This script will be run multiple times to expand all 10 documents
    # Next batch will include: organisation_designer, swot_synthesizer, tech_stack_agent
    # Then: marketing_strategy, operations, financial_modelling
    # Finally: launch_contingency, exit_strategy


if __name__ == "__main__":
    main()
