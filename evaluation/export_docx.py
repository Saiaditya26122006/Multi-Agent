"""
Export evaluation results to Word document business plan with professional styling.
"""

import json
import logging
import re
from datetime import datetime
from pathlib import Path
from typing import Dict, Any, List

from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

logger = logging.getLogger(__name__)

# Brand Colors
PRIMARY_COLOR = RGBColor(37, 99, 235)      # Blue
SECONDARY_COLOR = RGBColor(16, 185, 129)   # Green
ACCENT_COLOR = RGBColor(245, 158, 11)      # Amber
ERROR_COLOR = RGBColor(239, 68, 68)        # Red
MUTED_COLOR = RGBColor(107, 114, 128)      # Gray

# Confidence Colors
CONFIDENCE_COLORS = {
    "high": RGBColor(16, 185, 129),        # Green
    "medium-high": RGBColor(52, 211, 153), # Light green
    "medium": RGBColor(245, 158, 11),      # Amber
    "medium-low": RGBColor(251, 146, 60),  # Orange
    "low": RGBColor(239, 68, 68),          # Red
}


def _extract_readable_text(item: Any) -> str:
    """
    Extract readable text from a dict, list, or string.

    For dicts: tries common field names in order of preference.
    For lists: joins items.
    For strings: returns as-is.
    """
    if isinstance(item, dict):
        # Try common field names for readable content
        for field in ["statement", "uncertainty", "item", "description", "text", "content"]:
            if field in item and item[field]:
                return str(item[field])

        # If no common field found, format as key-value pairs
        # but skip internal/technical fields
        skip_keys = {"id", "because", "evidence_quality", "section_number", "confidence"}
        readable_parts = []
        for k, v in item.items():
            if k not in skip_keys and v:
                readable_parts.append(f"{k.replace('_', ' ')}: {v}")

        if readable_parts:
            return "; ".join(readable_parts)

        # Last resort: convert whole dict
        return str(item)

    elif isinstance(item, list):
        return ", ".join([_extract_readable_text(i) for i in item])

    else:
        return str(item)


def _strip_markdown(text: str) -> str:
    """Remove markdown formatting from LLM output for clean docx rendering."""
    text = re.sub(r"^#{1,6}\s*", "", text, flags=re.MULTILINE)
    text = re.sub(r"\*\*([^*]+)\*\*", r"\1", text)
    text = re.sub(r"\*([^*]+)\*", r"\1", text)
    text = re.sub(r"`([^`]+)`", r"\1", text)
    text = re.sub(r"```[^\n]*\n?", "", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


# Section metadata mapping
SECTION_NAMES = {
    "1": "Opportunity Analysis",
    "3": "Environment Research",
    "4": "Organisation Design",
    "5": "SWOT Synthesis",
    "6.5": "Tech Stack & Data Privacy",
    "8": "Marketing Strategy",
    "10": "Operations",
    "12": "Financial Modelling",
    "13": "Launch & Contingency",
    "14": "Exit Strategy & Cap Table",
    "executive_summary": "Executive Summary",
}


def export_to_docx(results_path: str) -> str:
    """
    Export grounded eval results to a Word document business plan.

    Args:
        results_path: Path to the results JSON file

    Returns:
        Full path to the saved .docx file
    """
    # Read results
    with open(results_path) as f:
        data = json.load(f)

    # Create outputs directory
    outputs_dir = Path("outputs")
    outputs_dir.mkdir(exist_ok=True)

    # Generate filename
    idea_name = data.get("idea_name", "business_plan").replace(" ", "_")
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    filename = f"business_plan_{idea_name}_{timestamp}.docx"
    output_path = outputs_dir / filename

    # Create document
    doc = Document()

    # Build the document
    _add_cover_page(doc, data)
    _add_executive_summary(doc, data)
    _add_agent_sections(doc, data)
    _add_gaps_page(doc, data)

    # Save
    doc.save(str(output_path))
    return str(output_path)


def _add_cover_page(doc: Document, data: Dict[str, Any]):
    """Add professional cover page with modern design."""
    # Top spacing
    for _ in range(4):
        doc.add_paragraph()

    # Company/Project Name
    company = doc.add_paragraph()
    company.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = company.add_run("🚀 EpistemicOS")
    run.font.size = Pt(20)
    run.font.color.rgb = PRIMARY_COLOR
    run.bold = True

    # Spacing
    doc.add_paragraph()

    # Title - Business Name
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run(data.get("idea_name", "Business Plan"))
    title_run.font.size = Pt(36)
    title_run.font.color.rgb = RGBColor(17, 24, 39)  # Dark gray
    title_run.bold = True

    # Underline
    underline = doc.add_paragraph("_" * 60)
    underline.alignment = WD_ALIGN_PARAGRAPH.CENTER
    underline.runs[0].font.color.rgb = PRIMARY_COLOR

    # Spacing
    doc.add_paragraph()

    # Subtitle
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = subtitle.add_run("BUSINESS PLAN")
    sub_run.font.size = Pt(18)
    sub_run.font.color.rgb = MUTED_COLOR
    sub_run.font.italic = True

    # Spacing
    for _ in range(3):
        doc.add_paragraph()

    # Date box
    date_para = doc.add_paragraph()
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_run = date_para.add_run(
        f"📅 Generated: {datetime.now().strftime('%B %d, %Y')}"
    )
    date_run.font.size = Pt(12)
    date_run.font.color.rgb = MUTED_COLOR

    # Spacing
    for _ in range(2):
        doc.add_paragraph()

    # Info box with background
    info_box = doc.add_paragraph()
    info_box.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_shading(info_box, RGBColor(243, 244, 246))  # Light gray background

    info_text = (
        "🤖 AI-Generated Business Plan\n\n"
        "This plan was created by a multi-agent AI system with explicit confidence levels, "
        "assumptions, and uncertainty tracking. All gaps in available data are flagged "
        "for further research."
    )
    info_run = info_box.add_run(info_text)
    info_run.font.size = Pt(10)
    info_run.font.italic = True
    info_run.font.color.rgb = MUTED_COLOR

    # Spacing
    for _ in range(2):
        doc.add_paragraph()

    # Metadata table
    table = doc.add_table(rows=4, cols=2)
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER
    table.style = 'Light Grid Accent 1'

    # Total tokens
    table.rows[0].cells[0].text = "📊 Total Tokens"
    table.rows[0].cells[1].text = f"{data.get('total_input_tokens', 0) + data.get('total_output_tokens', 0):,}"

    # Processing time
    table.rows[1].cells[0].text = "⏱️  Processing Time"
    table.rows[1].cells[1].text = f"{data.get('total_latency_seconds', 0):.1f}s"

    # Grounded data
    table.rows[2].cells[0].text = "🎯 CEO Data Grounding"
    table.rows[2].cells[1].text = "✅ Yes" if data.get('grounded') else "❌ No"

    # Confidence
    overall_conf = _calculate_overall_confidence(data)
    table.rows[3].cells[0].text = "📈 Overall Confidence"
    table.rows[3].cells[1].text = overall_conf.upper()

    # Style table cells
    for row in table.rows:
        for cell in row.cells:
            cell.paragraphs[0].runs[0].font.size = Pt(10)
            if cell == row.cells[0]:  # First column
                cell.paragraphs[0].runs[0].bold = True
                cell.paragraphs[0].runs[0].font.color.rgb = PRIMARY_COLOR

    doc.add_page_break()


def _add_executive_summary(doc: Document, data: Dict[str, Any]):
    """Add executive summary section with professional styling."""
    # Section header with icon
    heading = doc.add_heading("", level=1)
    run = heading.add_run("📋 Executive Summary")
    run.font.color.rgb = PRIMARY_COLOR

    sections = data.get("sections", {})
    exec_section = sections.get("executive_summary", {})
    output = exec_section.get("output", {})

    if not output:
        doc.add_paragraph("Executive summary not available.")
        doc.add_page_break()
        return

    # Confidence badge at top
    confidence = output.get("confidence_score", "unknown")
    _add_confidence_badge(doc, confidence)

    doc.add_paragraph()  # spacing

    # Main content in highlighted box
    summary_text = output.get("executive_summary", "")
    if summary_text:
        summary_para = doc.add_paragraph(_strip_markdown(summary_text))
        summary_para.runs[0].font.size = Pt(11)
        _add_shading(summary_para, RGBColor(249, 250, 251))  # Very light gray
    else:
        doc.add_paragraph("No executive summary generated.")

    doc.add_paragraph()  # spacing

    # Key metrics in table
    key_metrics = output.get("key_metrics", [])
    if key_metrics:
        doc.add_heading("🎯 Key Metrics", level=2)

        table = doc.add_table(rows=len(key_metrics), cols=1)
        table.style = 'Light List Accent 1'

        for idx, metric in enumerate(key_metrics):
            cell = table.rows[idx].cells[0]
            cell.text = f"✓ {_strip_markdown(_extract_readable_text(metric))}"
            cell.paragraphs[0].runs[0].font.size = Pt(10)

        doc.add_paragraph()  # spacing

    # Uncertainties in warning box
    uncertainties = output.get("uncertainties", [])
    if uncertainties:
        doc.add_heading("⚠️  Key Uncertainties", level=2)

        for unc in uncertainties:
            if isinstance(unc, dict):
                text = _strip_markdown(_extract_readable_text(unc))
                severity = unc.get("severity", "")
                if severity:
                    text = f"{text} [{severity.upper()}]"
            else:
                text = _strip_markdown(str(unc))

            p = doc.add_paragraph(f"• {text}")
            p.runs[0].font.italic = True
            p.runs[0].font.color.rgb = ACCENT_COLOR
            p.runs[0].font.size = Pt(10)
            _add_shading(p, RGBColor(254, 252, 232))  # Light yellow

    doc.add_page_break()


def _add_agent_sections(doc: Document, data: Dict[str, Any]):
    """Add all agent output sections with professional styling."""
    sections = data.get("sections", {})
    section_order = ["1", "3", "4", "5", "6.5", "8", "10", "12", "13", "14"]

    # Section icons mapping
    section_icons = {
        "1": "🔍",
        "3": "🌍",
        "4": "🏢",
        "5": "⚖️",
        "6.5": "💻",
        "8": "📢",
        "10": "⚙️",
        "12": "💰",
        "13": "🚀",
        "14": "🚪",
    }

    for section_num in section_order:
        section_data = sections.get(section_num)
        if not section_data:
            continue

        section_name = SECTION_NAMES.get(section_num, f"Section {section_num}")
        icon = section_icons.get(section_num, "📄")

        # Section header with icon and color
        heading = doc.add_heading("", level=1)
        run = heading.add_run(f"{icon} {section_num}. {section_name}")
        run.font.color.rgb = PRIMARY_COLOR

        output = section_data.get("output", {})
        if not output:
            doc.add_paragraph("Section not available.")
            doc.add_page_break()
            continue

        # Confidence badge
        confidence = output.get("confidence_score", "unknown")
        _add_confidence_badge(doc, confidence)

        doc.add_paragraph()  # spacing

        # Add main content fields
        _add_section_content(doc, output, section_num)

        # Assumptions in styled box
        assumptions = output.get("assumptions_used", [])
        if assumptions:
            doc.add_heading("💡 Key Assumptions", level=2)

            for assumption in assumptions:
                # Extract main statement from dict or use as-is if string
                if isinstance(assumption, dict):
                    statement = assumption.get("statement", _extract_readable_text(assumption))
                    confidence_level = assumption.get("confidence", "")
                    source = assumption.get("source", "")

                    p = doc.add_paragraph()
                    # Bullet
                    p.add_run("▸ ").font.color.rgb = SECONDARY_COLOR
                    # Statement
                    stmt_run = p.add_run(statement)
                    stmt_run.font.size = Pt(10)

                    # Metadata in smaller italic text
                    if confidence_level or source:
                        meta_parts = []
                        if confidence_level:
                            meta_parts.append(f"confidence: {confidence_level}")
                        if source:
                            meta_parts.append(f"source: {source}")
                        meta_p = doc.add_paragraph(f"      {' | '.join(meta_parts)}")
                        meta_p.runs[0].font.italic = True
                        meta_p.runs[0].font.size = Pt(8)
                        meta_p.runs[0].font.color.rgb = MUTED_COLOR
                else:
                    p = doc.add_paragraph()
                    p.add_run("▸ ").font.color.rgb = SECONDARY_COLOR
                    p.add_run(str(assumption)).font.size = Pt(10)

            doc.add_paragraph()  # spacing

        # Uncertainties in warning style
        uncertainties = output.get("uncertainties", [])
        if uncertainties:
            doc.add_heading("⚠️  Uncertainties", level=2)

            for unc in uncertainties:
                if isinstance(unc, dict):
                    text = _strip_markdown(_extract_readable_text(unc))
                    severity = unc.get("severity", "")
                    if severity:
                        text = f"{text} [Severity: {severity.upper()}]"
                else:
                    text = _strip_markdown(str(unc))

                p = doc.add_paragraph(f"• {text}")
                p.runs[0].font.italic = True
                p.runs[0].font.color.rgb = ACCENT_COLOR
                p.runs[0].font.size = Pt(10)
                _add_shading(p, RGBColor(254, 252, 232))  # Light yellow

            doc.add_paragraph()  # spacing

        # Unresolved challenges in error box
        challenges = output.get("_unresolved_challenges", [])
        if challenges:
            doc.add_heading("🚨 Unresolved Challenges", level=2)

            for challenge in challenges:
                _render_challenge(doc, challenge)

        doc.add_page_break()


def _add_section_content(doc: Document, output: Dict[str, Any], section_num: str):
    """Add the main content fields for a section with styled formatting."""
    # Skip metadata fields
    skip_fields = {
        "section_number",
        "confidence_score",
        "assumptions_used",
        "uncertainties",
        "_unresolved_challenges",
        "input_tokens",
        "output_tokens",
    }

    for key, value in output.items():
        if key in skip_fields:
            continue

        # Format field name with emoji
        field_name = key.replace("_", " ").title()
        heading = doc.add_heading("", level=2)
        run = heading.add_run(f"▶ {field_name}")
        run.font.color.rgb = SECONDARY_COLOR

        # Handle different value types
        if isinstance(value, list):
            for item in value:
                if isinstance(item, dict):
                    text = _strip_markdown(_extract_readable_text(item))
                    p = doc.add_paragraph()
                    bullet_run = p.add_run("• ")
                    bullet_run.font.color.rgb = PRIMARY_COLOR
                    text_run = p.add_run(text)
                    text_run.font.size = Pt(10)
                else:
                    p = doc.add_paragraph()
                    bullet_run = p.add_run("• ")
                    bullet_run.font.color.rgb = PRIMARY_COLOR
                    text_run = p.add_run(_strip_markdown(str(item)))
                    text_run.font.size = Pt(10)
        elif isinstance(value, dict):
            # For dict values, show as key: value pairs in a table
            table = doc.add_table(rows=len(value), cols=2)
            table.style = 'Light Grid Accent 1'

            row_idx = 0
            for k, v in value.items():
                skip_keys = {"id", "because", "evidence_quality"}
                if k not in skip_keys:
                    table.rows[row_idx].cells[0].text = k.replace('_', ' ').title()
                    table.rows[row_idx].cells[1].text = _strip_markdown(str(v))
                    table.rows[row_idx].cells[0].paragraphs[0].runs[0].bold = True
                    table.rows[row_idx].cells[0].paragraphs[0].runs[0].font.size = Pt(9)
                    table.rows[row_idx].cells[1].paragraphs[0].runs[0].font.size = Pt(9)
                    row_idx += 1
        else:
            # Plain text in styled paragraph
            p = doc.add_paragraph(_strip_markdown(str(value)))
            p.runs[0].font.size = Pt(10)
            _add_shading(p, RGBColor(249, 250, 251))  # Light gray background

        doc.add_paragraph()  # spacing between fields


def _add_shading(paragraph, color: RGBColor):
    """Add background shading to a paragraph."""
    try:
        shading_elm = OxmlElement('w:shd')
        # RGBColor stores RGB as tuple (r, g, b) - access directly
        r, g, b = color
        shading_elm.set(qn('w:fill'), f"{r:02x}{g:02x}{b:02x}")
        paragraph._element.get_or_add_pPr().append(shading_elm)
    except Exception as e:
        logger.warning(f"Failed to add shading: {e}")


def _render_challenge(doc: Document, challenge, section_badge: str = ""):
    """Render a single unresolved challenge with structured formatting."""
    if isinstance(challenge, dict):
        c_type = challenge.get("type", "unknown")
        location = challenge.get("location", "")
        problem = _strip_markdown(challenge.get("problem", str(challenge)))
        fix = _strip_markdown(challenge.get("fix", ""))
    else:
        c_type = "unknown"
        location = ""
        problem = _strip_markdown(str(challenge))
        fix = ""

    p = doc.add_paragraph()

    if section_badge:
        badge_run = p.add_run(f"[{section_badge}] ")
        badge_run.bold = True
        badge_run.font.color.rgb = PRIMARY_COLOR
        badge_run.font.size = Pt(9)

    type_label = c_type.upper().replace("_", " ")
    problem_run = p.add_run(f"[{type_label}] {problem}")
    problem_run.font.color.rgb = ERROR_COLOR
    problem_run.font.size = Pt(10)
    problem_run.bold = True

    _add_shading(p, RGBColor(254, 242, 242))

    if location:
        loc_p = doc.add_paragraph()
        loc_run = loc_p.add_run(f"    Location: {_strip_markdown(location)}")
        loc_run.font.size = Pt(9)
        loc_run.font.italic = True
        loc_run.font.color.rgb = MUTED_COLOR

    if fix:
        fix_p = doc.add_paragraph()
        fix_run = fix_p.add_run(f"    Suggested fix: {fix}")
        fix_run.font.size = Pt(9)
        fix_run.font.color.rgb = MUTED_COLOR


def _add_confidence_badge(doc: Document, confidence: str):
    """Add a styled confidence level badge."""
    conf_para = doc.add_paragraph()

    # Label
    label_run = conf_para.add_run("CONFIDENCE: ")
    label_run.font.size = Pt(9)
    label_run.font.bold = True
    label_run.font.color.rgb = MUTED_COLOR

    # Badge
    badge_run = conf_para.add_run(f" {str(confidence).upper()} ")
    badge_run.font.size = Pt(10)
    badge_run.font.bold = True

    # Color based on confidence level
    conf_lower = str(confidence).lower()
    badge_color = CONFIDENCE_COLORS.get(conf_lower, MUTED_COLOR)
    badge_run.font.color.rgb = badge_color

    # Add visual meter
    meter_text = _get_confidence_meter(conf_lower)
    meter_run = conf_para.add_run(f"  {meter_text}")
    meter_run.font.size = Pt(12)

    # Background shading for the whole paragraph
    _add_shading(conf_para, RGBColor(249, 250, 251))


def _get_confidence_meter(confidence: str) -> str:
    """Get a visual meter for confidence level."""
    meters = {
        "high": "█████",
        "medium-high": "████░",
        "medium": "███░░",
        "medium-low": "██░░░",
        "low": "█░░░░",
    }
    return meters.get(confidence, "░░░░░")


def _calculate_overall_confidence(data: Dict[str, Any]) -> str:
    """Calculate overall confidence from all sections."""
    sections = data.get("sections", {})
    confidence_scores = []

    confidence_map = {
        "high": 5,
        "medium-high": 4,
        "medium": 3,
        "medium-low": 2,
        "low": 1,
    }

    for section_data in sections.values():
        output = section_data.get("output", {})
        conf = str(output.get("confidence_score", "")).lower()
        if conf in confidence_map:
            confidence_scores.append(confidence_map[conf])

    if not confidence_scores:
        return "unknown"

    avg_score = sum(confidence_scores) / len(confidence_scores)

    # Map back to label
    if avg_score >= 4.5:
        return "high"
    elif avg_score >= 3.5:
        return "medium-high"
    elif avg_score >= 2.5:
        return "medium"
    elif avg_score >= 1.5:
        return "medium-low"
    else:
        return "low"


def _add_gaps_page(doc: Document, data: Dict[str, Any]):
    """Add final page listing all gaps and uncertainties with professional styling."""
    heading = doc.add_heading("", level=1)
    run = heading.add_run("🔍 Data Gaps & Knowledge Limitations")
    run.font.color.rgb = ERROR_COLOR

    intro = doc.add_paragraph(
        "This section aggregates all uncertainties and unresolved challenges "
        "identified during the business plan generation. Addressing these gaps "
        "would improve the plan's accuracy and confidence."
    )
    intro.runs[0].font.size = Pt(11)
    _add_shading(intro, RGBColor(254, 252, 232))  # Light yellow

    doc.add_paragraph()  # spacing

    sections = data.get("sections", {})
    all_uncertainties = []
    all_challenges = []

    for section_num, section_data in sections.items():
        output = section_data.get("output", {})
        section_name = SECTION_NAMES.get(section_num, f"Section {section_num}")

        uncertainties = output.get("uncertainties", [])
        for unc in uncertainties:
            all_uncertainties.append((section_name, unc))

        challenges = output.get("_unresolved_challenges", [])
        for challenge in challenges:
            all_challenges.append((section_name, challenge))

    # Summary stats table
    stats_table = doc.add_table(rows=2, cols=2)
    stats_table.style = 'Light Grid Accent 1'

    stats_table.rows[0].cells[0].text = "⚠️  Total Uncertainties"
    stats_table.rows[0].cells[1].text = str(len(all_uncertainties))
    stats_table.rows[1].cells[0].text = "🚨 Critical Challenges"
    stats_table.rows[1].cells[1].text = str(len(all_challenges))

    for row in stats_table.rows:
        row.cells[0].paragraphs[0].runs[0].bold = True
        row.cells[1].paragraphs[0].runs[0].font.color.rgb = ERROR_COLOR
        row.cells[1].paragraphs[0].runs[0].bold = True

    doc.add_paragraph()  # spacing

    # List uncertainties
    if all_uncertainties:
        doc.add_heading("⚠️  All Uncertainties", level=2)

        for section_name, unc in all_uncertainties:
            p = doc.add_paragraph()

            # Section badge
            badge_run = p.add_run(f"[{section_name}] ")
            badge_run.bold = True
            badge_run.font.color.rgb = PRIMARY_COLOR
            badge_run.font.size = Pt(9)

            if isinstance(unc, dict):
                text = _strip_markdown(_extract_readable_text(unc))
                severity = unc.get("severity", "")
                if severity:
                    text = f"{text} [Severity: {severity.upper()}]"
            else:
                text = _strip_markdown(str(unc))

            text_run = p.add_run(text)
            text_run.font.size = Pt(10)
            text_run.font.italic = True

            _add_shading(p, RGBColor(254, 252, 232))  # Light yellow

        doc.add_paragraph()  # spacing

    # List unresolved challenges
    if all_challenges:
        doc.add_heading("🚨 Critical Unresolved Challenges", level=2)

        for section_name, challenge in all_challenges:
            _render_challenge(doc, challenge, section_badge=section_name)

        doc.add_paragraph()  # spacing

    # Recommendation box
    doc.add_heading("💡 Recommendations", level=2)
    rec = doc.add_paragraph(
        "Priority Actions:\n\n"
        "1. Address CRITICAL challenges first — these block plan viability\n"
        "2. Research HIGH-severity uncertainties using suggested databases:\n"
        "   • Passport, GlobalData — market sizing\n"
        "   • CB Insights, FACTIVA — competitive intelligence\n"
        "   • Statista, WARC — customer data\n"
        "   • Mergermarket — valuation comps\n\n"
        "3. Update the Knowledge Base with findings and re-run evaluation\n"
        "4. Review MEDIUM-severity gaps for investor deck preparation"
    )
    rec.runs[0].font.size = Pt(10)
    _add_shading(rec, RGBColor(239, 246, 255))  # Light blue
