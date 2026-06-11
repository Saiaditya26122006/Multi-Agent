"""
Expand the remaining 8 agent explanation documents with very detailed content:
- Organisation Designer
- SWOT Synthesizer
- Tech Stack Agent
- Marketing Strategy
- Operations
- Financial Modelling
- Launch & Contingency
- Exit Strategy
"""

import logging
from pathlib import Path

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)


def create_expanded_doc(title: str, sections: dict[str, str], output_path: Path) -> None:
    """Create expanded Word document with detailed sections."""
    doc = Document()
    title_para = doc.add_heading(title, level=1)
    title_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    for section_title, content in sections.items():
        doc.add_heading(section_title, level=2)
        paragraphs = content.strip().split("\n\n")
        for para_text in paragraphs:
            if para_text.strip():
                para = doc.add_paragraph(para_text.strip())
                para.style.font.size = Pt(11)

    doc.save(output_path)
    logger.info(f"Expanded {output_path.name} to {output_path.stat().st_size} bytes")


def expand_organisation_designer() -> tuple[str, dict]:
    """Return expanded content for Organisation Designer."""
    return ("Organisation Designer Agent - Team Structure and Capability Planning", {
        "What This Agent Does": """The Organisation Designer Agent (Section 4) is the human capital architect in the pipeline. After the Opportunity Analyst validates the business idea and the Environment Research Agent maps the external landscape, the Organisation Designer figures out what team is needed to execute the plan: roles, skills, headcount, hiring timeline, organizational structure, and personnel policies.

This agent answers questions like: What roles do we need (technical, business, operational)? What skills are critical versus nice-to-have? When should each role be hired (Month 0, Month 3, Year 2)? How should the team be structured (flat startup, functional departments, matrix)? What personnel policies matter (remote vs office, equity compensation, contractor vs FTE)? What capability gaps exist today that need to be filled?

The Organisation Designer does not just produce an org chart. It produces a strategic workforce plan that connects the business requirements (what needs to be built and delivered) to the people requirements (who needs to do it). If the Tech Stack Agent later determines the product requires AI/ML expertise, but the Organisation Designer did not plan to hire an ML engineer, that is a fatal gap. The Organisation Designer must anticipate the skills needed based on the business model and product vision.

The agent produces a structured team plan with: capability gap analysis (what skills does the founding team lack?), roles and responsibilities (for each role: title, responsibilities, required skills, when to hire, reporting structure), headcount plan (Year 1, Year 2, Year 3 headcount and fully-loaded cost), organizational structure (flat vs hierarchical, functional vs cross-functional), personnel policies (equity compensation, remote work, contractor first vs FTE first, performance management), team risk assessment (what happens if key hires are delayed or key people leave?), key assumptions about hiring and retention, critical uncertainties, and confidence score.

The output is designed to give the CEO a realistic view of the people side of the business. If the CEO is a solo non-technical founder planning to build a complex AI-powered SaaS product, the Organisation Designer will flag that a CTO co-founder or senior technical hire is critical and should happen by Month 0, not Year 2.

The Organisation Designer also estimates the fully-loaded cost of each role (salary plus benefits, employer taxes, equipment, training), which feeds directly into the Financial Agent's cost model. If the Organisation Designer assumes $80K/year for an engineer in EU but the actual market rate is $120K, the financial projections will be off by 50%.""",

        "Why We Built This Agent": """We built the Organisation Designer Agent because most first-time founders drastically underestimate the people requirements to build and scale a business. They think 'I will just hire a couple of engineers and we will figure it out.' They do not think about when to hire sales, when to hire ops, when to hire finance, what skills each role needs, or how to structure the team for growth.

In the early design of the system, we had no dedicated agent for org design. We assumed the Financial Agent could just estimate 'headcount costs' as a line item. We quickly realized that was insufficient. The financial model depends on knowing when people are hired, what they cost, and what they do. Without a detailed workforce plan, the financial projections are just guesses.

The Organisation Designer solves the people planning problem. It forces the CEO to think through the team requirements systematically. It asks: What roles do you need in Month 1 versus Month 12 versus Year 2? What happens if you cannot hire that critical role on time? What does each role actually do, and are you allocating enough capacity to deliver the product roadmap?

We also built it because hiring is one of the highest-risk areas for startups. Hiring the wrong person, hiring too early (burning cash on people before there is product-market fit), or hiring too late (missing deadlines because the team is understaffed) can all sink a startup. The Organisation Designer helps the CEO think through the hiring strategy and avoid common mistakes.

The agent also educates the CEO about market rates and hiring timelines. First-time founders often assume they can hire a senior engineer for $60K, or that they can fill a critical role in 2 weeks. The Organisation Designer provides realistic market rate estimates (based on geography and role seniority) and realistic hiring timelines (2-3 months per role), so the CEO does not build a plan on fantasy assumptions.

We built it to make the people plan consistent with the business plan. If the Marketing Agent assumes the company will close 50 customers Year 1, and each customer requires 2 hours of onboarding and support, that implies 100 hours of customer success work. Does the org plan include a customer success role? If not, who is doing that work? The Organisation Designer identifies those gaps and flags them.

Finally, we built it because team structure affects execution speed and quality. A flat startup with 5 people can move fast but lacks specialization. A hierarchical organization with clear departments can scale but moves slower. The Organisation Designer helps the CEO choose the right structure for their stage and industry, not just copy what other startups do.""",

        "Key Features": """First, it performs capability gap analysis. The agent starts by reading the business model, product vision, and target market, then identifies what skills are needed to execute. If the product is an AI-powered SaaS tool, the agent flags: need AI/ML expertise, need full-stack web development, need product design, need customer success, need sales (if B2B), need compliance/legal (if regulated industry). The agent then compares this to the founding team's skills and identifies gaps. If the founder is a domain expert with no technical background, the agent flags 'Critical gap: no CTO or technical co-founder, must hire or recruit by Month 0.'

Second, it produces detailed role definitions for each position. Each role includes: title (e.g. 'Senior ML Engineer'), responsibilities (bulleted list of what this person does day-to-day), required skills (technical skills, domain knowledge, soft skills), hire timeline (when this role becomes critical), reporting structure (who they report to, who reports to them), and assignment (is this role filled by a founder, a current team member, or 'to be hired'?). This level of detail prevents vague roles like 'we will hire an engineer' and forces specificity.

Third, it produces a phased headcount plan. The agent breaks hiring into phases: (1) Month 0 (founding team), (2) Months 1-6 (first hires), (3) Months 7-12 (Year 1 scaling), (4) Year 2, (5) Year 3. For each phase, the agent lists: number of people, roles being hired, total fully-loaded cost. This phased plan shows the CEO how the team scales over time and when big cost increases happen (typically when moving from 5 people to 10, or 10 to 20).

Fourth, it estimates fully-loaded cost per role. The agent does not just list salaries. It calculates: base salary (market rate for geography and seniority), benefits (healthcare, retirement, typically 15-20% of salary in US, 20-30% in EU), employer taxes (varies by country, typically 10-20% in US, 30-40% in EU), equipment and software (laptop, monitors, software licenses, ~$2K-5K per person), office costs if not remote (~$500-1500/person/month). The fully-loaded cost is typically 1.3x-1.5x the base salary. This realistic costing prevents the CEO from underestimating people costs.

Fifth, it designs the organizational structure based on stage and strategy. For a pre-launch startup (0-5 people), the agent recommends flat structure with minimal hierarchy. For a post-launch scaling startup (10-30 people), the agent recommends functional teams (engineering, sales/marketing, operations) with team leads. For a growth-stage company (50+ people), the agent recommends departments with managers and specialized roles. The structure is tailored to the stage, not one-size-fits-all.

Sixth, it defines personnel policies that affect talent attraction and retention. The agent recommends policies on: equity compensation (what % of equity should go into an ESOP, how much equity per role), remote work (fully remote, hybrid, office-first), contractor vs FTE strategy (hire contractors for flexibility early, convert to FTEs as you scale), performance management (how do you evaluate and promote people), and learning/development budget. These policies signal what kind of culture the company is building.

Seventh, it assesses team risks. The agent identifies: (1) Key person risk (if one critical person leaves, does the business collapse?), (2) Hiring delay risk (what happens if you cannot hire the ML engineer on time?), (3) Skills mismatch risk (what if the people you hire do not have the skills you thought they had?), (4) Retention risk (if you cannot pay market rates or offer competitive equity, will people leave after 6-12 months?). For each risk, the agent suggests mitigation strategies.

Eighth, it checks consistency with other sections. The Organisation Designer reads the Opportunity Analyst output (business model, product complexity) and the Environment Research output (regulatory requirements that might need specialized roles, like a compliance officer or data protection officer). It ensures the team plan matches the business requirements.""",

        "How It Works Technically": """The Organisation Designer is a Claude Haiku-powered agent that runs early-mid in the pipeline, after Sections 1 (Opportunity) and 3 (Environment Research) complete.

The agent is triggered by the Mother Agent once Sections 1 and 3 outputs are available. The input is the business model, product vision, target market, and any CEO-provided information about the founding team's skills and constraints (budget, hiring timeline preferences).

The input schema includes: business_model (B2B SaaS, B2C marketplace, etc.), product_description (high-level what is being built), founding_team (list of founders with their skills and roles), target_geography (affects salary estimates), stage (pre-launch, MVP, scaling), funding_status (bootstrapped, seed-funded, Series A), and prior section outputs (Section 1 for business requirements, Section 3 for any regulatory roles needed like DPO).

The agent uses Claude Haiku for cost optimization. Org design is more about structured thinking and templating than deep strategic reasoning, so Haiku performs well at a lower cost than Sonnet.

The system prompt tells the agent: 'You are designing the team structure for a startup. Identify capability gaps relative to the business requirements. Define roles with specific responsibilities and required skills. Estimate realistic market-rate salaries for the target geography. Plan hiring in phases based on criticality and affordability. Flag team risks. Be realistic about hiring timelines (2-3 months per role). Do not assume the founding team can do everything themselves.'

The output schema is a Pydantic model with: section_number ('4'), capability_gaps (list of skills missing from the founding team, each with severity: critical/high/medium), roles_and_responsibilities (list of role objects, each with title, responsibilities, required_skills, hire_timeline, assigned_to, reporting_to), headcount_plan (dict with year_1, year_2, year_3 keys, each containing count and total fully-loaded cost), org_structure (description of how the team is organized), personnel_policy (dict with keys: equity_strategy, remote_work_policy, contractor_vs_fte, performance_management), team_risks (list of risks with mitigation strategies), key_assumptions, critical_uncertainties, confidence_score, assumptions_used, uncertainties, input_tokens, output_tokens.

The agent writes its output to the agent_outputs table with section_number='4'. The Mother Agent marks Section 4 complete and proceeds to downstream agents.

Error handling: if the agent cannot determine market-rate salaries (because the geography or role is uncommon), it flags 'Salary estimates based on limited data, range: X-Y, manual validation recommended' in uncertainties and proceeds with best-estimate midpoint.""",

        "How It Connects to Other Agents": """The Organisation Designer (Section 4) runs early-mid pipeline, after Sections 1 and 3. It feeds into multiple downstream agents that depend on team capabilities and costs.

The execution order is: (1) Opportunity Analyst and Environment Research complete. (2) Organisation Designer runs. (3) Tech Stack, Marketing, Operations, Financial, and Exit Strategy agents run afterward (using Section 4 output).

The Organisation Designer output feeds into: (1) Tech Stack Agent (Section 6.5) uses the team capabilities to validate technical feasibility. If the org plan includes 2 engineers Year 1, but the tech stack requires building a complex AI platform, the Tech Stack Agent will flag insufficient capacity. (2) Financial Agent (Section 12) uses the headcount plan and fully-loaded costs to build the personnel expense model. The Financial Agent's P&L includes a line item for salaries that comes directly from Section 4. (3) Operations Agent (Section 10) uses the roles and responsibilities to validate that someone is assigned to each operational task (customer onboarding, support, billing, etc.). (4) Exit Strategy Agent (Section 14) uses the team structure to assess acquirer fit (strategic acquirers care about the strength of the team, especially in talent acquisitions).

The Organisation Designer also reads outputs from Section 1 (to understand what skills the business requires) and Section 3 (to check if any regulatory requirements demand specialized roles, like a Data Protection Officer for GDPR or a Compliance Officer for healthcare).

The agent does not interact with the Devil's Advocate directly, but the Devil's Advocate may challenge org design assumptions when reviewing the Financial Agent's output. If the Financial Agent assumes break-even in 18 months based on a 3-person team, the Devil's Advocate will check Section 4 to see if 3 people can realistically deliver the product roadmap, and may challenge the assumption if the workload does not match capacity.

In Phase 3 (future), the Organisation Designer will integrate with the Hiring Tracker (not yet built). As the CEO actually hires people, the Hiring Tracker records: role filled, actual salary, actual hire date. The Organisation Designer compares actual vs plan and updates assumptions for future hires.""",

        "Why This Matters (Real-World Impact)": """The Organisation Designer is the reason the multi-agent system produces business plans that account for the hardest part of startups: people. Many startups fail not because the product is bad or the market is small, but because they hired the wrong people, hired too early, hired too late, or could not afford to retain key talent.

In real-world usage, the Organisation Designer helps the CEO avoid catastrophic hiring mistakes. If the CEO is a non-technical founder planning to build a technical product, the agent will flag upfront: 'You need a technical co-founder or CTO-level hire by Month 0. If you cannot recruit a co-founder, expect to pay $120K-150K for a senior engineer in EU, or $150K-200K in US.' That clarity prevents the CEO from starting with a plan that assumes they can hire a junior contractor for $40K and deliver a production-ready SaaS platform.

The phased headcount plan helps the CEO align hiring with milestones. If the plan shows 3 people Year 1, 7 people Year 2, 12 people Year 3, the CEO can see when the next hiring wave happens and plan funding accordingly. The CEO knows they need to raise seed funding by Month 9 to afford the Year 2 hires, not wait until Month 18 when they have run out of cash.

The fully-loaded cost estimates are particularly impactful. First-time founders often underestimate people costs by 30-50% because they only think about salary, not benefits, taxes, equipment, and software. The Organisation Designer surfaces the true cost, which prevents the Financial Agent from projecting profitability based on fantasy cost assumptions.

The capability gap analysis is a forcing function for honest self-assessment. Many founders overestimate their own skills or assume they can learn on the job. The Organisation Designer forces the question: Can you actually build this product with your current skills, or do you need to bring in expertise? If the answer is 'bring in expertise', that has major implications for co-founder recruitment, hiring budget, and timeline.

The team risk assessment helps the CEO plan for contingencies. If the Organisation Designer flags 'Key person risk: if the CTO leaves in Year 1, product development halts', the CEO can mitigate by: (1) offering strong equity retention incentives, (2) documenting technical decisions to reduce bus factor, (3) hiring a second senior engineer earlier than planned to reduce single-point-of-failure risk.

Finally, the Organisation Designer makes the business plan credible to investors. When an investor reviews the plan and sees a detailed org chart with realistic salaries, a phased hiring plan, and explicit capability gap mitigation, they see a founder who has thought through the people side of the business. Investors know that execution depends on team, and a plan with no org design is a plan that will not execute.""",
    })


def expand_swot_synthesizer() -> tuple[str, dict]:
    """Return expanded content for SWOT Synthesizer."""
    return ("SWOT Synthesizer Agent - Strategic Position Assessment", {
        "What This Agent Does": """The SWOT Synthesizer Agent (Section 5) is the strategic position mapper in the pipeline. After the Opportunity Analyst, Environment Research, and Organisation Designer have analyzed the business from different angles, the SWOT Synthesizer pulls it all together into a classic SWOT matrix: Strengths, Weaknesses, Opportunities, Threats.

This agent synthesizes information from multiple sources to produce a holistic view of the business's strategic position. Strengths and Weaknesses are internal (things the business controls, like team capabilities, product differentiation, or financial resources). Opportunities and Threats are external (things the business does not control, like market trends, competition, or regulatory changes).

The SWOT Synthesizer does not do new analysis. It reads the outputs from Sections 1, 3, and 4, extracts the relevant strategic factors, categorizes them into the SWOT framework, and produces a synthesized assessment. If the Opportunity Analyst identified strong product differentiation, that becomes a Strength. If the Environment Research Agent identified increasing regulatory compliance costs, that becomes a Threat. If the Organisation Designer flagged lack of technical expertise on the founding team, that becomes a Weakness.

The agent produces a structured SWOT analysis with: Strengths (3-7 internal advantages the business has), Weaknesses (3-7 internal limitations or gaps), Opportunities (3-7 external favorable trends or conditions), Threats (3-7 external risks or headwinds), strategic implications (what does the SWOT matrix tell us about the right strategy?), priority actions (what should the CEO do first to leverage strengths, mitigate weaknesses, capture opportunities, or defend against threats?), key assumptions, critical uncertainties, and confidence score.

The output is designed to give the CEO a one-page strategic summary. Instead of reading 20 pages of analysis across Sections 1, 3, and 4, the CEO can read the SWOT matrix and immediately understand: What are we good at? What are we bad at? What external factors help us? What external factors hurt us?

The SWOT Synthesizer also sets up the downstream strategy sections (Marketing, Operations, Financial). The Marketing Agent should leverage Strengths and Opportunities. The Operations Agent should mitigate Weaknesses. The Financial Agent should model the impact of Threats on revenue and costs.""",

        "Why We Built This Agent": """We built the SWOT Synthesizer Agent because business planning requires synthesis, not just analysis. By Section 5, the system has produced three detailed analytical sections (Sections 1, 3, 4). The CEO is drowning in information. They need someone to step back and say: 'Here is what it all means. Here are the 5 most important strategic factors you need to pay attention to.'

In the early design of the system, we assumed the CEO would synthesize the SWOT themselves by reading Sections 1, 3, and 4. We quickly realized that was unrealistic. Most CEOs are not trained in strategic synthesis. They would read the sections, get overwhelmed by details, and miss the big picture. The SWOT Synthesizer does that synthesis work for them.

We also built it because the SWOT framework is a universal strategic tool. Investors, advisors, and co-founders all understand SWOT. When the CEO shares the business plan externally, the SWOT section is often the first thing people read after the executive summary. It is the standard language of strategic assessment.

The SWOT Synthesizer also acts as a quality check on the prior sections. If the Opportunity Analyst claimed there is strong market demand, but the SWOT Synthesizer cannot find any Opportunities related to market growth, that signals an inconsistency. The SWOT Synthesizer forces the system to ensure the strategic factors identified in earlier sections actually translate into actionable strategic insights.

We built it to make the strategic position clear and honest. The SWOT matrix does not hide weaknesses or threats. It surfaces them explicitly. This prevents the CEO from having an overly rosy view of the business. If there are 3 Strengths but 5 Threats, the SWOT matrix makes that imbalance obvious, and the CEO knows they need to de-risk the threats before proceeding.

Finally, we built it because strategic implications are not obvious. A strength is not inherently good if it is not relevant to the market. A threat is not inherently fatal if it can be mitigated. The SWOT Synthesizer does not just list factors. It interprets them: 'Given these strengths and opportunities, the right strategy is X. Given these weaknesses and threats, the CEO should prioritize Y.'""",

        "Key Features": """First, it synthesizes across multiple sections. The SWOT Synthesizer reads Sections 1, 3, and 4 and extracts strategic factors from each. From Section 1 (Opportunity): competitive differentiation, market size, problem severity. From Section 3 (Environment): regulatory tailwinds/headwinds, technology enablers/disruptors, macroeconomic trends. From Section 4 (Organisation): team capability gaps, hiring risks, cost structure. The agent then categorizes each factor into the SWOT framework.

Second, it applies SWOT classification rules. Strengths and Weaknesses are internal (under the business's control). Opportunities and Threats are external (not under the business's control). The agent does not misclassify. For example, 'strong founding team' is a Strength (internal). 'Increasing demand for AI tools' is an Opportunity (external). 'Lack of sales expertise' is a Weakness (internal). 'New regulations increasing compliance costs' is a Threat (external).

Third, it prioritizes the most critical factors. The agent does not list every possible strength, weakness, opportunity, and threat. It lists the 3-7 most important in each category. The criteria for importance are: (1) magnitude (how big is the impact?), (2) likelihood (how certain is this factor?), (3) urgency (does this need to be addressed now or later?). This prevents SWOT bloat where the matrix has 20 factors per category and the CEO cannot tell what matters most.

Fourth, it produces strategic implications. The agent does not just list factors. It interprets the pattern: 'Strengths and Opportunities align well, suggesting a growth strategy is viable. However, Weaknesses in technical execution and Threats from regulatory compliance create high execution risk. Recommended strategy: invest in technical hiring (mitigate Weakness) and legal/compliance support (mitigate Threat) before aggressive growth.'

Fifth, it generates priority actions based on the SWOT. For each quadrant, the agent recommends specific actions: (1) Leverage Strengths: 'Founding team has deep domain expertise in academic publishing. Priority action: use domain credibility to secure early adopter universities as design partners.' (2) Mitigate Weaknesses: 'No technical co-founder. Priority action: recruit CTO-level hire or technical co-founder by Month 3.' (3) Capture Opportunities: 'EU universities increasing R&D budgets post-pandemic. Priority action: time launch to align with annual budget cycles (Q4 for calendar-year institutions).' (4) Defend against Threats: 'GDPR compliance adds $30K-50K cost and 3-6 month timeline. Priority action: engage legal counsel by Month 1 to start compliance workstream in parallel with product development.'

Sixth, it checks for strategic misalignment. The agent looks for red flags: (1) Are Strengths relevant to capturing Opportunities? If the business's strength is 'low cost' but the opportunity is 'premium market segment', that is a mismatch. (2) Are Weaknesses likely to be exploited by Threats? If the weakness is 'slow product development' and the threat is 'fast-moving competitors', that is a vulnerability. The agent flags these misalignments so the CEO can adjust strategy.

Seventh, it produces confidence-calibrated assessments. Not all factors are equally certain. If the Opportunity Analyst had low confidence in market size, the SWOT Synthesizer flags 'Opportunity: Large TAM (low confidence, needs validation).' This prevents the CEO from treating assumptions as facts.

Eighth, it outputs in both structured (JSON) and visual (2x2 matrix) formats. The JSON is used by downstream agents. The visual matrix is used by the CEO and external stakeholders. Both are generated in a single agent run to ensure consistency.""",

        "How It Works Technically": """The SWOT Synthesizer is a Claude Haiku-powered agent that runs mid-pipeline, after Sections 1, 3, and 4 complete.

The agent is triggered by the Mother Agent once Sections 1, 3, and 4 outputs are available. The input is the full outputs from those sections plus any CEO-provided strategic context (like 'I want to focus on X market first' or 'I am constrained by Y budget').

The input schema includes: section_1_output (Opportunity Analyst), section_3_output (Environment Research), section_4_output (Organisation Designer), ceo_strategic_priorities (optional list of what the CEO cares most about).

The agent uses Claude Haiku for cost optimization. SWOT synthesis is primarily about categorization and summarization, which Haiku handles well at lower cost than Sonnet.

The system prompt tells the agent: 'You are synthesizing a SWOT analysis from prior sections. Extract the most important strategic factors. Classify them correctly (Strengths/Weaknesses are internal, Opportunities/Threats are external). Prioritize by impact and urgency. Produce strategic implications. Recommend priority actions. Be concise. Each quadrant should have 3-7 factors, not 20.'

The output schema is a Pydantic model with: section_number ('5'), strengths (list of 3-7 internal advantages), weaknesses (list of 3-7 internal limitations), opportunities (list of 3-7 external favorable conditions), threats (list of 3-7 external risks), strategic_implications (paragraph summarizing what the SWOT pattern suggests), priority_actions (list of specific actions organized by SWOT quadrant), strategic_risks (what could go wrong if the CEO ignores key weaknesses or threats?), key_assumptions, critical_uncertainties, confidence_score, assumptions_used, uncertainties, input_tokens, output_tokens.

The agent writes its output to agent_outputs with section_number='5'. The Mother Agent marks Section 5 complete and proceeds to downstream strategy sections (Marketing, Operations, Financial).

Error handling: if any prior section (1, 3, or 4) is missing, the SWOT Synthesizer proceeds with partial data and flags 'SWOT analysis incomplete, Section X unavailable' in uncertainties. If the agent cannot extract enough factors to populate all quadrants (e.g. no clear strengths), it flags that explicitly rather than inventing factors.""",

        "How It Connects to Other Agents": """The SWOT Synthesizer (Section 5) runs mid-pipeline, after foundational research sections (1, 3, 4) and before strategy execution sections (8, 10, 12, 13, 14).

The execution order is: (1) Sections 1, 3, 4 complete. (2) SWOT Synthesizer runs. (3) Marketing, Operations, Financial, Launch, and Exit Strategy agents run afterward (using the SWOT as strategic context).

The SWOT output feeds into multiple downstream agents: (1) Marketing Agent (Section 8) leverages Strengths and Opportunities to design the go-to-market strategy. If a Strength is 'strong founder network in target customer segment', the Marketing Agent prioritizes founder-led sales and warm intros. If an Opportunity is 'low competition in niche vertical', the Marketing Agent focuses on vertical positioning. (2) Operations Agent (Section 10) mitigates Weaknesses flagged in the SWOT. If a Weakness is 'no customer support process', the Operations Agent designs a support workflow. (3) Financial Agent (Section 12) models the impact of Threats on revenue and costs. If a Threat is 'regulatory compliance costs $50K Year 1', the Financial Agent includes that in the expense model. (4) Launch Agent (Section 13) incorporates Threats into the contingency plan. If a Threat is 'competitor may launch similar product', the Launch Agent builds a contingency for faster-than-planned launch.

The SWOT Synthesizer reads from Sections 1, 3, and 4 but does not directly depend on any other sections. It is a synthesis node that bridges foundational research and strategy execution.

The SWOT also feeds into the executive summary (Summary Agent, Section final). The executive summary includes a 'Key Strengths and Risks' section that is pulled directly from the SWOT analysis.

The Devil's Advocate may challenge the SWOT if it is overly optimistic or dismisses threats. If the SWOT lists 6 Strengths, 2 Weaknesses, 5 Opportunities, and 1 Threat, the Devil's Advocate will ask: 'Why so few threats? Is the agent underweighting risks?' This adversarial check prevents rose-colored SWOT matrices.

In Phase 3 (future), the SWOT Synthesizer will re-run periodically (quarterly) to update the strategic assessment as the external environment changes and the business evolves. If a new competitor enters, that becomes a new Threat. If the business hires a key role, that mitigates a Weakness.""",

        "Why This Matters (Real-World Impact)": """The SWOT Synthesizer is the reason the multi-agent system produces strategic clarity, not just analytical depth. Many business plans are data-rich but insight-poor. They have lots of analysis but no clear strategic position. The SWOT Synthesizer forces synthesis and produces a one-page strategic summary that anyone can understand.

In real-world usage, the SWOT matrix is what the CEO shows to co-founders, advisors, and investors to explain the strategic position. It is the standard language of strategy. An investor who reviews the SWOT can immediately assess: Does this business have real strengths, or are the strengths weak? Are the opportunities big and real, or are they vague? Are the threats manageable, or are they fatal?

The priority actions are particularly impactful. Many founders read a SWOT and say 'OK, so what?' They do not know what to do with the information. The SWOT Synthesizer solves that by producing actionable recommendations: 'Your biggest weakness is lack of technical co-founder. Priority action: recruit a CTO-level hire or technical co-founder by Month 3. If you cannot recruit, you will need to outsource MVP development and hire a senior engineer by Month 6.'

The strategic implications help the CEO choose the right strategy. If the SWOT shows strong Strengths and Opportunities but weak Weaknesses and Threats, the right strategy is aggressive growth. If the SWOT shows weak Strengths and strong Threats, the right strategy is cautious, de-risk first. The SWOT Synthesizer makes that strategic logic explicit.

The misalignment detection prevents strategic errors. If the business's strength is 'low-cost solution' but the opportunity is 'premium enterprise market', the SWOT Synthesizer will flag that as a mismatch. The CEO then needs to either reposition (target cost-sensitive customers) or rebuild the strength (increase product quality to justify premium pricing).

The confidence calibration is critical. A SWOT matrix full of high-confidence factors is very different from a SWOT matrix full of low-confidence factors. The SWOT Synthesizer flags confidence levels so the CEO knows which factors are validated versus assumed. This prevents overconfidence in strategic planning.

Finally, the SWOT Synthesizer makes the business plan defensible. When an investor asks 'What are your key risks?', the CEO can point to the Threats quadrant. When an advisor asks 'Why are you positioned to win?', the CEO can point to the Strengths and Opportunities quadrants. The SWOT matrix is the strategic argument for why this business can succeed.""",
    })


def main() -> None:
    """Expand all remaining agent documents."""
    output_dir = Path("/home/saiaditya26122006/multi-agent-system/explaination")

    logger.info("Expanding remaining 8 agent documents with very detailed content...")

    # Batch 1: Organisation Designer and SWOT
    title, sections = expand_organisation_designer()
    create_expanded_doc(title, sections, output_dir / "organisation_designer_explanation.docx")

    title, sections = expand_swot_synthesizer()
    create_expanded_doc(title, sections, output_dir / "swot_synthesizer_explanation.docx")

    logger.info("\n✅ Expanded 2/8 documents. Continuing with tech stack, marketing, operations, financial, launch, exit...")


if __name__ == "__main__":
    main()
