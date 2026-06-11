"""
Generate Word document explanations for each agent.

Creates detailed, conversational explanations of agent code in DOCX format.
"""

import os
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH


def add_styled_heading(doc, text, level=1):
    """Add a heading with custom styling."""
    heading = doc.add_heading(text, level=level)
    heading.style.font.color.rgb = RGBColor(31, 78, 121)
    return heading


def add_styled_paragraph(doc, text, bold=False, italic=False):
    """Add a paragraph with optional styling."""
    para = doc.add_paragraph(text)
    if bold:
        para.runs[0].bold = True
    if italic:
        para.runs[0].italic = True
    para.style.font.size = Pt(11)
    return para


def add_code_block(doc, code_text):
    """Add a code block with monospace font."""
    para = doc.add_paragraph(code_text)
    para.style.font.name = 'Courier New'
    para.style.font.size = Pt(9)
    para.paragraph_format.left_indent = Inches(0.5)
    return para


# Agent explanations
AGENT_EXPLANATIONS = {
    "opportunity_analyst": {
        "title": "Opportunity Analyst Agent",
        "purpose": """
The Opportunity Analyst is the first agent in the pipeline. Think of it as the strategist who looks at a raw business idea and figures out if there's actually a real opportunity here. It's not about being optimistic or pessimistic. It's about being honest.

When someone says "I want to build a SaaS for universities," this agent asks the hard questions: Who exactly is going to buy this? What problem are you solving that isn't already solved? How will you compete? What does success look like in Year 1?
        """,
        "why_written": """
We wrote this agent because too many business plans start with vague ideas. They say things like "target market is anyone who needs better productivity" or "we'll capture 10% market share." Those statements mean nothing without specifics.

This agent forces specificity. It won't let you get away with "competitive advantage is better UX" without explaining why that UX matters to your specific buyer. It won't accept revenue projections without defining who the first 10 customers actually are.
        """,
        "key_features": [
            ("ICP Hypothesis", "Defines the Ideal Customer Profile with actual buyer roles, budget processes, and decision timelines. Not just demographics, but how they actually buy."),
            ("Competitive Strategy", "Requires you to state exactly how you'll win. Not generic advantages, but specific barriers that competitors can't easily copy."),
            ("Year 1 Objectives", "Sets concrete, measurable goals tied to revenue milestones. If you say you'll have 50 customers, it asks how that traces to your sales capacity."),
            ("Evidence Grading", "Every assumption gets labeled: validated (you have data), alex_provided (CEO gave it), agent_inferred (we derived it), or assumed (we're guessing)."),
        ],
        "technical_approach": """
The agent uses the Intelligence Engine, which runs a 4-step reasoning process:

1. Decompose: Break the business idea into components (who, what, why, how)
2. Produce: Generate the initial business analysis
3. Challenge: Run Devil's Advocate to find weak spots
4. Revise: Fix the weaknesses and produce final output

It uses Claude Sonnet (the smarter, more expensive model) because this is a high-stakes section. Getting the opportunity wrong here means the entire rest of the plan is built on sand.
        """,
        "integration": """
This agent feeds into almost every other agent:

Section 3 (Environment Research) needs the competitive_strategy to know what market forces matter.
Section 8 (Marketing) needs the icp_hypothesis to build a go-to-market plan.
Section 12 (Financial) needs the objectives to set Year 1 revenue targets.

If this agent produces weak output (lots of "assumed" labels, vague descriptions), every downstream agent inherits that uncertainty. That's why it has strict quality gates.
        """,
        "why_it_matters": """
Imagine a CEO brings you an idea: "Academic manuscript validation software." Without this agent, you'd jump straight to building financial models. But wait - who actually buys this? Is it individual researchers (B2C)? Is it universities (B2B)? Is it journals (B2B2C)?

This agent forces that question upfront. It says "Your buyer is research-intensive business schools, specifically research deans and doctoral program directors." Now every other agent knows who the customer is.

That specificity cascades. Marketing knows to target academic conferences, not Reddit ads. Financial knows to model 6-12 month sales cycles, not 1-day impulse purchases. Operations knows to plan for enterprise procurement workflows.

One clear ICP hypothesis in Section 1 prevents 10 wrong assumptions in Sections 8-12.
        """
    },

    "financial_modelling": {
        "title": "Financial Modelling Agent",
        "purpose": """
The Financial Modelling agent is where the rubber meets the road. All the strategy and marketing plans and team designs from the earlier sections have to turn into actual numbers: How much money do you need? When do you run out of cash? When do you break even? What's the realistic path to profitability?

This isn't a spreadsheet generator. It's a reality checker.
        """,
        "why_written": """
We built this agent because most financial projections are fantasy. They show hockey-stick growth with no explanation. They project break-even in Month 12 but the underlying assumptions would require 100 sales per month from a 2-person team.

This agent traces every number back to its source. Revenue doesn't just "happen" in the model. It comes from: (price from Section 8) times (volume from Section 8) times (conversion rate from Section 8). If any of those inputs are "assumed" (no evidence), the financial model's confidence score drops to "low."
        """,
        "key_features": [
            ("Three Statement Model", "P&L, Balance Sheet, Cash Flow - all internally consistent. If revenue increases, so do receivables. If you hire, cash flow reflects the salary payments."),
            ("Break-Even Analysis", "Not just 'Month 18' - it shows baseline/optimistic/pessimistic scenarios and explains what assumptions have to hold for each."),
            ("Monte Carlo Simulation", "Runs 1000 simulations with variations in key assumptions. Gives you probability distributions: 'There's a 70% chance you break even between Month 14-22.'"),
            ("Upstream Traceability", "Every line item points back to its source section. Headcount costs come from Section 4. CAC comes from Section 8. No orphaned numbers."),
            ("Confidence Ceiling", "If marketing assumptions are weak, financial confidence is automatically capped at 'medium' even if the math is perfect. Bad inputs = uncertain outputs."),
        ],
        "technical_approach": """
This agent uses multiple specialized components:

Financial Skills: Loads methodology documents from skills/financial/ that define how to build each statement correctly. These are mini-guides on P&L construction, cash flow waterfalls, DCF valuation.

SimPy Simulation: Runs Monte Carlo analysis in simulation/financial_sim.py. Varies price, volume, CAC, churn rate simultaneously and tracks how many scenarios lead to cash-out vs profitability.

Intelligence Engine: Orchestrates the reasoning. But unlike other agents, this one does heavy programmatic validation before sending to the LLM. It checks: Does revenue = price × volume? Do costs add up? Is the balance sheet balanced?

Hypothesis Testing (P1-2): Two-tier validation. Programmatic tier catches math errors instantly (LTV:CAC ratio mismatch, percentage >100%). LLM tier catches semantic issues (funnel math doesn't match sales process, timeline unrealistic).
        """,
        "integration": """
This agent is one of the most connected in the system:

DEPENDS ON:
- Section 1: Objectives (sets Year 1 revenue target)
- Section 4: Headcount plan (determines personnel costs)
- Section 8: Revenue assumptions (price, volume, sales cycle)
- Section 10: Cost structure (fixed costs, variable costs, COGS)

FEEDS INTO:
- Section 13: Launch plan (needs funding requirements, cash runway)
- Section 14: Exit strategy (needs Year 3 ARR for valuation multiples)
- Summary: Uses probability distribution for risk assessment

If upstream sections have weak assumptions, this agent's confidence score reflects that. It's the canary in the coal mine for business model viability.
        """,
        "why_it_matters": """
Here's the difference this agent makes:

WITHOUT IT:
CEO: "We'll break even in 18 months."
Investor: "Based on what?"
CEO: "Um, our projections."

WITH IT:
CEO: "Our baseline model shows break-even at Month 18 with these assumptions: 30 customers at $5K each, CAC of $2K, 3-person team. But there's a 60% probability we actually break even between Month 14-22, with the biggest risk being CAC variance. If CAC hits $3K instead of $2K, break-even pushes to Month 24 and we need an extra $150K funding."

The second answer is honest. It shows you understand your business model and its risks. Investors trust that.

More importantly: YOU need to know this before you start spending money. If the simulation shows a 40% chance of running out of cash before break-even, that's not a financing problem. That's a business model problem. This agent surfaces that upfront.
        """
    },

    "marketing_strategy": {
        "title": "Marketing Strategy Agent",
        "purpose": """
The Marketing Strategy agent builds the complete go-to-market plan. Not just "we'll do content marketing and paid ads." It builds the engine that turns money and effort into customers: What channels? What conversion rates? What does one customer actually cost to acquire? How does revenue scale over 3 years?

Every number it produces has to trace to a mechanism. If it says "Year 1 revenue is $300K," it shows: 60 customers × $5K average deal = $300K. If you ask where 60 customers comes from, it shows: 1200 trials × 5% conversion = 60 customers. If you ask where 1200 trials comes from, it shows: 40K ad impressions at 3% CTR = 1200 trials.

It won't let you handwave the middle steps.
        """,
        "why_written": """
We built this because marketing sections in most business plans are fiction. They say "leverage social media for brand awareness" without naming a platform. They claim "30% year-over-year growth" without explaining what changes to drive that growth.

This agent forces you to build the acquisition machine explicitly:
- What channel brings in leads?
- What's the conversion rate at each funnel stage?
- What's the true all-in cost (CAC)?
- How do unit economics work (LTV vs CAC)?

If you can't answer those questions with evidence or at least reasonable assumptions, your revenue projections are guesses.
        """,
        "key_features": [
            ("CAC-to-Conversion Chain", "CAC isn't a single number. It's: (ad spend / impressions) × (impressions / clicks) × (clicks / trials) × (trials / customers). Every step is explicit."),
            ("Channel Selection with Evidence", "Doesn't accept 'social media marketing' without specifying: which platform, why that ICP is there, what conversion rate you're assuming, what comparable companies achieved."),
            ("Unit Economics (LTV:CAC)", "Calculates lifetime value vs acquisition cost. Enforces the magic ratio: LTV:CAC ≥ 3:1. If ratio is lower, agent escalates or requires written justification."),
            ("Pricing Logic", "Price isn't 'competitive pricing.' It's: pain severity → willingness to pay → competitive alternatives → cost floor. If you price at $5K, agent shows why that's the right number."),
            ("Growth Justification", "If volume grows 100% year-over-year, agent demands the driver: new channel? higher conversion? expanded market? 'Growing market' isn't enough."),
        ],
        "technical_approach": """
This agent uses several specialized capabilities:

Live Market Data: Searches for real pricing benchmarks, CAC payback periods, and competitor data using the search service. For academic SaaS, it looks up "institutional SaaS pricing universities Europe 2025."

Unit Economics Calculator: Programmatic validation in P1-2. Checks: Is LTV:CAC ≥ 3? Is payback period <18 months? Are margins realistic? Catches broken economics before wasting LLM tokens.

Magic Ratio Guardrail: If LTV:CAC <3:1, agent CANNOT proceed without either: (1) fixing the economics, (2) providing an acceptable exception (e.g., land-and-expand model), or (3) escalating to CEO.

Reasoning Framework in SYSTEM_PROMPT: Uses Claude Sonnet with a 1500-line prompt that encodes best practices: CAC chain mechanics, channel evidence requirements, pricing derivation logic, growth driver taxonomy.
        """,
        "integration": """
Marketing is the linchpin between opportunity analysis and financial modeling:

DEPENDS ON:
- Section 1: ICP hypothesis (who to target), competitive strategy (how to position)
- Section 5: SWOT matrix (strengths to leverage, weaknesses to address)

FEEDS INTO:
- Section 12: Revenue assumptions (price, volume, sales cycle directly flow to financial model)
- Section 10: Operations (delivery model must support the promised sales volume)
- Section 14: Exit strategy (Year 3 revenue determines valuation multiples)

If marketing produces weak unit economics, financial model will show cash-out risk. If marketing assumes 6-month sales cycle but ICP suggests 12 months, operations will under-staff.

This agent is the translation layer between strategy and execution.
        """,
        "why_it_matters": """
Let me show you what weak vs strong marketing analysis looks like:

WEAK (what most plans do):
- Target market: Small and medium businesses
- Pricing: Competitive pricing in the $50-200/month range
- Marketing channels: Content marketing, paid ads, partnerships
- Year 1 revenue: $500K

STRONG (what this agent produces):
- Target market: European business schools with 500+ doctoral students (ICP: research deans, €20K-50K annual budget for research tools)
- Pricing: €5,000/year per school (justified by: pain point is research rejection rates costing €30K in wasted time per year, competitor pricing €3K-8K for adjacent tools)
- Marketing channels: (1) Direct outreach at EURAM/AOM conferences (2) LinkedIn ads targeting research deans (3) Referrals from pilot schools. Conversion funnel: 40 conference conversations → 12 demos → 6 pilots → 3 paid (50% trial-to-paid conversion).
- Year 1 revenue: €90K (18 customers × €5K), growing to €450K Year 2 (assumes 5x volume from repeatability in outreach + 2 satisfied pilot references).
- CAC: €2,000 (conference costs €10K for 40 convos, demo costs negligible, pilot costs €5K time → total €15K / 3 customers / 2 founders selling = €2K blended CAC Year 1).
- LTV: €17,500 (€5K annual × 3.5 year average contract length × 80% gross margin = €14K, but initial customers likely churn faster so using €17.5K conservative LTV → 8.75:1 LTV:CAC ratio = healthy).

The second version is 10x longer. But every number is defensible. When an investor asks "How did you get to €450K Year 2?" you don't stammer. You say "18 customers Year 1, each referring 0.5 new customers on average, plus our sales capacity doubles from 2 to 3 reps, and our conversion rate improves from 50% to 60% as we refine the pitch. That gives us 90 customers Year 2 at €5K each."

This agent builds that defensibility.
        """
    },

    "tech_stack_agent": {
        "title": "Tech Stack and Data Privacy Agent",
        "purpose": """
The Tech Stack agent answers a critical question that most business plans ignore until it's too late: What does it actually cost to run this software, and is it legal where you're selling it?

For a European SaaS business, data privacy isn't optional. GDPR violations can kill your business before you sign your first customer. This agent doesn't let you handwave "we'll be GDPR compliant." It asks: Where is your data stored? What encryption do you use? Do you have Data Processing Agreements with vendors? Do you need a Data Protection Officer?

On the cost side, it builds the real tech budget. Not "cloud costs" as a vague line item. It breaks down: compute, database, LLM API calls, third-party services, and ties each cost to expected usage.
        """,
        "why_written": """
We built this because tech costs and compliance risks are consistently underestimated. Founders assume "AWS will be like $200/month" without calculating actual traffic or storage needs. They assume "GDPR compliance is just a privacy policy" without understanding data residency requirements or consent management.

This agent forces reality:

On costs: If your product uses LLM APIs, it calculates: users × sessions/month × tokens/session × $3 per million tokens = actual monthly cost. If that's 40% of revenue, you have a unit economics problem.

On compliance: If you're selling to EU universities and storing data in us-east-1 (Virginia), you're violating GDPR Article 44. This agent flags that as CRITICAL and explains why.
        """,
        "key_features": [
            ("Infrastructure Design with GDPR Awareness", "Chooses cloud provider and region based on: LLM availability (Bedrock in eu-west-1 for Claude), data residency (EU only for GDPR), team expertise (AWS vs GCP vs Azure)."),
            ("Usage-Based Cost Estimation", "Not vague guesses. Actual calculation: 100 users × 10 sessions/month × 50K tokens/session = 50M tokens = $150/month LLM cost. Adds compute, database, storage."),
            ("Compliance Checklist", "Returns specific GDPR requirements met/missing: data residency (where), encryption (AES-256 at rest, TLS 1.3 in transit), user rights (which ones implemented), DPA status (signed or not), DPO (appointed or threshold not met)."),
            ("Tech Cost to Revenue Validation", "Flags if tech cost >30% of revenue as 'unsustainable.' If tech cost exceeds revenue, escalates as FATAL."),
            ("Third-Party Vendor Risk Assessment", "Lists every API/service with its purpose, cost, and GDPR compliance status. Flags US-only vendors (Stripe US, Mixpanel US) as risks for EU data."),
        ],
        "technical_approach": """
This agent uses Claude Haiku (faster, cheaper model) because tech stack decisions are more deterministic than strategy. The reasoning framework is structured:

1. Region Selection: Checks compliance requirements from input. If GDPR in list, forces EU regions. If CCPA only, US regions acceptable.

2. Provider Selection: Cross-references region requirements with LLM availability. Example: Claude via Bedrock only in us-east-1 and eu-west-1. If you need Claude + GDPR, answer is eu-west-1.

3. Cost Estimation: Loads typical SaaS cost structure from skills or training data. Scales to expected usage from Section 4 (headcount) and Section 8 (customer volume).

4. Compliance Validation: Runs through GDPR checklist. For each requirement, checks if it's addressed in the stack. Example: "Encryption at rest" → checks database choice. If RDS, yes. If MongoDB Atlas with default settings, no.

5. Vendor Risk Scoring: For each third-party service, checks: Does vendor have DPA available? Is vendor on EU Adequacy list? Has vendor certified under EU-US Data Privacy Framework?
        """,
        "integration": """
Tech Stack connects to operational and financial planning:

DEPENDS ON:
- Section 1: Business type (determines whether you need LLMs, vector databases, etc.)
- Section 4: Team size and capabilities (3 engineers can self-host; 2 founders need managed services)
- Section 8: Expected usage (customer volume, sessions/month for cost calculation)

FEEDS INTO:
- Section 10: Operations cost structure (tech costs are fixed costs in the financial model)
- Section 12: Financial model (monthly tech cost flows into burn rate and break-even analysis)

If this agent estimates $5K/month tech costs but financial model only budgets $1K, there's a mismatch. The financial agent will catch that and flag it.

If this agent flags GDPR gaps (no DPO appointed, data in US region) but marketing targets EU universities, operations will face procurement blockers when universities audit vendor compliance.
        """,
        "why_it_matters": """
Here's a real scenario this agent prevents:

BAD PATH (without this agent):
- Month 0: Launch SaaS product in Europe using OpenAI API (US-based) and AWS us-east-1 for lower costs.
- Month 3: First enterprise customer (German university) asks for GDPR compliance documentation.
- Month 4: Customer's legal team reviews: data is in US, no DPA with OpenAI, no data residency guarantees. Deal blocked.
- Month 5: Scramble to migrate to EU region. OpenAI doesn't offer EU data residency. Switch to Bedrock Claude in eu-west-1. Migration costs €15K in engineering time.
- Month 6: Costs spike 30% due to higher EU pricing. Unit economics now broken.

GOOD PATH (with this agent):
- Month 0: Agent flags: "Target market is EU universities. GDPR Article 44 requires EU data residency. OpenAI is US-only. Use Claude via Bedrock eu-west-1. Cost delta: 2% higher but legally compliant. Budget accordingly."
- Decision made before launch. No migration. No broken deal. No surprise costs.

The agent's value is in preventing the €15K mistake, not in the €10/month it saves by optimizing instance sizes.

For founders, this agent answers: "Can I legally sell this product in my target market?" and "What will it really cost to operate?" before you build anything.
        """
    },

    "exit_strategy": {
        "title": "Exit Strategy and Contingency Plan Agent",
        "purpose": """
The Exit Strategy agent thinks about the endgame from day one. It's not premature. It's realistic. Every business needs an exit path: acquisition, IPO, bootstrap to profitability, or even wind-down if fundamentals don't work.

This agent designs that path with actual numbers: Who would acquire you? At what valuation multiple? What returns would investors see? How does your cap table evolve through funding rounds? If things go wrong, what are the specific triggers for pivoting or shutting down?

It's the antidote to "we'll figure it out later" thinking.
        """,
        "why_written": """
We built this because exit planning is usually either missing or fantasy. Plans say "IPO in 5 years" for businesses that will do $2M revenue. They show founders keeping 70% equity after three funding rounds (impossible math). They say "pivot if needed" without defining what "needed" means.

This agent enforces reality:

Acquisition: You must name 3-5 plausible acquirers with rationale. Not "strategic interest" but "Company X would buy us because we fill their product gap in European markets and they've acquired similar companies before (citations: Acquisition Y in 2022, Acquisition Z in 2020)."

Cap Table: You must show the math. Raise $2M at $10M post-money = 20% dilution. Existing shareholders dilute proportionally. Can't show founders at 80% post-Series A after raising that much.

Contingency: You must set observable triggers. "If CAC exceeds $500 after 6 months with 50+ trials, pivot to partner channel." Not "if things don't work out."
        """,
        "key_features": [
            ("Realistic Exit Path Selection", "Acquisition for most startups (3-7 years). IPO only if projections show >$100M ARR with strong growth. Bootstrap-to-profitability if no external funding."),
            ("Named Acquirers with Rationale", "Not generic. Specific companies: 'Instructure (Canvas LMS) because they're expanding EU university presence and acquired Portfolium in 2020.' With probability assessment: low/medium/high based on precedent."),
            ("Cap Table Evolution Math", "Shows ownership % at pre-seed, post-seed, post-Series A, exit. Validates dilution is mathematically possible. Founders can't keep 60% after raising 3 rounds."),
            ("Investor Return Calculation", "For each funding round: seed investors own X%, exit valuation is $Y, return multiple is Z. If seed investors get <3x return at baseline exit valuation, flags as 'not venture-scale.'"),
            ("Observable Contingency Triggers", "Not 'pivot if market conditions worsen' but 'if <$50K MRR by Month 18 and <6 months runway, initiate wind-down.' Tied to SimPy P10 scenario (pessimistic case from financial model)."),
        ],
        "technical_approach": """
This agent uses Claude Sonnet (complex reasoning required) with specialized logic:

Exit Path Determination: Looks at Year 3 revenue from Section 12. If <$5M ARR, rules out IPO (minimum viable scale not met). If >$20M ARR with 80%+ growth, considers IPO path. Otherwise defaults to acquisition.

Acquirer Research: Uses live market data searches. For academic SaaS, searches "academic software M&A exit valuations 2025" and "education technology strategic acquirers." Looks for real comparable transactions (who bought what, when, for how much).

Cap Table Math: Takes funding_strategy from input (seed round $1M at $5M post, Series A $5M at $20M post). Calculates dilution at each round. Validates that final ownership % sums to 100% (not 110% because founders forgot to account for ESOP).

Valuation Multiple Logic: Applies industry standards (5-10x ARR for SaaS, 3-5x revenue for marketplaces) to Year 3 or Year 5 projections. Produces exit valuation range. Backcalculates investor returns.

Contingency Integration: Reads probability_distribution from Section 12 (SimPy output). P10 scenario (worst 10% outcomes) becomes the contingency scenario. "If we hit P10: 24-month break-even instead of 18, need extra $150K, pivot at Month 15 to reduce burn."
        """,
        "integration": """
Exit Strategy is the culmination of the entire plan:

DEPENDS ON:
- Section 1: Target market, competitive positioning (determines acquirer landscape)
- Section 8: Revenue trajectory (determines valuation)
- Section 12: Financial projections (Year 3 revenue, profitability, cash runway)
- Section 13: Launch milestones (funding timing, capital needs)

FEEDS INTO:
- Executive Summary: Uses exit valuation and investor returns for final recommendation
- Risk assessment: Contingency scenarios from P10 become summary risks

This section synthesizes everything. If financials show $1.2M ARR by Year 3, it might calculate: "$1.2M × 6x SaaS multiple = $7.2M exit valuation. Seed investors (15% ownership) get $1.08M return on $1M investment = 1.08x (below venture threshold, raises red flag)."

That's honest. Maybe the business is great for bootstrapping but not venture-backable. This agent surfaces that upfront.
        """,
        "why_it_matters": """
Here's what this agent prevents:

NAIVE SCENARIO:
Founder raises seed round with pitch: "We'll be acquired by a big edtech company in 5 years."
Investor: "Which company?"
Founder: "Could be Instructure, D2L, Blackboard, whoever."
Investor: "What's the exit valuation?"
Founder: "Probably $50M+"
Investor: "Based on what revenue?"
Founder: "Our projections show $3M ARR by Year 5."
Investor: "So 17x ARR multiple? For a small B2B SaaS? That's not realistic. Pass."

REALISTIC SCENARIO (this agent's output):
Founder: "Our exit strategy is strategic acquisition, most likely by an LMS vendor expanding EU presence or an academic publisher building SaaS adjacency. We've identified 5 plausible acquirers. Year 3 projection is $1.2M ARR, Year 5 is $4M ARR assuming we execute on our go-to-market plan. Conservative exit valuation is $24M (6x Year 5 ARR), upside is $40M (10x if we hit aggressive targets). Your seed investment of $1M for 15% equity would return $3.6M at baseline ($24M exit) or $6M at upside, so 3.6x-6x multiple over 5 years. That's venture-returnable but not a home run. If we hit our numbers but can't find an acquirer, we're profitable by Year 3 with 40% margins, so we can bootstrap and distribute dividends instead."

Second founder gets the meeting. Why? Because they thought through the exit realistically. They know the numbers, the timeline, the buyers, the alternatives.

This agent builds that clarity. It doesn't promise a unicorn exit. It promises a thoughtful, defensible path to liquidity that matches the business's actual scale and market position.
        """
    }
}


def generate_all_explanations():
    """Generate Word documents for all agents."""
    output_dir = Path("/home/saiaditya26122006/multi-agent-system/explaination")
    output_dir.mkdir(exist_ok=True)

    for agent_key, content in AGENT_EXPLANATIONS.items():
        print(f"Generating explanation for {content['title']}...")

        doc = Document()

        # Title
        title = doc.add_heading(content['title'], 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_paragraph()  # Spacing

        # Purpose section
        add_styled_heading(doc, "What This Agent Does", level=1)
        add_styled_paragraph(doc, content['purpose'].strip())

        doc.add_paragraph()  # Spacing

        # Why written section
        add_styled_heading(doc, "Why We Built This Agent", level=1)
        add_styled_paragraph(doc, content['why_written'].strip())

        doc.add_paragraph()  # Spacing

        # Key features section
        if 'key_features' in content:
            add_styled_heading(doc, "Key Features", level=1)
            for feature_name, feature_desc in content['key_features']:
                para = doc.add_paragraph(style='List Bullet')
                para.add_run(feature_name + ": ").bold = True
                para.add_run(feature_desc)

        doc.add_paragraph()  # Spacing

        # Technical approach section
        if 'technical_approach' in content:
            add_styled_heading(doc, "How It Works Technically", level=1)
            add_styled_paragraph(doc, content['technical_approach'].strip())

        doc.add_paragraph()  # Spacing

        # Integration section
        if 'integration' in content:
            add_styled_heading(doc, "How It Connects to Other Agents", level=1)
            add_styled_paragraph(doc, content['integration'].strip())

        doc.add_paragraph()  # Spacing

        # Why it matters section
        if 'why_it_matters' in content:
            add_styled_heading(doc, "Why This Matters (Real-World Impact)", level=1)
            add_styled_paragraph(doc, content['why_it_matters'].strip())

        # Save document
        filename = f"{agent_key}_explanation.docx"
        filepath = output_dir / filename
        doc.save(str(filepath))
        print(f"  ✅ Saved: {filename}")

    print(f"\n✅ All explanations generated in {output_dir}")
    return len(AGENT_EXPLANATIONS)


if __name__ == "__main__":
    count = generate_all_explanations()
    print(f"\n📄 Generated {count} agent explanation documents")
