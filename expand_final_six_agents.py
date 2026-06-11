"""
Comprehensive expansion of the final 6 agent explanation documents with very detailed content:
- Tech Stack Agent (Section 6.5)
- Marketing Strategy Agent (Section 8)
- Operations Agent (Section 10)
- Financial Modelling Agent (Section 12)
- Launch & Contingency Agent (Section 13)
- Exit Strategy Agent (Section 14)

Each document expanded to 42-43KB with extensive detail across all 5 sections.
"""

import logging
from pathlib import Path
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)


def create_doc(title: str, sections: dict, path: Path) -> None:
    """Create detailed Word document."""
    doc = Document()
    t = doc.add_heading(title, level=1)
    t.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    for sec_title, content in sections.items():
        doc.add_heading(sec_title, level=2)
        for para in content.strip().split("\n\n"):
            if para.strip():
                p = doc.add_paragraph(para.strip())
                p.style.font.size = Pt(11)

    doc.save(path)
    logger.info(f"Expanded {path.name} to {path.stat().st_size} bytes")


def expand_tech_stack(output_dir: Path) -> None:
    """Expand Tech Stack Agent explanation."""
    sections = {
        "What This Agent Does": """The Tech Stack Agent (Section 6.5) is the infrastructure architect in the pipeline. After the Opportunity Analyst validates the business idea, the Environment Research maps regulatory requirements, and the Organisation Designer plans the team, the Tech Stack Agent designs the technical infrastructure needed to build and run the product: cloud provider selection, database architecture, AI/ML stack, third-party APIs, authentication system, and data privacy compliance implementation.

This agent answers critical technical questions: What cloud infrastructure do we use and why? What is the total monthly and annual infrastructure cost? Which AI/ML services or models do we integrate? What database technology fits our data model and scale? How do we handle authentication and user management? How do we ensure GDPR, CCPA, and other data privacy regulations are met with our technical architecture?

The Tech Stack Agent does not write code or design detailed system architecture. It makes the high-level technology choices that define what gets built with what tools. If the agent chooses AWS over Google Cloud, that affects everything downstream: pricing models, regional availability, compliance certifications, and ecosystem integrations.

The agent produces a structured technical design document with: infrastructure architecture (cloud provider, regions, availability zones, rationale for choices), estimated infrastructure costs (broken down by service: compute, storage, database, networking, monitoring), AI/ML stack (if applicable: which LLM APIs, model serving infrastructure, vector databases, fine-tuning strategy), database selection (SQL vs NoSQL, managed vs self-hosted, replication and backup strategy), third-party API integrations (payment processing, email delivery, analytics, customer support, each with cost estimates and vendor lock-in assessment), authentication and authorization system (OAuth, SSO, MFA, role-based access control), data privacy and compliance architecture (GDPR, CCPA, DPDP compliance: data residency, encryption at rest and in transit, audit logging, data retention and deletion, consent management), total monthly and annual tech costs, tech risk assessment (vendor lock-in, single points of failure, scalability bottlenecks, compliance gaps), key assumptions about technical requirements, critical uncertainties, and confidence score.

The output is designed to give the CEO a realistic technical plan and cost estimate before any code is written. If the CEO assumed the product could be built for $200/month in infrastructure costs, but the Tech Stack Agent calculates $2,500/month due to database costs, AI API usage, and compliance requirements, that is critical information that changes the financial model.

The Tech Stack Agent also ensures compliance with regulatory requirements identified in Section 3 (Environment Research). If GDPR applies, the agent designs the tech stack with data residency in EU, encryption, and audit logging. If HIPAA applies (healthtech), the agent ensures the cloud provider is HIPAA-compliant and designs the architecture to meet HIPAA technical safeguards.""",

        "Why We Built This Agent": """We built the Tech Stack Agent because first-time founders, especially non-technical founders, drastically underestimate infrastructure costs and complexity. They think building a SaaS product means spinning up a $10/month server and writing some code. They do not think about databases, backups, monitoring, security, compliance, AI API costs, third-party integrations, or scaling.

In early designs of the system, we assumed the Organisation Designer's cost model would include infrastructure costs as a line item. We quickly realized that was insufficient. Infrastructure costs are highly variable depending on architecture choices (managed databases vs self-hosted, serverless vs containers, US region vs EU region for GDPR). The Financial Agent's cost model depends on accurate infrastructure cost estimates, which requires an agent that understands technical architecture.

The Tech Stack Agent solves the technical cost estimation problem. It takes the business requirements (product description, scale, compliance needs) and produces a realistic infrastructure architecture with bottom-up cost estimates. This prevents the Financial Agent from projecting profitability based on fantasy infrastructure costs.

We also built it because technology choices have long-term consequences. Choosing vendor A over vendor B because A is cheaper today might lead to expensive migration costs later if A does not scale or does not meet compliance requirements. The Tech Stack Agent evaluates technology choices with a multi-year lens, not just 'what is cheapest right now?'

The agent also ensures compliance is designed in from the start, not bolted on later. Many startups build their product and then discover that achieving GDPR or HIPAA compliance requires re-architecting the entire system, which costs 6-12 months and $100K-500K. The Tech Stack Agent identifies compliance requirements upfront (from Section 3 Environment Research) and designs the architecture to meet them from day one.

We built it to make technical decisions transparent to non-technical founders. A non-technical CEO does not know what AWS Lambda is, why Postgres might be better than MongoDB for their use case, or why EU data residency adds 15-20% to infrastructure costs. The Tech Stack Agent explains each technology choice in plain language with clear rationale, so the CEO understands what is being built and why.

Finally, we built it because infrastructure costs scale non-linearly, and founders need to plan for that. A product that costs $500/month to serve 100 users might cost $5,000/month to serve 1,000 users, or $50,000/month to serve 10,000 users, depending on architecture. The Tech Stack Agent models those scaling dynamics so the CEO knows what infrastructure costs look like at Year 1, Year 2, and Year 3 volumes.""",

        "Key Features": """First, it performs architecture design based on business requirements, not generic templates. The agent reads the product description, target customer, scale projections (from Section 12 Financial), and compliance requirements (from Section 3 Environment Research), and designs an architecture that fits those requirements. If the product is a B2B SaaS tool with 100-500 customers Year 1, the agent recommends a simple managed database and serverless compute. If the product is a B2C marketplace with 10K-100K users Year 1, the agent recommends a more robust setup with load balancing, caching, and horizontal scaling.

Second, it estimates infrastructure costs with service-level granularity. The agent breaks costs down by: (1) Compute (EC2 instances, Lambda invocations, or equivalent), (2) Database (RDS, Aurora, DynamoDB, or equivalent), (3) Storage (S3, object storage for files/images), (4) Networking (data transfer, load balancer, CDN), (5) AI/ML APIs (if applicable: OpenAI, Anthropic, Google Vertex AI per-token costs), (6) Third-party SaaS (Stripe, SendGrid, Segment, Intercom), (7) Monitoring and logging (CloudWatch, Datadog, Sentry), (8) Security and compliance (WAF, audit logging, encryption key management). For each service, the agent estimates monthly and annual costs based on projected usage volume.

Third, it evaluates data residency and compliance architecture. If the business targets EU customers and GDPR applies, the agent recommends: (1) All infrastructure in EU regions (AWS eu-west-1, Google europe-west1, Azure westeurope), (2) No data transfer to non-EU regions without explicit consent, (3) Encryption at rest (AES-256) and in transit (TLS 1.3), (4) Audit logging of all data access, (5) Data retention and deletion workflows, (6) Consent management system. The agent calculates the cost premium for EU residency vs US (typically 10-20% higher cloud costs) and explains the trade-off: higher cost but eliminates regulatory risk.

Fourth, it designs AI/ML stack based on product requirements. If the product includes LLM-powered features, the agent specifies: (1) Which LLM API (OpenAI GPT-4, Anthropic Claude, Google Gemini, open-source via HuggingFace), (2) Expected token usage per user per month, (3) Estimated monthly API costs (at scale), (4) Latency and availability requirements, (5) Fallback strategy if primary API fails, (6) Whether fine-tuning is needed (and cost of that), (7) Whether embedding/vector search is needed (and vector database choice: Pinecone, Weaviate, Qdrant, pgvector). The agent models API costs at multiple scale points: 100 users, 1K users, 10K users, so the CEO understands how costs scale.

Fifth, it evaluates vendor lock-in risks. For each technology choice, the agent assesses: (1) How easy is it to migrate away from this vendor? (2) Are there open standards or is this proprietary? (3) What is the migration cost if we outgrow this vendor or they raise prices 10x? (4) Is there a second-choice vendor we could switch to? The agent flags high lock-in risks (like AWS proprietary services that do not have equivalents elsewhere) and recommends mitigation strategies (use open-source alternatives where possible, or accept lock-in if switching cost is worth the operational benefits).

Sixth, it identifies single points of failure and designs redundancy. The agent checks: (1) Is the database replicated across availability zones? (2) Is there a backup strategy and tested restore process? (3) Are there multiple regions for disaster recovery? (4) If a third-party API goes down, does the product break completely or degrade gracefully? For each single point of failure, the agent recommends mitigation: multi-AZ database, automated backups, circuit breakers for external APIs, monitoring and alerts for downtime.

Seventh, it produces confidence-calibrated cost estimates. Infrastructure costs are uncertain, especially at scale. The agent provides: (1) Baseline estimate (most likely cost), (2) Optimistic estimate (if usage is lower than expected), (3) Pessimistic estimate (if usage is higher than expected), (4) Cost sensitivity analysis (what drives costs most: database, AI API, data transfer?). This range helps the Financial Agent model optimistic/pessimistic scenarios.

Eighth, it integrates with the team plan from Section 4. The agent checks: does the team have the skills to manage this tech stack? If the tech stack includes Kubernetes and the team has 2 junior engineers with no DevOps experience, the agent flags 'Tech risk: team lacks Kubernetes expertise, recommend managed service (AWS ECS/Fargate or Google Cloud Run) instead, or hire DevOps engineer by Month 3.'""",

        "How It Works Technically": """The Tech Stack Agent is a Claude Haiku-powered agent that runs mid-pipeline, after Sections 1, 3, and 4 complete. It uses Haiku for cost optimization since infrastructure design is more about selection and estimation than deep creative reasoning.

The agent is triggered by the Mother Agent once foundational sections are available. The input includes: product description, target customer type (B2B/B2C), projected scale (from Section 1 and Section 12 if available, or estimated), compliance requirements (from Section 3: GDPR, HIPAA, AI Act, etc.), team technical capabilities (from Section 4: do they have ML engineers, DevOps engineers, or just full-stack engineers?), and budget constraints (if the CEO has flagged infrastructure budget limits).

The input schema includes: business_type (B2B SaaS, B2C marketplace, etc.), product_description, target_geography (affects data residency requirements), compliance_requirements (list of regulations that apply), projected_users_year1, projected_users_year3, team_capabilities (from Section 4), budget_constraints (optional).

The system prompt tells the agent: 'You are designing the technical infrastructure for a startup. Choose technologies that are proven, cost-effective, and match the team's skill level. Prioritize managed services over self-hosted to reduce operational burden for small teams. Ensure compliance with regulations identified in Section 3. Estimate costs at Year 1 and Year 3 scale. Flag vendor lock-in risks. Explain trade-offs clearly. Be realistic about costs. Do not under-estimate infrastructure costs. Include AI API costs if the product uses LLMs. Include third-party SaaS costs for essential services like payments, email, analytics.'

The agent uses a decision tree for cloud provider selection: (1) If EU data residency required and no US presence: AWS eu-west-1 or Google Cloud europe-west1 (both have strong GDPR compliance certifications). (2) If US-only: AWS us-east-1 or Google Cloud us-central1 (cheaper than EU). (3) If multi-region: AWS or Google Cloud (Azure less commonly used by startups). (4) If extreme cost sensitivity: consider cheaper providers (DigitalOcean, Hetzner, Vultr) but flag that they have fewer managed services and weaker compliance certifications.

The agent uses a cost estimation model with bottom-up calculations: (1) Compute: estimate instance type and count based on projected load. For a B2B SaaS with 100-500 users Year 1, typically 1-2 t3.medium instances ($30-60/month). For B2C with 10K users, 5-10 instances ($150-300/month). (2) Database: estimate database size (GB) and query load (IOPS). For B2B SaaS, typically RDS t3.micro with 20-50 GB ($20-40/month Year 1). For B2C, larger instance with read replicas ($100-300/month Year 1). (3) AI API costs: if using Claude API, estimate tokens per user per month. For a product with light usage (1-2 queries per user per day, 1K tokens per query), 100 users = 6M tokens/month = $18-30/month. 1K users = $180-300/month. 10K users = $1,800-3,000/month. (4) Third-party SaaS: Stripe (2.9% + 30c per transaction), SendGrid ($15-100/month depending on volume), Segment ($120/month), Intercom ($74/month starter). Sum all services.

The output schema is a Pydantic model with: section_number ('6.5'), infrastructure (cloud_provider, primary_region, architecture_description, rationale, monthly_cost, annual_cost), ai_ml_stack (if applicable: primary_service, model, token_cost_estimate, future_expansion), database (type, provider, replication_strategy, backup_strategy, monthly_cost), third_party_apis (list of integrations, each with provider, purpose, monthly_cost, lock_in_risk), authentication (approach, MFA, SSO, rationale), data_privacy_compliance (regulations_covered, data_residency, encryption, audit_logging, consent_management, estimated_compliance_cost), total_tech_cost_monthly, total_tech_cost_annual, tech_risk_assessment (list of risks: vendor lock-in, single points of failure, scalability concerns, team skill gaps), key_assumptions, critical_uncertainties, confidence_score, assumptions_used, uncertainties, input_tokens, output_tokens.

The agent writes its output to agent_outputs with section_number='6.5'. The Mother Agent marks Section 6.5 complete and passes the tech cost estimates to the Financial Agent (Section 12) to incorporate into the financial model.

Error handling: if compliance requirements are unclear (e.g. the business targets a regulated industry but Section 3 did not identify specific regulations), the agent proceeds with general best practices (encryption, audit logging) and flags 'Compliance architecture incomplete, manual legal and technical review required' in uncertainties.""",

        "How It Connects to Other Agents": """The Tech Stack Agent (Section 6.5) runs mid-pipeline, after Sections 1, 3, and 4, and before Section 12 (Financial). It bridges the business requirements (what needs to be built) and the financial model (how much it costs to build and run).

The execution order is: (1) Sections 1, 3, 4 complete. (2) Tech Stack Agent runs. (3) Section 12 (Financial) runs using tech cost estimates from Section 6.5.

The Tech Stack output feeds into: (1) Financial Agent (Section 12) uses the total_tech_cost_monthly and total_tech_cost_annual estimates as input to the cost model. The Financial Agent's P&L includes a line item 'Infrastructure and Technology Costs' that comes directly from Section 6.5. If the Tech Stack Agent underestimates costs, the Financial Agent projects false profitability. (2) Operations Agent (Section 10) uses the tech stack design to plan operational workflows. If the tech stack includes a managed database with automated backups, the Operations Agent does not need to plan manual backup procedures. If the tech stack includes Stripe for payments, the Operations Agent designs the billing workflow around Stripe's APIs and webhooks. (3) Launch Agent (Section 13) uses the tech stack to plan deployment and go-live. If the tech stack is serverless (AWS Lambda), the launch plan is simpler than if the tech stack requires managing Kubernetes clusters.

The Tech Stack Agent reads from: (1) Section 3 (Environment Research) to identify compliance requirements. If Section 3 flagged GDPR, the Tech Stack Agent ensures EU data residency. If Section 3 flagged HIPAA, the Tech Stack Agent ensures HIPAA-compliant infrastructure. (2) Section 4 (Organisation Designer) to check team capabilities. If the team has no DevOps expertise, the Tech Stack Agent recommends managed services over self-hosted. (3) Section 1 (Opportunity Analyst) to understand product complexity and scale projections.

The Tech Stack Agent does not interact with the Marketing Agent (Section 8) directly, but the tech stack can affect marketing strategy. If the tech stack includes expensive AI API costs that only make sense at scale, the Marketing Agent needs to plan for aggressive user acquisition to reach unit economics. If the tech stack is low-cost and scales linearly, the Marketing Agent has more flexibility.

The Devil's Advocate may challenge the Tech Stack Agent's cost estimates. If the agent assumes $500/month infrastructure costs but does not account for data transfer costs, backup storage costs, or monitoring costs, the Devil's Advocate will flag that. If the agent recommends a proprietary vendor with high lock-in risk and no fallback, the Devil's Advocate will challenge that choice.

In Phase 3 (future), the Tech Stack Agent will re-run as the product scales. When the business reaches 1K users, 10K users, 100K users, the agent re-evaluates the architecture and recommends optimizations or migrations (e.g. move from managed database to self-hosted for cost savings, add caching layer to reduce database load, move to multi-region for latency and availability).""",

        "Why This Matters (Real-World Impact)": """The Tech Stack Agent is the reason the multi-agent system produces business plans with realistic technical and financial projections. Many startup failures are not due to bad ideas or bad teams, but due to underestimating infrastructure costs, choosing the wrong technology and getting locked in, or failing to meet compliance requirements.

In real-world usage, the Tech Stack Agent prevents the CEO from launching with a technical architecture that does not scale or does not comply with regulations. If the CEO plans to target EU universities and needs GDPR compliance, but their initial tech stack is US-only with no EU data residency, they will face a $50K-100K re-architecture cost and 6-12 month delay when they try to sell to EU customers. The Tech Stack Agent surfaces that upfront and designs the architecture correctly from day one.

The cost estimates are particularly impactful. First-time founders often assume infrastructure is cheap ('it is just a few servers'). The Tech Stack Agent shows the real cost: $500/month Year 1 becomes $2,500/month Year 2 as usage grows, and $10,000/month Year 3 if AI API usage is heavy. If the Financial Agent projects break-even in 18 months based on $200/month infrastructure costs, but the real cost is $2,000/month, the business will run out of money before reaching profitability. The Tech Stack Agent prevents that.

The compliance architecture design prevents catastrophic regulatory failures. Many startups build products and later discover they cannot sell to their target market because they do not meet compliance requirements. A healthtech startup without HIPAA compliance cannot sell to US hospitals. A SaaS startup without GDPR compliance cannot sell to EU enterprises. The Tech Stack Agent identifies compliance requirements from Section 3 and designs the architecture to meet them, so the CEO does not waste 6-12 months building a product they cannot legally sell.

The vendor lock-in assessment helps the CEO make informed technology choices. Choosing AWS Lambda over Google Cloud Functions might seem like a minor choice, but if the startup later wants to migrate (due to cost, reliability, or acquisition considerations), the migration cost can be $100K-500K in engineering time. The Tech Stack Agent flags high lock-in risks and recommends mitigation strategies (use open-source alternatives, or accept lock-in if the operational benefits are worth it).

The scalability planning prevents technical debt. Many startups choose the cheapest simplest architecture for MVP, then hit a scaling wall at 1K-10K users and need to re-architect. The Tech Stack Agent designs the architecture to scale from 100 users to 10K users without a complete rewrite. This saves 3-6 months of engineering time that would otherwise be spent on scaling rewrites.

Finally, the Tech Stack Agent makes technical decisions transparent to non-technical stakeholders. When the CEO pitches to an investor and the investor asks 'how will you handle data privacy?', the CEO can point to Section 6.5, which explains the GDPR-compliant architecture with EU data residency, encryption, and audit logging. That level of technical rigor builds investor confidence that the founding team understands what they are building.""",
    }

    create_doc("Tech Stack Agent - Infrastructure and Compliance Architecture", sections, output_dir / "tech_stack_agent_explanation.docx")


def expand_marketing(output_dir: Path) -> None:
    """Expand Marketing Strategy Agent explanation."""
    sections = {
        "What This Agent Does": """The Marketing Strategy Agent (Section 8) is the customer acquisition architect in the pipeline. After the business model is validated, the market is researched, the team is planned, and the tech stack is designed, the Marketing Agent figures out how to actually acquire customers: target customer profile, channel strategy, customer acquisition cost (CAC), lifetime value (LTV), funnel conversion rates, messaging and positioning, and go-to-market timeline.

This agent answers critical commercial questions: Who exactly is the target customer (not just 'universities', but 'PhD students at research-intensive EU universities writing journal submissions')? What marketing channels will reach them (paid ads, content marketing, partnerships, direct sales, events)? How much does it cost to acquire one customer (CAC)? How much revenue does one customer generate over their lifetime (LTV)? What is the LTV:CAC ratio, and is it viable (need 3:1 or better for sustainable business)? What is the conversion funnel (awareness to trial to paid to retained)? What is the core messaging and competitive positioning?

The Marketing Agent does not create actual marketing materials (ads, landing pages, blog posts). It creates the marketing strategy that defines what to build and what to say. If the agent determines the best channel is founder-led sales to warm introductions, that is very different from performance marketing with paid ads. The Marketing Agent makes that strategic choice and backs it with reasoning.

The agent produces a structured go-to-market plan with: target customer profile (ICP: ideal customer profile with demographics, psychographics, pain points, buying behavior, budget authority), channel strategy (for each channel: paid ads, SEO, content, social, partnerships, direct sales, events, conferences - viability assessment, estimated CAC, timeline to first customers), customer acquisition cost (CAC) estimate with breakdown by channel, lifetime value (LTV) calculation (average revenue per customer over retention period), LTV:CAC ratio analysis (is unit economics viable?), funnel conversion model (top of funnel to awareness to trial to paid to retained), messaging and positioning (core value prop, differentiation from competitors, positioning statement), go-to-market timeline (Month 1-12 milestones: when to launch which channels, when to expect first paying customers), marketing budget estimate (Year 1, Year 2, Year 3), marketing risk assessment (what if CAC is 2x higher than estimated, what if conversion rate is 50% lower, what if primary channel does not work), key assumptions about customer acquisition, critical uncertainties, and confidence score.

The output is designed to give the CEO a realistic view of how customers will be acquired and whether the unit economics work. If the Marketing Agent calculates CAC of $800 and LTV of $2,000, the LTV:CAC ratio is 2.5:1, which is marginal (need 3:1 or better). The CEO then knows customer acquisition is the highest-risk area and needs to be validated with real pilots before scaling.

The Marketing Agent also feeds directly into the Financial Agent (Section 12). The revenue model depends on the customer volume projections from the Marketing Agent. If the Marketing Agent estimates 50 customers Year 1, but the Financial Agent assumes 100 customers Year 1, that inconsistency will be caught.""",

        "Why We Built This Agent": """We built the Marketing Strategy Agent because most first-time founders drastically underestimate customer acquisition difficulty and cost. They assume 'if we build it, they will come' or 'we will just run some Facebook ads.' They do not think about CAC, LTV, channel viability, funnel conversion rates, or whether the unit economics actually work.

In early designs of the system, we assumed the Opportunity Analyst's market size analysis was sufficient for commercial planning. We quickly realized that was insufficient. Knowing the market is $5B does not tell you how to acquire your first 100 customers, what it costs, or whether the economics make sense. The Marketing Agent solves the customer acquisition strategy problem.

We also built it because customer acquisition is the #1 startup killer. More startups fail because they cannot acquire customers cost-effectively than because the product is bad. A great product with a CAC of $2,000 and LTV of $1,500 is a dead business. A mediocre product with a CAC of $300 and LTV of $1,200 is a viable business. The Marketing Agent forces the CEO to confront the unit economics before launch, not after burning $50K-100K on failed customer acquisition experiments.

The agent also educates the CEO about realistic CAC and conversion rates. First-time founders often assume unrealistic funnel math: '1% conversion is standard, so if we get 10,000 website visitors, we will get 100 customers.' Real B2B SaaS conversion rates from website visitor to paid customer are typically 0.1-0.5%, not 1-2%. The Marketing Agent provides realistic conversion benchmarks based on business model, customer type, and pricing, so the CEO builds the revenue model on realistic assumptions.

We built it to make marketing strategy match the business model. If the product is high-touch enterprise SaaS ($20K-50K ACV), the right marketing strategy is direct sales with warm intros, not performance marketing ads. If the product is low-touch self-serve SaaS ($50-200/month), the right strategy is performance marketing, SEO, and PLG (product-led growth). The Marketing Agent tailors the strategy to the business model, not one-size-fits-all.

The agent also prevents premature scaling. Many founders start running paid ads on Day 1, before validating messaging, before understanding the customer, before having any conversion data. They burn $10K-50K with zero ROI. The Marketing Agent produces a phased go-to-market plan: Month 1-3 is customer discovery and messaging validation with 5-10 pilot customers via warm intros. Month 4-6 is testing channels with small budgets to measure CAC. Month 7-12 is scaling the channels that work. This prevents wasted spend on unvalidated marketing.

Finally, we built it because marketing strategy needs to be consistent with financial projections. If the Financial Agent projects $500K revenue Year 1, that implies a specific number of customers at a specific price point. The Marketing Agent reverse-engineers that: 'To reach $500K revenue with $5K ACV, you need 100 customers. To get 100 customers with 5% trial-to-paid conversion, you need 2,000 trials. To get 2,000 trials with 2% website-to-trial conversion, you need 100,000 website visitors. Can you realistically drive 100K visitors Year 1?' If the answer is no, the revenue target is fantasy.""",

        "Key Features": """First, it defines the Ideal Customer Profile (ICP) with specificity. The agent does not say 'universities'. It says 'PhD students at research-intensive universities in EU (specifically Germany, UK, Spain, France) writing journal submissions in social sciences and humanities, in the final year of their PhD program, experiencing high desk-rejection rates, supervised by advisors who value publication output, with personal or grant budget authority for $50-200 software tools, active on academic Twitter and attending 2-3 conferences per year.' This level of specificity enables targeted marketing.

Second, it evaluates multiple marketing channels with viability scoring. For each channel (paid ads, SEO, content marketing, social media, partnerships, direct sales, events, conferences, PR, influencer marketing, community building, product-led growth), the agent assesses: (1) Viability: can this channel realistically reach the ICP? (2) CAC estimate: what does it cost to acquire one customer via this channel? (3) Timeline to first customer: how long until this channel generates the first paying customer? (4) Scale potential: can this channel scale to 100 customers, 1000 customers, 10K customers? (5) Control: do we control this channel or depend on platforms/partners? The agent ranks channels by viability and recommends a prioritized list.

Third, it estimates CAC with channel-specific breakdown. The agent calculates: (1) Paid ads CAC = (CPC * clicks per customer) + (creative cost / customers from that creative). For B2B SaaS, typical $2-10 CPC, 100-500 clicks per customer, CAC $200-2000. (2) Content marketing CAC = (content production cost + distribution cost) / customers from content. Typically $50-500 depending on virality. (3) Direct sales CAC = (sales salary + tools + travel) / customers closed. For B2B enterprise, typically $5K-50K per customer. (4) Partnerships CAC = (partnership dev cost + revenue share or referral fee) / customers from partners. The agent produces a blended CAC estimate across all channels.

Fourth, it calculates LTV with retention modeling. LTV = (Average Revenue Per Customer) * (Retention Period) - (Cost to Serve). For a $5K/year SaaS product with 70% annual retention, LTV = $5K * (1/0.3) = $16.7K over ~3.3 years. If churn is higher (50% annual retention), LTV drops to $10K over 2 years. The agent models LTV at different retention rates to show the CEO how retention affects unit economics. The agent also includes cost to serve (customer support, account management, infrastructure per user), which reduces net LTV by 10-30%.

Fifth, it produces LTV:CAC ratio analysis with viability threshold. The agent calculates LTV:CAC ratio and flags: (1) <1:1 = not viable, burning money on every customer. (2) 1:1 to 2:1 = marginal, only works if payback period is very short (3-6 months). (3) 2:1 to 3:1 = viable but tight, need to improve either LTV (increase price, reduce churn, upsell) or CAC (better targeting, better conversion). (4) 3:1+ = good unit economics, can scale profitably. (5) 5:1+ = great unit economics, rare but achievable with strong product-market fit and efficient go-to-market.

Sixth, it models the conversion funnel with realistic benchmarks. The agent builds a multi-stage funnel: (1) Awareness (how many people become aware of the product via marketing), (2) Consideration (how many visit the website or engage), (3) Trial (how many sign up for free trial or demo), (4) Paid (how many convert to paying customers), (5) Retained (how many stay customers after 3 months, 6 months, 12 months). For each stage, the agent provides conversion rate benchmarks based on industry data: B2B SaaS website-to-trial is 2-5%, trial-to-paid is 10-25%, annual retention is 60-80%. The agent models optimistic/baseline/pessimistic scenarios to show sensitivity.

Seventh, it develops messaging and positioning based on SWOT and competitive landscape. The agent reads Section 1 (Opportunity: competitive differentiation) and Section 5 (SWOT: strengths) to craft the core value prop and positioning statement. If the strength is 'only tool that validates epistemic claims, not just formatting', the positioning is 'The only manuscript validation tool that checks your reasoning, not just your references.' The agent produces a positioning statement format: 'For [target customer], who [pain point], our product is [category] that [unique differentiation]. Unlike [alternatives], we [key benefit].'

Eighth, it produces a phased go-to-market timeline. The agent does not say 'launch all channels on Day 1.' It produces a phased plan: (1) Month 0-3: Customer discovery via warm intros, 5-10 design partner customers, validate messaging and pain points, no paid marketing. (2) Month 3-6: Launch website, run small-budget experiments on 2-3 channels ($500-2K/month), measure CAC and conversion, iterate messaging. (3) Month 6-12: Scale the 1-2 channels that work, increase budget to $5K-10K/month, aim for 30-50 customers by Month 12. This phased approach prevents premature scaling and capital waste.""",

        "How It Works Technically": """The Marketing Strategy Agent is a Claude Sonnet-powered agent (requires strategic thinking and market sense) that runs mid-late pipeline, after Sections 1, 3, 4, 5, and sometimes 6.5 (though 6.5 can run in parallel).

The agent is triggered by the Mother Agent once foundational sections are available. The input includes: target customer segment (from Section 1), competitive landscape (from Section 1), SWOT analysis (from Section 5), external market trends (from Section 3), team capabilities (from Section 4: is there a marketing/sales hire or is this founder-led?), product pricing (if known, or estimated), and any CEO-provided marketing constraints or preferences.

The input schema includes: business_model (B2B/B2C, SaaS/marketplace/etc), target_customer_segment, product_pricing_estimate (ACV or monthly subscription), competitive_landscape (from Section 1), swot_strengths (from Section 5), team_marketing_capability (from Section 4), geography (affects channel availability and costs), stage (pre-launch, MVP, scaling).

The agent uses Claude Sonnet (not Haiku) because marketing strategy requires creative strategic thinking, understanding customer psychology, and evaluating trade-offs across many possible channels. Sonnet significantly outperforms Haiku on this type of open-ended strategic reasoning.

The system prompt tells the agent: 'You are designing the go-to-market strategy for a startup. Define the ideal customer profile with specificity. Evaluate marketing channels realistically (paid ads, SEO, content, partnerships, direct sales). Estimate CAC and LTV based on business model and pricing. Calculate LTV:CAC ratio and flag if unit economics are not viable. Model the conversion funnel with realistic benchmarks (B2B SaaS trial-to-paid is typically 10-25%, not 50%). Prioritize channels that match the business model (B2B high-touch = direct sales, B2B low-touch = PLG + paid ads, B2C = performance marketing + virality). Produce a phased go-to-market timeline. Be realistic about customer acquisition difficulty. Do not assume viral growth or word-of-mouth without evidence. Flag if CAC is too high relative to LTV.'

The agent uses industry benchmarks for CAC and conversion rates, parameterized by business model: (1) B2B SaaS $100-10K ACV: CAC $200-2K, website-to-trial 2-5%, trial-to-paid 10-25%, annual retention 70-85%, payback period 6-18 months. (2) B2B SaaS $10K-100K ACV: CAC $5K-50K, direct sales driven, close rate 10-30%, annual retention 80-95%, payback period 12-24 months. (3) B2C freemium SaaS: CAC $10-100, website-to-signup 5-15%, free-to-paid 1-5%, annual retention 40-60%, payback period 3-12 months. The agent selects the appropriate benchmark based on business model.

The output schema is a Pydantic model with: section_number ('8'), ideal_customer_profile (demographics, psychographics, pain_points, buying_behavior, budget_authority), channel_strategy (list of channels, each with viability_score, estimated_cac, timeline_to_first_customer, scale_potential, priority_rank), estimated_cac_blended (weighted average across channels), estimated_ltv (calculation breakdown), ltv_cac_ratio (value and viability assessment), funnel_conversion_model (awareness -> consideration -> trial -> paid -> retained, with conversion rates at each stage), messaging_and_positioning (core value prop, differentiation, positioning statement, 3-5 key messages), go_to_market_timeline (Month 1-12 milestones), marketing_budget_estimate (Year 1, Year 2, Year 3 monthly spend by channel), marketing_risk_assessment (what if CAC is 2x, what if conversion is 50% lower, what if primary channel fails), key_assumptions, critical_uncertainties, confidence_score, assumptions_used, uncertainties, input_tokens, output_tokens.

The agent writes its output to agent_outputs with section_number='8'. The Mother Agent marks Section 8 complete and passes the customer volume projections and CAC/LTV estimates to the Financial Agent (Section 12) to incorporate into the revenue and cost model.

Error handling: if the agent cannot estimate CAC with confidence (because the business model is novel or the customer segment is niche), it provides a range (CAC $200-2000) and flags 'CAC estimate has high uncertainty, range spans 10x, validate with pilot campaigns' in uncertainties and proceeds with baseline estimate.""",

        "How It Connects to Other Agents": """The Marketing Strategy Agent (Section 8) runs mid-late pipeline, after Sections 1, 3, 4, 5 and before Section 12 (Financial). It bridges the market opportunity (what could be sold) and the financial model (what revenue is realistic).

The execution order is: (1) Sections 1, 3, 4, 5 complete. (2) Marketing Agent runs. (3) Financial Agent (Section 12) runs using customer volume projections and CAC from Section 8.

The Marketing output feeds into: (1) Financial Agent (Section 12) uses the customer acquisition projections (how many customers Year 1, Year 2, Year 3) to build the revenue model. The Financial Agent's P&L revenue line is directly derived from Section 8's customer volume * pricing. The Financial Agent also uses the marketing budget estimates and CAC as input to the cost model. If the Marketing Agent estimates $5K/month marketing spend Year 1, the Financial Agent includes that in opex. (2) Operations Agent (Section 10) uses the go-to-market channels to plan operational workflows. If the Marketing Agent recommends direct sales, the Operations Agent designs a sales CRM workflow, outreach sequences, and demo scheduling process. If the Marketing Agent recommends self-serve PLG, the Operations Agent designs an automated onboarding flow.

The Marketing Agent reads from: (1) Section 1 (Opportunity) for competitive landscape, market size, and customer segment definition. (2) Section 3 (Environment Research) for external trends that affect channels (e.g. if there is a trend toward decreased ad effectiveness, the Marketing Agent deprioritizes paid ads). (3) Section 4 (Organisation Designer) to check if there is a marketing/sales hire in the team plan. If not, the marketing strategy must be founder-led, which limits channel options. (4) Section 5 (SWOT) for strengths to leverage in positioning and opportunities to capture in go-to-market timing.

The Marketing Agent does not interact with the Tech Stack Agent directly, but tech stack can affect marketing strategy. If the tech stack includes expensive AI API costs that only make sense at scale (>1K users), the Marketing Agent knows the business needs aggressive growth to reach unit economics, not slow steady growth.

The Devil's Advocate will aggressively challenge the Marketing Agent's CAC and LTV estimates. If the Marketing Agent assumes CAC of $400 with no pilot data, the Devil's Advocate will ask: 'Based on what evidence? B2B SaaS in this space typically has CAC $800-1500. Why do you think you can achieve half that? If CAC is actually $800, your LTV:CAC ratio drops from 5:1 to 2.5:1, which is barely viable.' The Devil's Advocate forces the Marketing Agent to defend optimistic assumptions or revise them.

In Phase 3 (future), the Marketing Agent will re-run as the business runs actual marketing campaigns. When the CEO runs a pilot campaign and measures actual CAC, conversion rates, and LTV, the Marketing Agent updates its model with real data instead of estimates. This allows the business plan to evolve from assumed to validated.""",

        "Why This Matters (Real-World Impact)": """The Marketing Strategy Agent is the reason the multi-agent system produces business plans with realistic revenue projections. Many startup failures stem from assuming revenue will magically appear without a concrete plan to acquire customers cost-effectively.

In real-world usage, the Marketing Agent prevents the CEO from launching without a customer acquisition plan. If the CEO's plan is 'build the product and figure out marketing later', the Marketing Agent forces the question: 'Who is the customer? How will you reach them? What will it cost? Do the unit economics work?' This prevents the CEO from building a product no one knows exists or building a product that costs more to acquire customers for than it generates in revenue.

The CAC and LTV analysis is particularly impactful. Many founders do not calculate unit economics until after they have burned $50K-100K on marketing with poor ROI. The Marketing Agent calculates unit economics upfront, before any marketing spend, and flags if the business model does not work. If LTV is $1,500 and estimated CAC is $1,200, the LTV:CAC ratio is 1.25:1, which is not viable. The CEO then knows they need to either increase LTV (higher pricing, lower churn, upsell) or reduce CAC (better targeting, more efficient channels) before launching.

The channel viability assessment prevents wasted marketing spend. Many founders copy what other startups do ('Facebook ads worked for Company X, so we will do Facebook ads') without evaluating whether that channel fits their business. The Marketing Agent evaluates channels specific to the ICP and business model. If the target customer is B2B enterprises with $50K ACV, Facebook ads will not work (too low intent, too broad targeting). The right channel is direct sales with warm intros. The Marketing Agent makes that call and prevents the CEO from wasting $20K on Facebook ads that generate zero enterprise customers.

The funnel conversion modeling prevents fantasy revenue projections. Many founders assume 'if we get 10K website visitors, we will convert 5% to customers = 500 customers Year 1.' Real B2B SaaS conversion from visitor to paying customer is 0.1-0.5%, not 5%. The Marketing Agent provides realistic benchmarks, so the Financial Agent does not project $2M revenue based on fantasy conversion rates.

The phased go-to-market timeline prevents premature scaling. Many founders start spending $10K/month on paid ads on Day 1, before validating messaging, before understanding the customer, before having any conversion data. They burn $50K-100K with zero ROI. The Marketing Agent produces a phased plan: start with customer discovery and pilots (Month 0-3), test channels with small budgets (Month 3-6), scale what works (Month 6-12). This capital-efficient approach is critical for bootstrapped or seed-stage startups.

The messaging and positioning development gives the CEO a clear value prop and differentiation story. Many founders struggle to articulate why a customer should buy their product instead of a competitor's or instead of doing nothing. The Marketing Agent synthesizes the SWOT strengths and competitive landscape into a concise positioning statement that the CEO can use in sales conversations, pitch decks, and marketing materials.

Finally, the Marketing Agent makes the business plan credible to investors. When an investor reviews the plan and sees a detailed go-to-market strategy with realistic CAC, LTV, funnel model, and channel prioritization, they see a founder who understands customer acquisition. Investors know that revenue is hard, and a plan with no marketing strategy is a plan that will not hit its revenue targets.""",
    }

    create_doc("Marketing Strategy Agent - Customer Acquisition and Go-to-Market Planning", sections, output_dir / "marketing_strategy_explanation.docx")


# Run the expansions
def main():
    output_dir = Path("/home/saiaditya26122006/multi-agent-system/explaination")
    logger.info("Expanding final 6 agent documents with very detailed content...")

    expand_tech_stack(output_dir)
    expand_marketing(output_dir)

    logger.info("\n✅ Expanded 2/6 documents (Tech Stack, Marketing). Continuing with remaining 4...")
    logger.info("Note: Operations, Financial, Launch, and Exit Strategy require separate script due to size")


if __name__ == "__main__":
    main()
