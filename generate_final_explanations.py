"""
Generate detailed Word document explanations for the final 3 agents:
- Summary Agent
- Devil's Advocate Agent
- Council Agent

Each document explains the agent's purpose, implementation, and real-world impact
in conversational style with extensive detail.
"""

import logging
from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)


def create_detailed_explanation_doc(
    agent_name: str,
    title: str,
    sections: dict[str, str],
    output_path: Path,
) -> None:
    """Create a detailed Word document explanation for an agent."""
    doc = Document()

    # Title
    title_para = doc.add_heading(title, level=1)
    title_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Add sections
    for section_title, content in sections.items():
        doc.add_heading(section_title, level=2)

        # Split content into paragraphs for better formatting
        paragraphs = content.strip().split("\n\n")
        for para_text in paragraphs:
            if para_text.strip():
                para = doc.add_paragraph(para_text.strip())
                para.style.font.size = Pt(11)

    doc.save(output_path)
    logger.info(f"Created {output_path.name} ({output_path.stat().st_size} bytes)")


def generate_summary_agent_explanation(output_dir: Path) -> None:
    """Generate explanation for Summary Agent."""
    sections = {
        "What This Agent Does": """The Summary Agent is the final synthesizer in the entire business plan pipeline. After all the other agents have completed their specialized work on market research, financials, operations, technology, marketing, and exit strategy, this agent reads everything they produced and creates a single, CEO-ready executive summary.

Think of it as the translator between the detailed technical analysis done by 10+ specialized agents and what a busy CEO or investor actually needs to read. It takes hundreds of pages of structured analysis and distills it down to a 3-5 page executive summary that captures the essence of the entire business plan.

The Summary Agent does not do any new analysis. It does not make decisions. It does not add new recommendations. Its job is pure synthesis. It reads the outputs from all prior sections, identifies the key insights, the biggest risks, the most critical assumptions, and packages them into a narrative that flows logically and tells the complete story of the business opportunity.

It produces a structured executive summary with these components: business opportunity overview, market validation summary, competitive positioning, financial projections highlight (revenue, break-even, exit valuation), go-to-market strategy summary, key risks and mitigation, critical assumptions that need validation, recommended next actions, and an overall confidence assessment.

The agent is smart about what to include and what to leave out. It highlights the findings that matter most for decision-making. If the Financial Agent flagged that pricing is completely unvalidated, the Summary Agent makes sure that shows up prominently in the executive summary. If the Marketing Agent found a specific channel with strong early traction, that gets highlighted too.

The output is designed to be read in 10 minutes and give someone a complete picture of whether this business idea is worth pursuing, what the main risks are, and what needs to happen next to de-risk it.""",

        "Why We Built This Agent": """We built the Summary Agent because even the best analysis is useless if decision-makers do not read it. And decision-makers will not read 50 pages of detailed agent outputs.

In the early design of the system, we assumed the CEO would just read all the individual section outputs. We quickly realized that was unrealistic. Each agent produces 2-10 pages of structured output. Across 11 sections, that is over 100 pages of dense analysis. No CEO has time for that, and even if they did, they would have to do the synthesis work themselves to figure out what it all means.

The Summary Agent solves the cold start problem for decision-making. It gives the CEO the essential information they need to decide whether to move forward, what to validate first, and what the biggest risks are. If they want to dig deeper into any specific area, the detailed section outputs are available, but they do not have to read them all to get started.

It also solves the communication problem for external stakeholders. If the CEO needs to pitch this business idea to an investor, a co-founder, or a potential customer, they cannot share 100 pages of agent outputs. They need a clean, professional executive summary that stands on its own. The Summary Agent produces exactly that.

We also built it because synthesis is a distinct skill from analysis. The other agents are specialists. They go deep on their domain (financials, marketing, operations). But someone needs to step back and see the big picture. Someone needs to ask: what are the 5 things that matter most across all this analysis? What is the through-line that connects market opportunity to financial viability to operational feasibility? That is what the Summary Agent does.

Finally, we built it to preserve the human-in-the-loop design of the system. The Summary Agent does not make the final decision. It presents the information clearly so the CEO can make the decision. It highlights where assumptions are weak, where risks are high, and where validation is needed. It gives the CEO the tools to be the ultimate decision-maker, not a passive recipient of AI recommendations.""",

        "Key Features": """First, it reads outputs from all prior sections. The agent has access to the structured outputs from Opportunity Analyst, Environment Research, Organisation Designer, SWOT, Tech Stack, Marketing, Operations, Financial, Launch & Contingency, and Exit Strategy. It does not read the reasoning traces or intermediate steps. It reads the final validated outputs that each agent produced.

Second, it uses intelligent summarization, not just copy-paste. The agent uses Claude Haiku to synthesize the outputs, not just concatenate them. It identifies themes, connections, and patterns. If three different agents all flagged pricing uncertainty, the Summary Agent groups that into a single critical assumption rather than listing it three times.

Third, it produces a structured executive summary with specific sections: business opportunity (2-3 sentences), market validation (what we know vs what we assumed), competitive positioning, financial highlights (Year 3 revenue, break-even month, exit valuation), go-to-market strategy, top 3 risks, top 5 critical assumptions, recommended next actions (prioritized), and overall confidence score.

Fourth, it preserves source attribution. When the Summary Agent says 'pricing is unvalidated', it references which agent made that assessment (Financial Agent, Section 12). When it says 'CAC assumed at $500 with no pilot data', it cites the Marketing Agent. This lets the CEO trace any claim back to its source if they want more detail.

Fifth, it applies business-logic filters to highlight what matters. Not every assumption is equally important. The Summary Agent knows that unvalidated pricing is more critical than unvalidated hiring timelines. It knows that customer acquisition risk is more important than office location risk. It uses the confidence scores and uncertainty flags from each agent to prioritize what to highlight.

Sixth, it produces a confidence-calibrated overall assessment. Each agent produces a confidence score (high/medium/low). The Summary Agent aggregates these and produces an overall plan confidence score. If 8 out of 10 agents returned low confidence, the executive summary makes that very clear. If most agents had high confidence except for one critical area (e.g. financial projections), that gets flagged as the bottleneck.

Seventh, it outputs in both JSON and natural language. The structured JSON is used by downstream systems (dashboards, Notion sync, archival). The natural language executive summary is what the CEO reads. Both are generated in a single agent run to ensure consistency.

Eighth, it handles missing or incomplete section outputs gracefully. If a prior agent failed or returned partial data, the Summary Agent flags that in the executive summary ('Note: Exit Strategy section unavailable - manual review required'). It does not fail silently.""",

        "How It Works Technically": """The Summary Agent is a Haiku-powered agent that runs last in the pipeline. It uses the same Intelligence Engine (4-step reasoning: Decompose, Produce, Challenge, Revise) as all other agents, but with a synthesis-focused prompt.

The agent reads from a centralized prior_outputs dictionary that contains the validated JSON outputs from all earlier sections. Each entry in this dictionary is keyed by section number ('1', '3', '4', '5', '6.5', '8', '10', '12', '13', '14') and contains the full structured output from that agent.

The agent does not re-run any analysis. It does not call external APIs. It does not access the CEO data directly. Everything it needs is in the prior_outputs dictionary. This makes the Summary Agent fast (typically completes in 30-60 seconds) and deterministic (same inputs always produce same output structure).

The input schema for the Summary Agent includes: all prior section outputs, the CEO's original business idea description, the target market, and any high-level constraints or preferences the CEO provided. The agent does not need detailed CEO data because the other agents already incorporated that into their outputs.

The output schema is a Pydantic model with these fields: section_number ('executive_summary'), business_opportunity_summary (150-300 words), market_validation_summary (what is validated vs assumed), competitive_positioning_summary, financial_highlights (Year 3 revenue, break-even, exit valuation, funding required), go_to_market_summary, top_risks (list of 3-5 risks with severity and mitigation), critical_assumptions (list of 5-10 assumptions that need validation), recommended_next_actions (prioritized list), overall_confidence_score (high/medium/low), assumptions_used (inherited from prior sections), uncertainties (aggregated from all sections), input_tokens, output_tokens.

The agent uses Claude Haiku instead of Sonnet because synthesis is less cognitively demanding than original analysis. Haiku is fast and cheap, and for summarization tasks, it performs nearly as well as Sonnet. This keeps the cost of the final step low.

The prompt engineering is critical here. The system prompt tells the agent: 'You are synthesizing a business plan. Your job is to highlight what matters most for decision-making. Prioritize risks, assumptions, and validation needs over positive claims. Be concise. Be honest about uncertainty. Do not add new analysis. Do not make recommendations that are not already in the prior section outputs.'

The agent uses the P0-3 confidence calibration system to produce the overall_confidence_score. It looks at the confidence scores from all prior agents. If most returned low confidence, the overall score is low. If most returned medium or high, the overall score reflects that. The calibration also considers the criticality of each section - low confidence in financials is worse than low confidence in org design.

After the agent produces the structured JSON output, a post-processing step generates a natural language executive summary from the JSON. This uses a lightweight template that formats the JSON fields into paragraphs. The natural language version is what gets exported to Word documents, PDFs, or sent to the CEO via Telegram.

The Summary Agent writes its output to the agent_outputs table in Supabase with section_number = 'executive_summary'. It also updates the session state to COMPLETED (if all sections succeeded) or NEEDS_REVIEW (if any section failed or returned low confidence).

Error handling: if any prior section is missing from prior_outputs, the Summary Agent flags that in the uncertainties list ('Section X output unavailable - manual review required'). If it cannot parse a prior section output (corrupted JSON), it logs the error and excludes that section from the summary with a warning.""",

        "How It Connects to Other Agents": """The Summary Agent is the terminal node in the pipeline. It depends on all other agents and no agent depends on it. It runs last, after all sections (1, 3, 4, 5, 6.5, 8, 10, 12, 13, 14) have completed.

The execution order is enforced by the Mother Agent (Phase 2 orchestrator). The Mother Agent tracks which sections have completed and only triggers the Summary Agent once all required sections are done. If any section fails, the Mother Agent decides whether to proceed with a partial summary (flagging the missing section) or halt and request human intervention.

The Summary Agent reads outputs from: (1) Opportunity Analyst (business idea validation), (3) Environment Research (market landscape), (4) Organisation Designer (team and roles), (5) SWOT (strengths, weaknesses, opportunities, threats), (6.5) Tech Stack (infrastructure and compliance), (8) Marketing Strategy (channels, CAC, LTV), (10) Operations (workflows, vendors, timelines), (12) Financial Modelling (revenue, costs, break-even, valuation), (13) Launch & Contingency (go-live plan and fallback), (14) Exit Strategy (acquisition targets, investor returns).

It does not send its output to any other agent. The CEO is the consumer of the executive summary. The output is displayed in the Streamlit monitoring UI, exported to Word/PDF, and sent via Telegram to the CEO's phone.

The Summary Agent also feeds into the manual review workflow. If the overall confidence score is low, the system flags the session for human review. A human analyst reads the executive summary, checks the detailed section outputs for the low-confidence areas, and decides whether to approve, request clarification, or reject the plan.

In Phase 3 (future), the Summary Agent will also feed into the Council Agent. The Council Agent is a 5-persona deliberation system that reviews high-stakes business plans before they are finalized. The Council reads the executive summary first, then selectively reads detailed sections based on what it wants to challenge. The Summary Agent's job is to give the Council a fast overview so it can decide where to focus its scrutiny.

The Summary Agent also connects to the archival and versioning system. Every time a business plan is re-evaluated (because the CEO provided new data or changed assumptions), a new executive summary is generated. The system stores all versions with timestamps so the CEO can see how the plan evolved over time. The executive summary is the primary artifact used for version comparison.""",

        "Why This Matters (Real-World Impact)": """The Summary Agent is the reason the entire multi-agent system is usable by non-technical CEOs. Without it, the system would be a technical demo that produces tons of data but no actionable insights.

In real-world usage, the CEO interacts almost exclusively with the executive summary. They receive it via Telegram as a formatted message or as a Word document attachment. They read it in 10 minutes and decide: 'Yes, this looks promising, let me dig into the financials' or 'No, the customer acquisition risk is too high, I will pivot the idea' or 'Maybe, but I need to validate these 3 assumptions first before I invest more time'.

The Summary Agent also makes the system useful for external stakeholders. If the CEO wants to share the analysis with a potential co-founder, investor, or advisor, they forward the executive summary. It is professional, self-contained, and tells the complete story without requiring the recipient to understand the multi-agent architecture.

The agent also acts as a quality gate for the system. If the executive summary reads incoherently or contradicts itself (e.g. 'low customer acquisition risk' in the summary but 'CAC completely unvalidated' in the Marketing section), that indicates a bug in one of the upstream agents. The Summary Agent's output is the integration test for the entire pipeline.

It also reduces the CEO's cognitive load. Business planning is overwhelming. There are hundreds of variables to consider. The Summary Agent does the synthesis work so the CEO can focus on decision-making. It answers the questions CEOs actually ask: What is the opportunity? What are the risks? What do I need to validate? What should I do next?

The confidence calibration is particularly impactful. Instead of presenting the business plan as a fait accompli, the Summary Agent honestly flags uncertainty. It says: 'This plan assumes pricing of $5,000 per customer with no validation. If actual willingness-to-pay is 50% lower, the entire financial model breaks.' That level of transparency builds trust. The CEO knows the system is not overselling the idea.

Finally, the Summary Agent enables iteration. Because it runs fast (30-60 seconds) and produces a compact output, the CEO can iterate quickly. They can update an assumption, re-run the pipeline, and see how the executive summary changes. This turns the system into a planning tool, not just a one-time analysis engine. The CEO can explore different scenarios (What if I target US universities instead of EU? What if I charge $10K instead of $5K?) and immediately see the impact on the overall plan viability.""",
    }

    output_path = output_dir / "summary_agent_explanation.docx"
    create_detailed_explanation_doc(
        "summary_agent",
        "Summary Agent - Executive Summary Synthesizer",
        sections,
        output_path,
    )


def generate_devils_advocate_explanation(output_dir: Path) -> None:
    """Generate explanation for Devil's Advocate Agent."""
    sections = {
        "What This Agent Does": """The Devil's Advocate Agent is the contrarian in the system. Its job is to challenge, stress-test, and poke holes in the outputs produced by other agents. After an agent completes its work and produces a structured output, the Devil's Advocate reviews it with a skeptical eye and asks: What could go wrong? What assumptions are shaky? What evidence is missing? What has been overlooked?

Unlike the other agents, which are focused on building the business plan, the Devil's Advocate is focused on breaking it. It looks for logical inconsistencies, weak evidence chains, overconfident claims, and blind spots. It does not produce new business plan content. It produces criticism.

The agent operates as an adversarial reviewer. After the Marketing Agent claims CAC will be $400, the Devil's Advocate asks: Based on what data? What if paid ads do not work in this market? What if organic is slower than assumed? It forces the system to confront uncomfortable questions before the CEO sees the final plan.

It produces a structured critique with these components: claim being challenged, evidence assessment (what evidence supports this claim and is it sufficient?), alternative scenarios (what if the opposite is true?), missing validation (what would need to be true for this claim to hold?), risk rating (low/medium/high if this claim is wrong), and recommendation (validate now vs acceptable assumption vs revise the claim).

The Devil's Advocate is not pessimistic for the sake of being pessimistic. It is realistic. It distinguishes between assumptions that are reasonable given the available data and assumptions that are wild guesses dressed up as analysis. It helps the system be honest with the CEO about what is known versus what is hoped.

The agent runs selectively, not on every section. It is invoked for high-stakes sections (financials, marketing, exit strategy) where overconfidence could lead to bad decisions. It does not run on low-stakes sections (org design, operations) where the cost of being wrong is lower.

The output is consumed by the Mother Agent (Phase 2 orchestrator), which decides whether the original agent needs to revise its work, whether the uncertainty should be flagged in the executive summary, or whether the critique is noted but the original output stands.""",

        "Why We Built This Agent": """We built the Devil's Advocate Agent because AI agents are overconfident by default. LLMs like Claude are trained to be helpful and confident. When you ask them to estimate customer acquisition cost, they will give you a number. They will not say 'I have no idea because you have not run a single marketing campaign yet.' They will make reasonable-sounding assumptions and present them as analysis.

This is dangerous in a business planning context. If the Financial Agent assumes $500 CAC with no pilot data, and the Marketing Agent assumes 5% conversion with no testing, and the Operations Agent assumes 2-week sales cycles with no customer feedback, the CEO ends up with a business plan that looks rigorous but is built on sand.

The Devil's Advocate Agent solves the overconfidence problem. It acts as the system's internal skeptic. After each agent produces its output, the Devil's Advocate reviews it and asks: How confident should we actually be in this? What would need to be true for this to work? What are we assuming that we should be validating?

We also built it because humans have confirmation bias. If the CEO wants to believe their business idea is great, they will read the agent outputs through a positive lens. They will focus on the promising projections and skim over the uncertainties. The Devil's Advocate forces the uncomfortable questions to the surface. It makes the uncertainties impossible to ignore.

The agent also improves the quality of other agents' outputs. Knowing that their work will be adversarially reviewed incentivizes agents to be more careful about evidence, more explicit about assumptions, and more conservative in their confidence scores. It is like having a peer review process built into the pipeline.

We built it because business planning is inherently uncertain, and the system needs to be honest about that uncertainty. The Devil's Advocate is the honesty enforcer. It prevents the system from presenting guesses as facts.

Finally, we built it to protect the CEO from bad decisions. If the CEO proceeds with a business plan based on unvalidated assumptions, they could waste months of time and tens of thousands of dollars. The Devil's Advocate is the early warning system that says: 'Before you commit resources, validate these 5 critical assumptions, because if any of them are wrong, the entire plan falls apart.'""",

        "Key Features": """First, it operates in adversarial mode. The Devil's Advocate is explicitly prompted to be skeptical, not balanced. Its job is not to say 'this could work or it could not work.' Its job is to say 'here is why this might not work.' The system has plenty of agents building the case for the business. The Devil's Advocate is the only agent building the case against it.

Second, it uses evidence-based critique. The agent does not just say 'I do not trust this assumption.' It says 'This assumption is based on zero pilot data, one unverified benchmark, and a claim from an agent that admitted low confidence. If this assumption is off by 50%, the break-even timeline extends from 18 months to 36 months, which exhausts the assumed funding runway.'

Third, it produces structured critiques with specific fields: claim (the exact statement being challenged), evidence_assessment (what supports this claim and is it sufficient?), alternative_scenarios (what if the opposite is true?), validation_gap (what is missing to be confident?), risk_if_wrong (low/medium/high impact), recommended_action (validate now, acceptable risk, or revise claim).

Fourth, it runs selectively on high-stakes sections. The Devil's Advocate does not challenge every single agent output. It focuses on sections where overconfidence is most dangerous: financials (Section 12), marketing (Section 8), and exit strategy (Section 14). These are the sections where bad assumptions can sink the business.

Fifth, it has access to cross-section context. The Devil's Advocate can spot inconsistencies across sections. If the Marketing Agent assumes 30 customers Year 1 and the Operations Agent assumes 50 customers Year 1, the Devil's Advocate flags that. If the Financial Agent assumes break-even at 18 months but the Marketing Agent's funnel math implies 24 months, the Devil's Advocate catches that.

Sixth, it distinguishes between fatal flaws and acceptable risks. Not every assumption needs to be validated immediately. Some risks are worth taking. The Devil's Advocate uses a risk prioritization framework: if the assumption is critical (financials depend on it) and the evidence is weak (no validation), it flags 'validate now.' If the assumption is low-impact or well-evidenced, it flags 'acceptable risk.'

Seventh, it integrates with the confidence calibration system. When the Devil's Advocate identifies a high-risk assumption, it can downgrade the confidence score of the section that made that assumption. If the Marketing Agent returned medium confidence but the Devil's Advocate found fatal flaws in the CAC estimate, the final recorded confidence for that section is downgraded to low.

Eighth, it produces actionable recommendations. The output is not just 'this is wrong.' The output is 'this is wrong because X, here is what would need to be true for it to be right, here is how to validate it, and here is what happens if we do not validate it.'""",

        "How It Works Technically": """The Devil's Advocate Agent is a Sonnet-powered agent that runs conditionally after high-stakes sections complete. It uses the same Intelligence Engine (4-step reasoning) as other agents, but with an adversarial prompt.

The agent is invoked by the Mother Agent (Phase 2 orchestrator) after sections 8 (Marketing), 12 (Financial), and 14 (Exit Strategy) complete. The Mother Agent passes the completed section output plus any cross-section context the Devil's Advocate might need (e.g. if challenging financials, it gets access to marketing assumptions about customer volume and pricing).

The input schema includes: section_number (which section is being challenged), section_output (the full validated JSON output from that agent), cross_section_context (outputs from other sections that might be relevant), ceo_data_summary (high-level CEO inputs, used to check if assumptions contradict stated constraints).

The output schema is a Pydantic model with: section_challenged, critiques (list of critique objects, each with claim, evidence_assessment, alternative_scenarios, validation_gap, risk_if_wrong, recommended_action), cross_section_inconsistencies (list of contradictions found across sections), overall_assessment (how confident should we be in this section after these critiques?), confidence_adjustment (should this section's confidence score be downgraded?), assumptions_used, uncertainties, input_tokens, output_tokens.

The agent uses Claude Sonnet, not Haiku, because adversarial reasoning is cognitively demanding. The agent needs to understand complex financial models, spot subtle logical flaws, and reason counterfactually ('what if the opposite is true?'). Sonnet is significantly better at this than Haiku.

The prompt engineering is critical. The system prompt says: 'You are a Devil's Advocate reviewing a business plan section. Your job is to challenge assumptions, stress-test claims, and identify risks. Be specific. Do not reject claims just because they are assumptions - all business plans have assumptions. Reject claims that are unsupported by evidence, inconsistent with other sections, or overly optimistic given the uncertainty. Prioritize your critiques - focus on the assumptions that, if wrong, would invalidate the entire section.'

The agent uses a critique framework with three severity levels: (1) Fatal flaw - this claim is unsupported and if wrong, the section's conclusions are invalid. Validate immediately. (2) Material risk - this claim is shaky and if wrong, the section's conclusions change significantly. Validate before launch. (3) Minor risk - this claim is uncertain but even if wrong, the impact is manageable. Document but proceed.

After the Devil's Advocate produces its output, the Mother Agent decides what to do with the critiques. If there are fatal flaws, the original agent is re-invoked to revise its output with the critiques incorporated. If there are material risks, those are flagged in the executive summary as critical assumptions needing validation. If there are only minor risks, those are noted but do not block progress.

The Devil's Advocate does not rewrite the original section output. It only produces critiques. The decision about whether to revise is made by the Mother Agent based on the severity of the critiques and the confidence_adjustment recommendation.

Error handling: if the Devil's Advocate fails (API error, timeout, invalid output), the system logs the failure but does not block progress. The original section output stands, but a warning is added to the uncertainties list: 'Devil's Advocate review unavailable for Section X - manual review recommended.'""",

        "How It Connects to Other Agents": """The Devil's Advocate Agent is a meta-agent. It does not produce business plan content. It reviews content produced by other agents. It runs after selected agents complete but before the Summary Agent synthesizes the final plan.

The execution order is: (1) Primary agent (Marketing, Financial, or Exit Strategy) completes and produces output. (2) Mother Agent triggers Devil's Advocate to review that output. (3) Devil's Advocate produces critiques. (4) Mother Agent decides whether to re-run the primary agent with the critiques incorporated. (5) If revised, the new output replaces the original. (6) Summary Agent reads the final (possibly revised) outputs.

The Devil's Advocate is invoked for sections 8, 12, and 14. These are the sections where overconfidence is most dangerous. It is not invoked for sections 1, 3, 4, 5, 6.5, 10, 13 because those sections are lower-stakes (if the org chart is slightly wrong, the business does not fail; if the financial projections are off by 3x, the business absolutely fails).

The agent reads from the same prior_outputs dictionary that other agents use. When challenging Section 12 (Financial), it has access to Section 8 (Marketing) to check if the revenue assumptions are consistent with the funnel math. When challenging Section 8 (Marketing), it has access to Section 1 (Opportunity) to check if the CAC assumptions are realistic given the competitive landscape.

The output is consumed by the Mother Agent, which decides the next action. If critiques are severe, the Mother Agent re-runs the primary agent with an augmented prompt: 'The Devil's Advocate identified these issues with your previous output: [critiques]. Revise your analysis to address these critiques. Be more conservative in your assumptions where evidence is weak.'

The Devil's Advocate also feeds into the Summary Agent. Even if the primary agent is not re-run, the critiques are passed to the Summary Agent. The Summary Agent incorporates the key critiques into the executive summary's 'Critical Assumptions' and 'Recommended Next Actions' sections.

In Phase 3 (future), the Devil's Advocate will also interact with the Council Agent. The Council is a 5-persona deliberation system for high-stakes decisions. The Devil's Advocate's critiques seed the Council's discussion. Each persona in the Council reads the critiques and decides which ones they find most compelling, which ones they disagree with, and what additional concerns they want to raise.

The Devil's Advocate does not interact with the CEO directly. The CEO sees the executive summary, which includes the key critiques distilled into the 'Critical Assumptions' section. If the CEO wants to see the full adversarial review, they can access the detailed Devil's Advocate output in the Streamlit monitoring UI.""",

        "Why This Matters (Real-World Impact)": """The Devil's Advocate Agent is the reason the multi-agent system can be trusted for real-world business decisions. Without it, the system would be an overconfident AI that produces polished-looking but ultimately unreliable business plans.

In real-world usage, the Devil's Advocate prevents the CEO from making expensive mistakes. If the Financial Agent projects break-even in 18 months based on unvalidated CAC and pricing assumptions, and the CEO commits $200K of seed funding based on that projection, and the actual break-even is 36 months, the CEO runs out of money before reaching profitability. The Devil's Advocate catches that risk early by saying: 'This break-even timeline assumes CAC of $500 with zero pilot data. If actual CAC is $1,000 (which is typical for B2B SaaS in this space), break-even extends to 30+ months. Validate CAC before committing capital.'

The agent also improves the credibility of the system with external stakeholders. If the CEO shares the business plan with an experienced investor, the investor will immediately spot weak assumptions. If the plan includes a Devil's Advocate critique that already flagged those weak assumptions and recommended validation steps, the investor sees that the CEO is being thoughtful and rigorous, not naive.

The Devil's Advocate also enables better iteration. When the CEO updates an assumption (e.g. changes pricing from $5K to $7K), the system re-runs the pipeline. The Devil's Advocate re-evaluates the updated plan and checks if the new assumption is better supported or introduces new risks. This gives the CEO confidence that they are not just moving the uncertainty around but actually improving the plan quality.

The confidence calibration impact is significant. Without the Devil's Advocate, an agent might return medium confidence even when its assumptions are shaky, because the agent is not actively trying to find flaws in its own reasoning. With the Devil's Advocate, that same section's confidence gets downgraded to low, and the CEO sees in the executive summary: 'Financial projections have low confidence due to unvalidated CAC and pricing assumptions.' That transparency prevents overcommitment.

The agent also educates the CEO about business planning rigor. By reading the Devil's Advocate critiques, the CEO learns what 'validated' versus 'assumed' means. They learn that 'we assume 5% conversion' is very different from 'we tested 3 channels and observed 4-6% conversion.' Over time, this makes the CEO a better business planner.

Finally, the Devil's Advocate makes the multi-agent system more defensible. If a decision goes wrong, the CEO can go back and check: Did the system flag this risk? Was I warned? In most cases, the answer will be yes. The Devil's Advocate leaves a clear audit trail of what risks were identified, what assumptions were flagged as needing validation, and what the system recommended. This protects the CEO (and the system developers) from the accusation that the AI gave bad advice. The AI gave honest advice, with clear warnings about uncertainty.""",
    }

    output_path = output_dir / "devils_advocate_explanation.docx"
    create_detailed_explanation_doc(
        "devils_advocate",
        "Devil's Advocate Agent - Adversarial Business Plan Reviewer",
        sections,
        output_path,
    )


def generate_council_explanation(output_dir: Path) -> None:
    """Generate explanation for Council Agent."""
    sections = {
        "What This Agent Does": """The Council Agent is a deliberation system that simulates a panel of 5 expert personas reviewing a high-stakes business plan. Unlike other agents that produce analysis, the Council produces judgment. After all the specialized agents have completed their work and the Devil's Advocate has challenged it, the Council convenes to decide: Should this plan proceed? Should it be revised? Or should it be killed?

The Council is modeled after real-world governance structures. In a startup, major decisions (like pivoting the business model or raising a funding round) are not made by one person alone. The CEO consults advisors, co-founders, domain experts, and investors. Each brings a different perspective. The technical co-founder cares about feasibility. The investor cares about ROI. The domain expert cares about market fit. The Council simulates that multi-perspective deliberation.

The 5 personas in the Council are: (1) The Operator - focused on execution, asks 'can we actually build and deliver this?' (2) The Investor - focused on returns, asks 'will this make money and when?' (3) The Market Strategist - focused on competition and positioning, asks 'why will customers choose us?' (4) The Risk Manager - focused on downside, asks 'what could sink this business?' (5) The Visionary - focused on long-term potential, asks 'is this idea big enough to matter?'

Each persona reads the executive summary and selected detailed sections. They form an initial opinion (Approve / Request Revisions / Reject). They then engage in a structured deliberation where they share their reasoning, challenge each other, and update their positions based on the discussion. The Council produces a final verdict with a vote tally (e.g. 4 Approve, 1 Reject), a summary of key points of agreement and disagreement, and specific conditions or validation steps required before launch.

The Council does not run on every business plan. It is invoked only for high-stakes decisions: plans requesting >$100K in funding, plans with low overall confidence scores, plans where the Devil's Advocate flagged fatal flaws, or plans where the CEO explicitly requested a rigorous review.

The output is a structured deliberation transcript with each persona's reasoning, the final vote, and a consolidated recommendation. This is presented to the CEO as the final gate before proceeding with the plan.""",

        "Why We Built This Agent": """We built the Council Agent because single-agent decision-making is brittle. Even the best AI agent has blind spots. An agent optimized for financial analysis might underweight operational complexity. An agent optimized for market analysis might overweight market size and underweight competitive intensity. The Council fixes this by forcing multiple perspectives to engage with each other.

We also built it because major business decisions should not be made by one AI or one human alone. In the real world, CEOs who make big decisions in isolation tend to make worse decisions than CEOs who consult diverse advisors. The Council simulates that advisory board deliberation process.

The Council also solves the 'AI says yes, but should I trust it?' problem. If a single agent says 'this business plan looks good, proceed', the CEO might wonder if the AI is overconfident or missed something. But if 5 different personas, each with a different lens, all independently review the plan and 4 out of 5 say 'proceed', the CEO can be more confident the plan has been rigorously vetted.

We built it because deliberation improves reasoning. When the Operator persona says 'we can build this in 6 months' and the Risk Manager persona challenges 'but the org design shows only 2 engineers Year 1, how do you deliver this feature set with 2 engineers?', the Operator has to defend its position. That back-and-forth produces better analysis than either persona would have produced alone.

The Council also makes the system more transparent. Instead of the system saying 'Confidence: Medium, proceed', the Council shows the reasoning behind that judgment. The CEO can read the Investor persona's logic for approving, the Risk Manager persona's concerns, and the Visionary persona's assessment of long-term potential. This lets the CEO understand why the system reached its conclusion, not just what the conclusion is.

We also built it as a safeguard against overcommitment. If the CEO is excited about a business idea, they might dismiss a single warning from the Devil's Advocate. But if the Council votes 2 Approve, 3 Reject, and the majority opinion cites unvalidated assumptions and high market risk, the CEO has to take that seriously. The Council is the system's strongest 'pump the brakes' mechanism.

Finally, we built it because Phase 3 of the system (future) will include autonomous decision-making in low-stakes contexts. For example, the system might autonomously approve small budget line items or operational changes. The Council is the training ground for that autonomous judgment. By simulating human-like deliberation now, we are building the foundation for safely delegating decisions to AI later.""",

        "Key Features": """First, it uses 5 distinct personas with different incentives and risk tolerances. The Operator cares about feasibility and does not want to commit to building something impossible. The Investor cares about ROI and does not want to fund a money pit. The Market Strategist cares about competition and does not want to enter a crowded market with no differentiation. The Risk Manager cares about downside and does not want to bet the company on unvalidated assumptions. The Visionary cares about upside and does not want to pursue small ideas. These personas often disagree, which is the point.

Second, it uses a structured deliberation process with 3 rounds: (1) Independent review - each persona reads the executive summary and forms an initial opinion without seeing what others think. (2) Deliberation - personas share their reasoning and challenge each other. The Investor might say 'I am concerned about break-even timeline' and the Operator responds 'that timeline assumes pessimistic CAC; if marketing hits their target, break-even is 6 months earlier.' (3) Final vote - after deliberation, each persona casts a final vote (Approve / Revise / Reject) and the majority wins.

Third, it produces a detailed transcript. The Council output includes each persona's initial reasoning, key points raised during deliberation, how opinions shifted (if at all), the final vote tally, points of consensus (what all personas agreed on), points of contention (what personas disagreed about), and minority opinions (if the vote was split, what was the minority's reasoning?).

Fourth, it handles tie-breaking and split decisions gracefully. If the vote is 3-2 or 2-2-1 (some Approve, some Revise, some Reject), the Council flags that as a 'contentious decision' and elevates it to the CEO with a clear summary of the disagreement. It does not force a false consensus.

Fifth, it uses cross-persona challenges to improve reasoning. The Risk Manager persona is explicitly prompted to challenge optimistic assumptions. If the Investor says 'this will generate $1M ARR by Year 3', the Risk Manager asks 'based on what validated CAC and conversion rate?' This forces the reasoning to be more rigorous.

Sixth, it integrates with all prior outputs. The Council has access to the executive summary, all detailed section outputs, Devil's Advocate critiques, and the CEO's original inputs. Each persona can selectively deep-dive into specific sections based on their focus area. The Investor reads financials in detail. The Operator reads tech stack and operations in detail. The Market Strategist reads market research and marketing strategy in detail.

Seventh, it produces actionable conditions. If the Council votes 'Approve with conditions', those conditions are specific: 'Approve pending validation of CAC below $600 via 4-week paid ads pilot in 2 channels' or 'Approve pending confirmation of 10 LOIs (letters of intent) from target customers.' These conditions become the CEO's next-actions checklist.

Eighth, it has a confidence-weighted voting system. Not all votes are equal. If a persona admits low confidence in their opinion, that vote carries less weight in the final tally. If a persona has high confidence, that vote carries more weight. This prevents weak opinions from overriding strong ones.""",

        "How It Works Technically": """The Council Agent is a multi-turn Claude Sonnet orchestration. Unlike single-shot agents that produce output in one LLM call, the Council makes multiple LLM calls to simulate the back-and-forth of a real deliberation.

The agent runs in 3 phases: (1) Independent review phase - 5 separate LLM calls, one per persona. Each persona is given the executive summary, selected detailed sections, and a persona-specific prompt ('You are the Operator. Assess whether this plan is operationally feasible. Consider team capabilities, timelines, and execution risks. Vote Approve, Revise, or Reject and explain your reasoning.'). Each persona returns an initial vote and reasoning. (2) Deliberation phase - 3-5 turns of discussion. In each turn, all personas see the current state of the discussion and respond. The Risk Manager might challenge the Investor's optimism. The Operator might defend the timeline. The Visionary might argue the market opportunity justifies the risk. (3) Final vote phase - each persona casts a final vote, potentially updated based on the deliberation. The system tallies the votes and produces a consolidated recommendation.

The input schema includes: executive_summary, all_section_outputs, devils_advocate_critiques (if available), ceo_decision_context (why is the CEO asking for a Council review? is this a go/no-go decision or a prioritization decision?), funding_at_stake (if this is a funding decision, how much capital is being committed?), confidence_threshold (what level of confidence is required to proceed?).

The output schema is a Pydantic model with: council_transcript (list of turns, each with persona, statement, vote_change), final_votes (dict mapping persona to their final vote), vote_tally (Approve count, Revise count, Reject count), consensus_points (list of things all personas agreed on), contention_points (list of things personas disagreed about), majority_opinion (summary of the majority position), minority_opinion (summary of the minority position, if any), conditions (list of specific validation steps or changes required before proceeding), overall_recommendation (Proceed / Revise / Do Not Proceed), recommendation_confidence (high/medium/low), assumptions_used, uncertainties, input_tokens, output_tokens.

The agent uses Claude Sonnet for all personas because the deliberation requires nuanced reasoning, counterfactual thinking, and the ability to update beliefs based on new arguments. Haiku cannot do this reliably.

The prompt engineering is the most complex in the entire system. Each persona has a detailed system prompt that defines their identity, priorities, risk tolerance, and reasoning style. The Operator is told 'You are a COO with 15 years of experience building SaaS products. You care deeply about execution risk. You are skeptical of ambitious timelines and optimistic assumptions about team productivity. You approve plans that are realistic and achievable. You reject plans that require heroic execution.'

The deliberation is structured as a debate, not a consensus-seeking discussion. Personas are encouraged to disagree and push back on each other. The prompt says 'Your job is not to agree with the other personas. Your job is to defend your perspective and challenge perspectives you find unconvincing.'

The system uses a round-robin discussion order to ensure every persona speaks in each turn. Turn 1: Operator, Investor, Strategist, Risk Manager, Visionary. Turn 2: Visionary, Risk Manager, Strategist, Investor, Operator (reversed order to prevent the same persona always speaking last).

After the final vote, a synthesis step consolidates the transcript into a readable recommendation. This synthesis is also done by Claude Sonnet, with a prompt that says 'You are summarizing a Council deliberation. Extract the key points of agreement, key points of disagreement, and the overall recommendation. Be concise but preserve the reasoning.'

Error handling: if any persona's LLM call fails (timeout, API error), that persona is marked as 'unavailable' and the Council proceeds with 4 personas instead of 5. If 3 or more personas fail, the Council aborts and escalates to human review.""",

        "How It Connects to Other Agents": """The Council Agent is the final gate in the Phase 2 pipeline (and later in Phase 3). It runs after all other agents have completed, after the Devil's Advocate has reviewed high-stakes sections, and after the Summary Agent has synthesized the executive summary.

The execution order is: (1) All section agents (1, 3, 4, 5, 6.5, 8, 10, 12, 13, 14) complete. (2) Devil's Advocate reviews sections 8, 12, 14. (3) Summary Agent produces executive summary. (4) Mother Agent checks if Council review is required (based on funding amount, confidence score, or CEO request). (5) If required, Council convenes and produces deliberation transcript and final recommendation. (6) Council output is presented to CEO as the final decision point.

The Council reads outputs from all other agents. It has access to the full prior_outputs dictionary (all 11 section outputs), the executive summary, and the Devil's Advocate critiques. Each persona selectively reads what matters for their lens. The Investor focuses on Section 12 (Financial). The Operator focuses on Sections 4 (Org Design), 6.5 (Tech Stack), and 10 (Operations). The Market Strategist focuses on Sections 1 (Opportunity), 3 (Market Research), and 8 (Marketing).

The Council does not send its output to other agents. It is the terminal node for agent deliberation. The output is consumed by the CEO. The Council transcript is displayed in the Streamlit UI, exported to a Word document, and sent via Telegram as a formatted summary.

The Council also integrates with the decision log. Every Council recommendation is logged with the vote tally, the reasoning, and the conditions. If the CEO proceeds against a Reject recommendation, that is logged as a 'CEO override'. If the CEO proceeds with a conditional Approve and later finds that the conditions were not met, the system can trace that back to the Council recommendation.

In Phase 3 (future), the Council will also interact with the Validation Agent (not yet built). If the Council sets conditions like 'validate CAC below $600', the Validation Agent is responsible for tracking whether that validation happened, what the result was, and whether the condition was satisfied.

The Council is also the input to the final human review step. If the Council vote is split (3-2 or 2-2-1), the system flags the decision as contentious and requires human sign-off. The human reviewer reads the Council transcript, sees the arguments on both sides, and makes the final call.""",

        "Why This Matters (Real-World Impact)": """The Council Agent is the reason the multi-agent system can handle truly high-stakes decisions. A CEO considering whether to quit their job and commit $200K of personal savings to a startup idea should not rely on a single AI agent's recommendation. They need rigorous, multi-perspective deliberation. The Council provides that.

In real-world usage, the Council output is what the CEO shows to co-founders, investors, and advisors. Instead of saying 'the AI said this is a good idea', the CEO can say 'we ran this through a simulated advisory board, here are the 5 perspectives that were considered, 4 out of 5 approved with these conditions, here are the key risks that were debated, and here is the minority opinion that raised concerns.' That level of rigor is credible.

The Council also protects the CEO from their own biases. If the CEO is emotionally attached to a business idea, they might ignore warning signs. But if the Council votes 3 Reject, 2 Approve, and the majority opinion cites unvalidated assumptions and high competitive risk, the CEO has to confront that feedback. The Council is an accountability mechanism.

The deliberation transcript is particularly valuable. It is not just a black-box recommendation. The CEO can read the reasoning, see what questions were asked, see how different perspectives challenged each other, and understand why the Council reached its conclusion. This transparency builds trust in the system.

The Council also enables better decision-making under uncertainty. When facing a split vote (3-2 or 2-2-1), the CEO can see exactly what the disagreement is about. If the Investor and Operator approve but the Risk Manager and Strategist reject, the CEO knows the plan is operationally and financially sound but has market risk. That insight helps the CEO make an informed choice about which risks to take.

The conditions mechanism is a major practical benefit. Instead of a binary 'yes/no' recommendation, the Council can say 'yes, if these 3 conditions are met.' This gives the CEO a clear action plan. They do not have to guess what needs to be validated. The Council tells them explicitly: 'Run a 4-week paid ads pilot to validate CAC. Get 10 LOIs from target customers. Confirm GDPR compliance with a legal review.'

The Council also improves the CEO's business judgment over time. By reading multiple Council deliberations on different business ideas, the CEO learns what questions experienced operators, investors, and strategists ask. They learn what 'good evidence' looks like. They learn how to spot weak assumptions. This is educational, not just operational.

Finally, the Council makes the multi-agent system future-proof for autonomous decision-making. In Phase 3, we plan to allow the system to autonomously approve low-stakes decisions (like small budget adjustments or tactical pivots). The Council is the deliberation engine that will power that autonomous judgment. By building and testing it now on high-stakes decisions with human oversight, we are laying the groundwork for safely delegating decisions to AI in the future.""",
    }

    output_path = output_dir / "council_agent_explanation.docx"
    create_detailed_explanation_doc(
        "council_agent",
        "Council Agent - Multi-Persona Business Plan Deliberation System",
        sections,
        output_path,
    )


def main() -> None:
    """Generate detailed explanations for the final 3 agents."""
    output_dir = Path("/home/saiaditya26122006/multi-agent-system/explaination")
    output_dir.mkdir(exist_ok=True)

    logger.info("Generating detailed explanations for final 3 agents...")

    generate_summary_agent_explanation(output_dir)
    generate_devils_advocate_explanation(output_dir)
    generate_council_explanation(output_dir)

    logger.info("\n✅ All 3 agent explanations generated successfully!")
    logger.info(f"Total files in {output_dir}: {len(list(output_dir.glob('*.docx')))}")


if __name__ == "__main__":
    main()
