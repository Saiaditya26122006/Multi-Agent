"""
Generate remaining agent explanations (Part 2).
"""

import os
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH


def add_styled_heading(doc, text, level=1):
    """Add a heading with custom styling."""
    heading = doc.add_heading(text, level=level)
    heading.style.font.color.rgb = RGBColor(31, 78, 121)
    return heading


def add_styled_paragraph(doc, text, bold=False, italic=False):
    """Add a paragraph with optional styling."""
    para = doc.add_paragraph(text)
    if bold:
        para.runs[0].bold = True
    if italic:
        para.runs[0].italic = True
    para.style.font.size = Pt(11)
    return para


# Remaining agent explanations
MORE_AGENT_EXPLANATIONS = {
    "environment_research": {
        "title": "Environment Research Agent",
        "purpose": """
The Environment Research agent looks outside your company. While other agents focus on your product, team, and finances, this one asks: What external forces could help or hurt you? What's happening in the broader market, regulatory landscape, and competitive environment?

It runs two classic frameworks: PEST analysis (Political, Economic, Social, Technological forces) and Porter's Five Forces (competitive intensity analysis). But it doesn't just fill in templates. It connects external factors to your specific business.
        """,
        "why_written": """
Most business plans treat the environment as static background. They say "the market is growing" without explaining why or how fast. They ignore regulatory changes until they become problems.

This agent makes external analysis actionable. It doesn't just say "GDPR is a political factor." It says "GDPR Article 44 requires EU data residency for EU customer data. If you're selling to European universities and hosting in US regions, you cannot legally operate. Impact: CRITICAL. Mitigation: Use EU cloud regions only."

That's the difference between academic exercise and business intelligence.
        """,
        "key_features": [
            ("PEST Analysis with Specificity", "Not generic 'political risk exists' but 'EU AI Act (effective 2025) classifies academic validation tools as limited-risk systems requiring transparency documentation.' Lists specific regulations, dates, compliance requirements."),
            ("Porter's Five Forces with Ratings", "Each force gets a rating (high/medium/low intensity) plus evidence. 'Threat of new entrants: LOW because academic domain expertise creates 18-24 month barrier to competitive feature parity.'"),
            ("Risk-Opportunity Pairing", "Every identified risk includes potential mitigation. Every opportunity includes exploitation strategy. Not just 'regulatory risk exists' but 'risk + here's how to address it.'"),
            ("Sector-Specific Analysis", "Adapts to business type. B2B SaaS gets different environmental scan than manufacturing or marketplace. Academic SaaS sees different forces than fintech SaaS."),
        ],
        "technical_approach": """
This agent uses Claude Haiku (faster model) because environmental analysis is more structured than strategy work. The SYSTEM_PROMPT encodes the frameworks with guidance on depth and specificity.

It searches for live market data: "EU AI Act academic research software compliance 2025" and "GDPR SaaS academic procurement requirements Europe." This grounds the analysis in current regulations, not outdated training data.

The reasoning engine still runs (decompose, produce, challenge, revise) but the challenge phase focuses on: Are these factors material to THIS business? Is the evidence current? Are the ratings justified?
        """,
        "integration": """
DEPENDS ON:
- Section 1: Business type, target market, competitive strategy (determines which environmental factors matter)

FEEDS INTO:
- Section 5: SWOT synthesis (external opportunities and threats come from this section)
- Section 8: Marketing strategy (regulatory constraints affect channel choices)
- Section 10: Operations (compliance requirements affect process design)

If this agent identifies "GDPR requires EU data residency" but tech stack puts servers in us-east-1, there's a contradiction. The coherence audit will catch that.
        """,
        "why_it_matters": """
Here's an example of why this matters:

Your business plan says: "We'll sell academic manuscript validation software to European universities."

WITHOUT this agent: You build for 6 months, launch, approach first customer. Their procurement team asks: "Are you GDPR compliant? Where's your data stored? Do you have a DPO appointed?" You hadn't thought about it. Deal blocked. You spend the next 3 months retrofitting compliance.

WITH this agent (Month 0): "PEST analysis identifies: GDPR Article 44 (political factor) requires EU data residency. EU AI Act (effective 2025, technological/political factor) requires transparency in AI systems. Impact: CRITICAL for go-to-market. Mitigation: (1) Host exclusively in eu-west-1, (2) Document validation logic for AI Act transparency requirements, (3) Appoint DPO if handling >250 employees or sensitive data."

Now you build correctly from day one. No retrofit. No blocked deals.

That's the value: surfacing external constraints before they become expensive surprises.
        """
    },

    "organisation_designer": {
        "title": "Organisation Designer Agent",
        "purpose": """
The Organisation Designer agent builds your team structure and hiring plan. Not just "we'll hire 5 people." It defines: What roles? With what skills? Hired when? At what cost? What gaps exist today? What hiring risks could delay execution?

It ties the headcount plan directly to revenue milestones. You don't hire Engineer #3 "eventually." You hire them at Month 6 when customer count hits 10, because that's when support load becomes unsustainable for the founders.
        """,
        "why_written": """
We built this because team planning in most business plans is vague. They show an org chart with boxes and titles. They say "Year 1: 3 people, Year 2: 5 people, Year 3: 8 people." But they don't explain:

What does each person do?
What happens if you can't hire that person?
What's the cost (salary plus benefits plus overhead)?
What skills are missing that could block execution?

This agent forces those questions. It says "Year 1 you need: CTO (technical co-founder, Month 0), ML Engineer (hired Month 3, cost $80K fully loaded), Customer Success Specialist (hired Month 9 when customer count hits 15)." Every role has a trigger condition.
        """,
        "key_features": [
            ("Roles Tied to Milestones", "Not 'hire someone eventually' but 'hire Customer Success at Month 9 (when 15 customers create 40hrs/week support load exceeding founder capacity).'"),
            ("Capability Gap Analysis", "Identifies missing skills that could block execution. Rates severity (critical/high/medium). 'No in-house ML expertise: CRITICAL gap because product depends on LLM integration. Resolution: Hire senior ML engineer by Month 3 or outsource to consultant.'"),
            ("Fully-Loaded Cost Calculation", "Not just base salary. Includes: salary + benefits (20-30%) + employer taxes + equipment + software licenses. $60K salary becomes $80K fully loaded."),
            ("Headcount Plan by Year", "Shows count and cost for Year 1/2/3. Tied to revenue milestones: 'Year 2 headcount increases to 5 when ARR hits $500K (enables hiring 2 sales reps at 10% of revenue).'"),
            ("Knowledge Gaps and Mitigation", "Not just skill gaps but knowledge gaps: 'No team member has academic domain expertise. Risk: building features researchers don't need. Mitigation: Advisory board with 2-3 professors, quarterly feedback sessions.'"),
        ],
        "technical_approach": """
This agent uses Claude Haiku with the Intelligence Engine. It's more constrained than strategy agents because org design follows patterns:

Early-stage SaaS: 2 founders + 1-2 engineers Year 1. Add sales/marketing Year 2 when PMF proven.
B2B: Need dedicated customer success earlier (customers expect support).
Technical product: Need strong engineering lead early. Marketing hire comes later.

The agent applies these heuristics but customizes to the specific business. For academic SaaS: "Need academic domain expertise for credibility with buyers. Either: hire PhD as advisor, or co-founder must have academic background."

Cost calculation is programmatic: base_salary × 1.25 (benefits) × 1.05 (equipment) = fully loaded cost.
        """,
        "integration": """
DEPENDS ON:
- Section 1: Product type, target market (determines required skills)
- Section 8: Sales process, customer volume (determines when to hire sales/CS)

FEEDS INTO:
- Section 6.5: Tech stack (team size affects build vs buy decisions)
- Section 10: Operations (headcount affects capacity planning)
- Section 12: Financial model (personnel costs are largest expense in most startups)

If this agent says "hire ML engineer at $80K Month 3" but financial model shows $30K burn rate in Month 3, there's a cash flow problem. Financial agent will catch that.

If this agent says "no domain expertise on team" but marketing targets niche academic buyers, marketing will struggle. SWOT should flag this as weakness.
        """,
        "why_it_matters": """
Let me show you the difference:

WEAK org planning:
- Year 1: 3 people (founders + 1 hire)
- Year 2: 5 people
- Year 3: 8 people
- Total cost Year 1: $200K

STRONG org planning (this agent):
- Month 0-3: 2 co-founders (CEO + CTO), equity only, burn $15K/month for tools/cloud
- Month 3: Hire ML Engineer ($80K/year = $6.7K/month) when LLM integration work exceeds founder capacity. Critical hire: blocks product launch if delayed.
- Month 9: Hire Customer Success Specialist ($50K/year = $4.2K/month) when customer count hits 15 (support load = 20hrs/week, exceeds founder capacity).
- Year 1 total personnel cost: $100K (pro-rated for part-year hires) + $15K founders' expenses = $115K
- Year 2: Add Sales Rep ($60K + $40K commission) and Marketing Specialist ($55K) when ARR hits $300K
- Year 2 total: $345K (5 people, mix of full-time and commission-based)

- Capability gaps: (1) No academic domain knowledge CRITICAL → Mitigation: form advisor network with 3 professors, meet quarterly. (2) No design expertise HIGH → Mitigation: use no-code tools (Webflow) + contractor for brand.

Second plan is 10x more actionable. When you hit Month 2 and your co-founder is drowning in LLM integration work, you know: hire the ML engineer now, it's budgeted, it's critical. You don't have vague "we need more people" feeling.

When you hit Month 8 with 12 customers and support is taking 15hrs/week, you know: hire CS in 1 month (at 15 customers), it's budgeted.

This agent builds that clarity. It turns "hire people" into a playbook: who, when, why, how much.
        """
    },

    "swot_synthesizer": {
        "title": "SWOT Synthesizer Agent",
        "purpose": """
The SWOT Synthesizer brings together internal capabilities (from org design) and external environment (from PEST and Five Forces) into one coherent strategic picture. It's the bridge between "what's happening out there" and "what we can do about it."

SWOT isn't just four boxes. It's a diagnosis of strategic position plus actionable implications. Strength + Opportunity = offensive strategy. Weakness + Threat = defensive priority. The synthesis is what matters.
        """,
        "why_written": """
We built this because SWOT analysis is usually shallow. Plans list generic strengths ("strong team," "innovative product") and generic threats ("competition," "market risk"). There's no synthesis, no prioritization, no action plan.

This agent forces depth and connection:

Strength isn't "strong technical team" but "Strong: CTO has 10 years LLM experience, built similar validation system at Google. Implication: Can build MVP in 3 months vs 6-9 for competitors without domain expertise."

Opportunity isn't "growing market" but "Opportunity: EU AI Act requires transparency in academic AI tools (effective 2025), no incumbents have compliance documentation ready. Implication: Early mover advantage if we launch with compliance built in."

Then it synthesizes: "Strategic priority: Leverage CTO's domain expertise (strength) to build compliance-first product (opportunity) before competitors wake up."

That's actionable strategy, not list-making.
        """,
        "key_features": [
            ("Connected SWOT Elements", "Every strength/weakness connects to specific capability from Section 4. Every opportunity/threat connects to specific factor from Section 3. No generic 'innovation' or 'competition' entries."),
            ("Strategic Implications", "After the four boxes, adds 'Strategic Implications' section: What does this SWOT tell you to do? Which strengths to leverage? Which weaknesses to address first? Which opportunities to chase?"),
            ("Priority Issues", "Not all SWOT items matter equally. Agent ranks: 'P0 issues (critical): Address technical team domain gap before launch. P1 issues (important): Build GDPR compliance before approaching EU customers.'"),
            ("Cross-References", "SWOT elements reference specific sections: 'Strength: Technical co-founder has ML expertise (from Section 4, role: CTO).' This proves the SWOT is grounded, not invented."),
        ],
        "technical_approach": """
This agent uses Claude Sonnet (strategic synthesis requires stronger reasoning) with careful prompt engineering:

1. Load outputs from Section 3 (environment) and Section 4 (org design)
2. Extract: What external factors are HIGH impact? What internal capabilities are CRITICAL?
3. Map external to internal: Which opportunities match our strengths? Which threats exploit our weaknesses?
4. Synthesize: What's the strategic game plan given this configuration?

The Intelligence Engine's "challenge" phase is critical here. It asks: Is this SWOT item generic or specific? Is it actionable or abstract? Is it supported by prior sections or invented?

If agent produces "Strength: innovative product," Devil's Advocate says "What specifically is innovative? How is that connected to team capabilities from Section 4?" Agent revises to something specific.
        """,
        "integration": """
DEPENDS ON:
- Section 3: PEST analysis, Five Forces (external opportunities and threats)
- Section 4: Org design, capability gaps (internal strengths and weaknesses)

FEEDS INTO:
- Section 8: Marketing strategy (strategic priorities inform positioning and channel choices)
- Section 13: Launch plan (priority issues become pre-launch requirements)

This is the strategic midpoint of the plan. Everything before it (Sections 1-4) is diagnosis. Everything after it (Sections 8-14) is execution. This section is the hinge: "Given this diagnosis, here's the strategy."

If this section produces weak synthesis, the rest of the plan lacks strategic coherence. Marketing might pursue opportunities that don't match your strengths. Operations might ignore critical weaknesses.
        """,
        "why_it_matters": """
Here's what good SWOT synthesis does:

EXAMPLE INPUT (from Sections 3 and 4):
- Section 3 (Environment): GDPR creates compliance barrier for US competitors. EU AI Act requires transparency. Universities have 12-18 month procurement cycles.
- Section 4 (Org): CTO has academic background + ML expertise. No sales team. No domain experts on go-to-market.

WEAK SWOT:
Strengths: Technical team, good product
Weaknesses: Limited resources
Opportunities: Growing market
Threats: Competition
Implications: Build good product and sell it.

STRONG SWOT (this agent):
Strengths:
- S1: CTO has dual advantage: ML expertise (can build LLM features competitors can't replicate quickly) + academic background (credibility with research buyers). Source: Section 4.
- S2: EU-based startup can offer native GDPR compliance (data never leaves EU). US competitors must retrofit. Source: Section 3, GDPR factor.

Weaknesses:
- W1: No sales team or enterprise sales experience. Risk: long university procurement cycles (12-18 months, Section 3) could lead to cash-out before first revenue. Severity: CRITICAL.
- W2: No academic domain expert on go-to-market team. Risk: building features researchers don't need. Severity: HIGH.

Opportunities:
- O1: EU AI Act (2025) requires transparency documentation for academic AI tools. No current incumbents have this ready. First mover advantage if we ship compliance-ready. Source: Section 3, regulatory analysis.
- O2: COVID accelerated shift to remote research validation. Universities budget for digital research tools increased 40% 2020-2024. Source: Section 3, economic trends.

Threats:
- T1: University procurement cycles (12-18 months) create cash flow risk in Year 1. Mitigation needed: faster revenue channel or extend runway. Source: Section 3, institutional buying behavior.
- T2: If Tur nitin or Grammarly add epistemic validation features, they have distribution advantage (already in universities). Timing: expect competitive response 18-24 months after our launch. Source: Section 3, competitive analysis.

Strategic Implications:
P0 (Critical): Address W1 (no sales capacity) by either (a) hiring enterprise sales rep with university experience Month 6, or (b) partnering with existing academic software vendor for distribution. Financial model must account for 12-month sales cycle.

P1 (Important): Leverage S1 + O1: Position as "compliance-first" epistemic validation platform. Ship with built-in EU AI Act documentation before US competitors catch up. Marketing should lead with "GDPR-native, AI Act compliant."

P2 (Important): Mitigate W2: Form academic advisory board (3 professors from target schools) by Month 3. Quarterly product feedback sessions. Budget €5K/year per advisor.

P3 (Monitor): Watch T2 (competitive response). If Turn itin announces validation features, timing shifts from "18-month window" to "race mode."

See the difference? Second version is a strategic roadmap. It tells you: What to do first (address W1), how to position (leverage S1+O1), what to watch (T2).

That's what synthesis means. This agent builds that synthesis.
        """
    },

    "operations": {
        "title": "Operations Agent",
        "purpose": """
The Operations agent designs how you actually deliver the product or service. For SaaS: What's the architecture? What's the delivery process when a customer signs up? Where are the bottlenecks at 2x scale? At 10x scale?

More importantly, it builds the cost structure. Fixed costs (rent, salaries, software licenses). Variable costs (hosting per customer, support time per user, payment processing fees). COGS (cost of goods sold) per unit.

These costs flow directly into the financial model. If operations says "COGS is $20 per customer per month" but marketing says "price is $50/month," finance calculates: $30 gross margin, 60% margin. That's good for SaaS. If COGS were $45, margin would be 10%. That's broken.

Operations is where pricing meets reality.
        """,
        "why_written": """
We built this because cost structure is usually an afterthought. Plans say "we'll operate efficiently" without defining what that means. They don't identify bottlenecks until they hit them in production.

This agent forces operational thinking:

At 10 customers: Founders can manually onboard, support via email, no automation needed.
At 100 customers: Need automated onboarding flow, ticketing system, dedicated CS person.
At 1000 customers: Need tiered support, self-service knowledge base, SLA tracking.

If you don't plan for these transitions, you break at scale. The agent identifies the inflection points upfront.
        """,
        "key_features": [
            ("Production/Delivery Process", "Step-by-step: How does a customer go from signup to value? For SaaS: Signup → Onboarding email → Product tour → First validation → Success milestone. Each step is explicit."),
            ("Cost Structure with Dollar Amounts", "Not 'hosting costs will be low' but 'Fixed costs: $850/month cloud hosting (AWS eu-west-1 for 2 instances + RDS). Variable costs: $5/customer/month (pro-rated compute + storage). COGS: $8/customer/month (LLM API calls for validation).'"),
            ("Capacity Planning with Bottleneck Analysis", "At 2x scale (20 customers to 40): No new bottlenecks, same infrastructure handles it. At 5x scale (100 customers): Need dedicated database instance, estimated cost +$300/month. At 10x scale (200 customers): Need load balancer + auto-scaling, support becomes full-time role."),
            ("Supplier Strategy", "For critical dependencies: 'LLM provider: Claude via AWS Bedrock. Vendor lock-in risk: MEDIUM. Mitigation: Abstraction layer allows switch to GPT-4 if needed. Expected effort: 2 weeks development.'"),
            ("Quality/Risk Analysis", "What could break? 'Single region deployment: If eu-west-1 has outage, service is down. Mitigation: Multi-region backup costs +$400/month, deferred until 50+ customers.'"),
        ],
        "technical_approach": """
This agent uses Claude Haiku (operational planning is more structured than strategy). It has domain knowledge encoded:

SaaS typical costs: Cloud hosting $500-2K/month (scales with users), support $50K/year per CS rep (handles ~50-100 customers), LLM APIs $0.01-0.10 per customer interaction.

Manufacturing typical costs: Different structure (tooling, materials, logistics).

The agent scales these templates to the specific business. For academic SaaS: "Validation requests per customer = 10/month. Tokens per validation = 50K. Claude Haiku cost = $0.25/1M tokens. Cost per customer = (10 × 50K / 1M) × $0.25 = $0.125/month. Round to $1/month with buffer for retries."

That's realistic costing, not guesswork.
        """,
        "integration": """
DEPENDS ON:
- Section 1: Product type (SaaS vs marketplace vs manufacturing)
- Section 4: Headcount (personnel costs are fixed costs)
- Section 8: Expected customer volume (determines variable costs)

FEEDS INTO:
- Section 12: Financial model (cost_structure becomes the 'costs' side of the P&L)
- Section 6.5: Tech stack (capacity constraints inform infrastructure sizing)

If this agent says "COGS is $15/customer" but marketing prices at $20/customer, that's 25% gross margin. For SaaS that's too low (should be 70-85%). Financial model will show the business can't scale profitably.

This agent is the reality check on pricing. It says "Given how we actually deliver this, here's what it costs." Marketing must price above that floor.
        """,
        "why_it_matters": """
Here's a real scenario:

You're building academic manuscript validation SaaS. Pricing is $5K/year per university.

WITHOUT operations analysis: You assume "cloud costs will be like $500/month, margins are great!"

Month 3: First customer signs up. They upload 50 manuscripts for validation. Each validation runs Claude Sonnet (expensive model) with 200K tokens (your prompts are verbose). Cost: 50 × 200K tokens × $3/1M tokens = $30 for that batch. Customer does this 10x per year. Your cost: $300/year. Your price: $5K. Margin: 94%. Great!

Month 6: Five customers now. One customer uploads 500 manuscripts (10x your assumption). Your cloud bill spikes to $2K that month. You investigate: their use case is different (validating entire dept's research, not just doctoral students). Your pricing doesn't account for usage tiers.

Month 9: Ten customers. Cloud bill: $4K/month. Your revenue: $50K/year = $4.2K/month. You're barely breaking even on gross margins before counting salaries. Business model is broken.

WITH operations analysis (this agent, Month 0):
"Delivery process: Customer uploads manuscript → LLM validation runs (estimated 50K tokens per manuscript at Claude Haiku $0.25/1M = $0.0125 per manuscript). Usage assumptions: 10 manuscripts per customer per month (validated with pilot).

Variable costs per customer: $0.0125 × 10 = $0.125/month (LLM) + $5/month (pro-rated compute/storage) = $5.125/month = $61.50/year.

COGS per customer: $61.50/year.
Price per customer: $5,000/year.
Gross margin: ($5,000 - $61.50) / $5,000 = 98.8%.

RISK: Usage assumption is based on pilot with 2 customers. If actual usage is 10x higher (100 manuscripts/month), COGS becomes $615/year. Margin drops to 87.7% (still healthy for SaaS).

MITIGATION: Implement usage tiers. Base tier: 100 manuscripts/year. Overage: $50 per 100 additional manuscripts. This protects margin while accommodating high-usage customers."

Second analysis catches the risk before launch. You build usage tiers into pricing. When customer uploads 500 manuscripts, they're on Enterprise tier at $10K/year, not Base tier. Your margins stay healthy.

That's what operations planning does. It forces you to think through the delivery mechanics and cost realities before they surprise you in production.
        """
    },

    "launch_contingency": {
        "title": "Launch and Contingency Planning Agent",
        "purpose": """
The Launch agent builds your startup program: the actual sequence of milestones from idea to first revenue to break-even. Not vague "launch in Q3" but "Month 3: Complete MVP, Month 6: First pilot customer, Month 9: First paid customer, Month 12: Break-even."

More importantly, it defines contingency scenarios. What if you don't hit Month 6 pilot milestone? What if first customer churns immediately? What if CAC is 2x your estimate? The agent sets observable triggers and pre-planned responses.

This is the "what if things go wrong" section. Most plans ignore it. This agent makes it explicit.
        """,
        "why_written": """
We built this because launch planning is usually optimistic timeline projection. "Month 1: MVP, Month 3: Launch, Month 6: Product-market fit, Month 12: Scale!" No account for delays, setbacks, or plan B.

This agent forces realism:

What are the prerequisite conditions for launch? (Need 5 pilot customers validated, GDPR compliance ready, pricing tested with 3 prospects.)

What's the critical path item? (Customer validation. If we can't get pilots to sign up, everything else is moot.)

What's the exit condition if fundamentals don't work? (If <$50K MRR by Month 18 and <6 months runway, initiate wind-down.)

That's hard to write. But it's essential. You need to know ahead of time: What does failure look like? At what point do you pivot vs persist?
        """,
        "key_features": [
            ("Milestone Timeline", "Not just dates but dependencies. 'Month 6: First pilot customer (depends on: MVP complete, 20 sales outreach conversations, pricing validated).' Shows what needs to happen before each milestone."),
            ("Prerequisite Conditions", "What must be true before you can launch? 'Prerequisites: (1) GDPR compliance documentation complete, (2) 3 pilot customers validated product, (3) CAC tested at <$2K with 10+ customer conversations.'"),
            ("Critical Path Item", "The one thing that, if it fails, kills the business. Usually it's customer validation. 'Critical path: Getting first 10 pilot customers to sign up. If we can't get pilots, we don't have product-market fit. All other work (team, tech, marketing) is wasted until this is proven.'"),
            ("Contingency Scenarios", "Observable triggers + planned responses. 'If CAC exceeds $2K by Month 6, pivot to partner channel (approach university consortiums for group deals instead of direct sales).'"),
            ("Exit Conditions", "Quantitative wind-down triggers. 'Exit if: (1) <$50K MRR by Month 18 and <6 months runway, OR (2) 3 consecutive quarters of negative growth with <6 months runway.'"),
        ],
        "technical_approach": """
This agent uses Claude Haiku with the Intelligence Engine. It synthesizes across all prior sections:

From Section 1: Objectives become milestones (if objective is '50 customers Year 1,' milestone is 'Month 12: Reach 50 customers').
From Section 8: Launch conditions include marketing prerequisites (validated CAC, tested channels).
From Section 12: Exit conditions reference financial model (break-even month, runway calculations).

The agent also integrates SimPy simulation results. P10 scenario (pessimistic case from financial model) becomes the contingency scenario. "If we hit P10: Break-even extends to Month 24, need $150K additional capital, trigger is '<$100K ARR by Month 12.'"
        """,
        "integration": """
DEPENDS ON:
- Section 1: Objectives (become milestones)
- Section 8: Marketing plan (determines customer acquisition milestones)
- Section 12: Financial projections (determines funding milestones and exit triggers)
- Section 14: Exit strategy (wind-down procedures come from exit conditions)

FEEDS INTO:
- Section 14: Exit strategy (contingency scenarios inform pivot/wind-down triggers)
- Executive Summary: Uses critical path and contingency as risk factors

This is one of the last sections because it needs inputs from almost everything else. You can't define milestones until you know what you're building (Sections 1-8). You can't set exit triggers until you know the financial model (Section 12).
        """,
        "why_it_matters": """
Here's why explicit contingency planning matters:

SCENARIO 1 (no contingency plan):
Month 6: You've spent $60K (runway from $100K seed money). You have 5 pilot customers. CAC is tracking at $2.5K (50% higher than plan).

Founder conversation:
"We're burning through cash faster than expected. Should we raise more? Cut costs? Pivot?"
"I don't know. Let's see how next month goes."

They drift for 2 months. Month 8: $20K left, no clarity. Panic mode. Scramble to raise emergency funding. Fail. Shut down.

SCENARIO 2 (with contingency plan from this agent):
Month 6: Same situation. $60K spent, 5 pilots, CAC $2.5K.

Founder conversation:
"We hit the contingency trigger: CAC exceeded $2K at Month 6. Our plan says: pivot to partner channel. Let's execute."

They know what to do because they thought it through at Month 0. They approach 3 university consortiums with group deal proposals. Month 8: One consortium signs, covers 10 universities, €40K contract. Runway extended.

The business survives because they had a pre-planned response to a foreseeable risk.

That's what this agent does. It makes you think through "what if X goes wrong?" while you still have time and money to adapt. Not when you're in crisis mode with 1 month of runway.

It's unglamorous planning. But it's often the difference between startups that survive Year 1 and startups that don't.
        """
    }
}


def generate_more_explanations():
    """Generate Word documents for remaining agents."""
    output_dir = Path("/home/saiaditya26122006/multi-agent-system/explaination")

    for agent_key, content in MORE_AGENT_EXPLANATIONS.items():
        print(f"Generating explanation for {content['title']}...")

        doc = Document()

        # Title
        title = doc.add_heading(content['title'], 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph()

        # Purpose
        add_styled_heading(doc, "What This Agent Does", level=1)
        add_styled_paragraph(doc, content['purpose'].strip())
        doc.add_paragraph()

        # Why written
        add_styled_heading(doc, "Why We Built This Agent", level=1)
        add_styled_paragraph(doc, content['why_written'].strip())
        doc.add_paragraph()

        # Key features
        if 'key_features' in content:
            add_styled_heading(doc, "Key Features", level=1)
            for feature_name, feature_desc in content['key_features']:
                para = doc.add_paragraph(style='List Bullet')
                para.add_run(feature_name + ": ").bold = True
                para.add_run(feature_desc)
        doc.add_paragraph()

        # Technical approach
        if 'technical_approach' in content:
            add_styled_heading(doc, "How It Works Technically", level=1)
            add_styled_paragraph(doc, content['technical_approach'].strip())
        doc.add_paragraph()

        # Integration
        if 'integration' in content:
            add_styled_heading(doc, "How It Connects to Other Agents", level=1)
            add_styled_paragraph(doc, content['integration'].strip())
        doc.add_paragraph()

        # Why it matters
        if 'why_it_matters' in content:
            add_styled_heading(doc, "Why This Matters (Real World Impact)", level=1)
            add_styled_paragraph(doc, content['why_it_matters'].strip())

        # Save
        filename = f"{agent_key}_explanation.docx"
        filepath = output_dir / filename
        doc.save(str(filepath))
        print(f"  ✅ Saved: {filename}")

    print(f"\n✅ Part 2 complete in {output_dir}")
    return len(MORE_AGENT_EXPLANATIONS)


if __name__ == "__main__":
    count = generate_more_explanations()
    print(f"\n📄 Generated {count} additional explanation documents")
