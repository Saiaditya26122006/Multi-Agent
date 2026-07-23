"""Build v2 — per-section invocation (Phase 1) + on-demand assembly.

"Work on section X" runs just that section's agent (installment model), instead
of the whole batch pipeline. Reuses the existing orchestrator agent machinery
(instantiate/register/dispatch over the MessageBus) but for ONE section, writing
the result into the section-state backbone.
"""

import asyncio
import json
import logging
import threading
import uuid
from typing import Optional

from services import section_state

logger = logging.getLogger(__name__)


def run_section(session_id: str, section_id: str, force: bool = False,
                feedback: Optional[str] = None, focus: Optional[str] = None) -> dict:
    """Start one section's agent in the background. Returns immediately.

    Returns {status, section_id, [reason]} where status is:
      - "started"  — agent dispatched
      - "blocked"  — dependencies not done (unless force=True)
      - "unknown_section"
    `feedback` (Adjust flow) steers a revision; `focus` (kickoff tap-clarify)
    tells a fresh run which angle Alex chose.
    """
    registry = section_state.effective_registry(session_id)
    if section_id not in registry:
        return {"status": "unknown_section", "section_id": section_id}

    section_state.init_sections(session_id)
    states = section_state.list_sections(session_id)

    if not force and not section_state.deps_met(section_id, states, registry):
        deps = registry[section_id].get("depends_on", [])
        pending = [d for d in deps if states.get(d, {}).get("status") != "done"]
        return {"status": "blocked", "section_id": section_id,
                "reason": f"waiting on sections {pending}", "pending_deps": pending}

    cur = (states.get(section_id) or {}).get("status")
    # re-open a done/needs_review section for a revision
    section_state.update_section(session_id, section_id,
                                 status="in_progress" if cur != "in_progress" else cur)

    def _worker() -> None:
        try:
            asyncio.run(_run_section_async(session_id, section_id, registry, feedback, focus))
        except Exception as e:  # noqa: BLE001
            logger.exception("[BuildV2] section %s crashed", section_id)
            try:
                section_state.update_section(session_id, section_id,
                                             status="failed", blocked_on=str(e)[:200])
            except Exception:
                pass

    threading.Thread(target=_worker, daemon=True, name=f"section-{section_id}").start()
    return {"status": "started", "section_id": section_id,
            "agent": registry[section_id].get("agent")}


def adjust_section(session_id: str, section_id: str, feedback: str) -> dict:
    """Re-open a done/needs_review section and re-run it with Alex's feedback.

    Replaces the old stubbed 'Adjust' — real section-level revision.
    """
    sec = section_state.get_section(session_id, section_id)
    if not sec:
        return {"status": "unknown_section"}
    if sec.get("status") not in ("done", "needs_review", "failed"):
        return {"status": "not_adjustable", "reason": f"section is {sec.get('status')}"}
    return run_section(session_id, section_id, force=True, feedback=feedback)


def export_plan(session_id: str) -> dict:
    """Compile the current section drafts into one markdown business plan.

    A view over section state (not a re-run). Includes only sections that have a
    draft; marks the rest as pending. DOCX export can adapt evaluation/export_docx.
    """
    section_state.init_sections(session_id)
    registry = section_state.effective_registry(session_id)
    states = section_state.list_sections(session_id)
    lines = ["# Business Plan\n"]
    included = 0
    for sid, meta in registry.items():
        sec = states.get(sid, {})
        draft = sec.get("draft")
        output = draft.get("output") if isinstance(draft, dict) else draft
        lines.append(f"\n## {sid}. {meta.get('title', sid)}\n")
        if output and sec.get("status") in ("done", "needs_review"):
            included += 1
            if isinstance(output, (dict, list)):
                import json as _json
                lines.append("```json\n" + _json.dumps(output, indent=2, ensure_ascii=False)[:4000] + "\n```")
            else:
                lines.append(str(output)[:4000])
        else:
            lines.append(f"_{sec.get('status', 'not_started')} — not yet drafted._")
    return {"markdown": "\n".join(lines), "sections_included": included, "total": len(registry)}


def _docx_render_output(doc, output) -> None:
    """Render a section's structured output into the docx reasonably."""
    if isinstance(output, dict):
        for key, val in output.items():
            if key.startswith("_") or key in ("task_id", "model_used", "input_tokens",
                                              "output_tokens", "section_number", "reasoning_trace"):
                continue
            label = key.replace("_", " ").title()
            if isinstance(val, list):
                doc.add_heading(label, level=3)
                for item in val[:20]:
                    doc.add_paragraph(
                        (json.dumps(item, ensure_ascii=False) if isinstance(item, (dict, list)) else str(item))[:600],
                        style="List Bullet")
            elif isinstance(val, dict):
                doc.add_heading(label, level=3)
                doc.add_paragraph(json.dumps(val, ensure_ascii=False, indent=2)[:1500])
            elif val:
                doc.add_heading(label, level=3)
                doc.add_paragraph(str(val)[:2000])
    else:
        doc.add_paragraph(str(output)[:4000])


def export_docx(session_id: str) -> Optional[bytes]:
    """Compile the current section drafts into a .docx business plan (bytes).

    A view over section state (not a re-run). Returns None if python-docx is
    unavailable. Non-drafted sections are marked pending.
    """
    try:
        import io
        import json  # noqa: F811 (local for clarity)
        from docx import Document
    except Exception as e:  # noqa: BLE001
        logger.warning("[BuildV2] docx export unavailable: %s", e)
        return None

    section_state.init_sections(session_id)
    registry = section_state.effective_registry(session_id)
    states = section_state.list_sections(session_id)

    doc = Document()
    doc.add_heading("Business Plan", 0)
    done = sum(1 for s in states.values() if s.get("status") == "done")
    doc.add_paragraph(f"{done}/{len(registry)} sections finalised · generated by Build v2")

    for sid, meta in registry.items():
        sec = states.get(sid, {})
        draft = sec.get("draft")
        output = draft.get("output") if isinstance(draft, dict) else draft
        doc.add_heading(f"{sid}. {meta.get('title', sid)}", level=1)
        if output and sec.get("status") in ("done", "needs_review"):
            _docx_render_output(doc, output)
            g = (draft or {}).get("grounding") if isinstance(draft, dict) else None
            if g and g.get("rate") is not None:
                doc.add_paragraph(f"[grounding: {int(g['rate']*100)}% of claims evidence-backed]")
        else:
            doc.add_paragraph(f"[{sec.get('status', 'not_started')} — not yet drafted]")

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


async def _run_section_async(session_id: str, section_id: str, registry: dict,
                             feedback: Optional[str] = None,
                             focus: Optional[str] = None) -> None:
    from services.pipeline_orchestrator import get_orchestrator
    from ceo_data.loader import get_relevant_ceo_data

    orch = get_orchestrator()
    agent_name = registry[section_id]["agent"]

    phase1 = orch._read_phase1_session(session_id) or {}
    ceo_data = get_relevant_ceo_data(phase1.get("idea", ""))

    # Prior section drafts become context (blackboard: read shared state).
    states = section_state.list_sections(session_id)
    prior_outputs = {
        sid: s.get("draft")
        for sid, s in states.items()
        if s.get("status") == "done" and s.get("draft")
    }

    await orch._ensure_agents_registered([agent_name])
    input_package = orch._build_input_package(agent_name, phase1, prior_outputs, ceo_data)
    # Tell the agent which section it is writing — required by the generic analyst
    # (custom sections), harmless to the topic-locked specialists (they ignore it).
    input_package["section_id"] = section_id
    input_package["section_title"] = registry.get(section_id, {}).get("title", section_id)
    if feedback:
        # Adjust flow: give the agent Alex's revision note + its prior draft.
        input_package["revision_feedback"] = feedback
        prior = section_state.get_section(session_id, section_id) or {}
        pd = prior.get("draft")
        input_package["prior_draft"] = pd.get("output") if isinstance(pd, dict) else pd
    if focus:
        # Kickoff tap-clarify: the focus angle Alex picked for a fresh run.
        input_package["focus_directive"] = focus

    # Phase 2 of two-phase filing: before generating, refine the suggested leaf
    # of any facts parked at this section's parent (pending_leaf_resolution).
    # Suggestion-only — nothing is moved; the agent just gets better leaf hints
    # in the facts it retrieves. Non-blocking: a failure never stops the build.
    try:
        from services.leaf_resolver import resolve_provisional

        resolve_provisional(section_prefix=section_id)
    except Exception as e:  # noqa: BLE001
        logger.warning(
            "[BuildV2] leaf resolution failed for %s: %s — continuing without refinement",
            section_id, e,
        )

    run_id = f"secrun_{uuid.uuid4().hex[:8]}"

    result = await orch._dispatch_to_agent(agent_name, session_id, run_id, input_package)

    if result:
        # Phase 3: grounding report travels with the draft so Alex sees how
        # evidence-backed it is before accepting. Best-effort — never blocks.
        grounding = None
        try:
            from services.grounding import check_draft
            grounding = check_draft(result, session_id)
        except Exception as e:  # noqa: BLE001
            logger.debug("[BuildV2] grounding skipped for %s: %s", section_id, e)
        # Phase 5: adversarial quality gate for gated sections (Skeptic/Operator).
        critique = None
        try:
            from services.quality_gate import critique_section
            critique = critique_section(section_id, result, session_id)
        except Exception as e:  # noqa: BLE001
            logger.debug("[BuildV2] critique skipped for %s: %s", section_id, e)
        draft = {"output": result, "grounding": grounding, "critique": critique}
        # If the agent fell back for lack of data, the "draft" is a placeholder —
        # park the section as blocked_on_data and raise a TARGETED data request so
        # Feed auto-resumes it (Phase 2 handshake) once Alex provides the evidence.
        used_fallback = isinstance(result, dict) and result.get("_fallback_used")
        if used_fallback:
            section_state.update_section(
                session_id, section_id, status="blocked_on_data", draft=draft,
                blocked_on="agent could only draft a placeholder — provide data via Feed")
            if section_id.isdigit():
                try:
                    from services.data_requests import create
                    title = registry[section_id].get("title", section_id)
                    create(session_id, section_id, [f"BP.{section_id}"],
                           f"Section {section_id} ({title}) needs supporting data — the agent "
                           f"could only produce a placeholder.",
                           why="agent fell back for insufficient evidence", agent=agent_name)
                except Exception as e:  # noqa: BLE001
                    logger.debug("[BuildV2] data request on fallback skipped: %s", e)
            logger.info("[BuildV2] section %s -> blocked_on_data (fallback; data request raised)", section_id)
        else:
            section_state.update_section(session_id, section_id,
                                         status="needs_review", draft=draft, blocked_on=None)
            rate = (grounding or {}).get("rate")
            verdict = (critique or {}).get("verdict")
            logger.info("[BuildV2] section %s -> needs_review (grounding=%s, critique=%s)",
                        section_id, rate, verdict)
    else:
        # No result = the agent escalated (typically weak/missing evidence for
        # its schema). That's a data gap, not a crash — mark blocked_on_data so
        # Alex can feed the missing evidence (which auto-resumes via the Phase 2
        # handshake), rather than a dead-end 'failed'.
        section_state.update_section(
            session_id, section_id, status="blocked_on_data",
            blocked_on="agent needs more evidence — provide supporting data via Feed",
        )
        logger.info("[BuildV2] section %s -> blocked_on_data (agent escalated)", section_id)


def build_all(session_id: str) -> dict:
    """Run every runnable section in dependency order (the 'build the whole plan'
    action, now served by the installment engine instead of the batch pipeline).

    Sections are run in dependency waves — a section runs once all its
    dependencies have a draft (done or needs_review). Each reaches needs_review /
    blocked_on_data; Alex reviews. Returns immediately (runs in the background).
    """
    section_state.init_sections(session_id)

    def _worker() -> None:
        try:
            asyncio.run(_build_all_async(session_id))
        except Exception:  # noqa: BLE001
            logger.exception("[BuildV2] build_all crashed")

    threading.Thread(target=_worker, daemon=True, name="build-all").start()
    return {"status": "started", "message": "Building all sections in dependency order"}


async def _build_all_async(session_id: str) -> None:
    registry = section_state.effective_registry(session_id)
    for _ in range(len(registry) + 5):  # bounded; each pass drafts >=1 or stops
        states = section_state.list_sections(session_id)
        runnable = [
            sid for sid in registry
            if states.get(sid, {}).get("status", "not_started") == "not_started"
            and all(states.get(d, {}).get("status") in ("done", "needs_review")
                    for d in registry[sid].get("depends_on", []))
        ]
        if not runnable:
            break
        for sid in runnable:
            try:
                section_state.update_section(session_id, sid, status="in_progress")
                await _run_section_async(session_id, sid, registry)
            except Exception as e:  # noqa: BLE001
                logger.error("[BuildV2] build_all section %s failed: %s", sid, e)
                try:
                    section_state.update_section(session_id, sid, status="failed",
                                                 blocked_on=str(e)[:200])
                except Exception:
                    pass
    logger.info("[BuildV2] build_all complete for %s", session_id)


def accept_section(session_id: str, section_id: str) -> dict:
    """Alex accepts a needs_review section -> done."""
    sec = section_state.get_section(session_id, section_id)
    if not sec:
        return {"status": "unknown_section"}
    section_state.update_section(session_id, section_id, status="done")
    return {"status": "done", "section_id": section_id}


def get_plan(session_id: str) -> dict:
    """On-demand assembly of the whole plan (done/WIP/blocked/not-started)."""
    section_state.init_sections(session_id)
    return section_state.assemble(session_id)


def quality_report(session_id: str) -> dict:
    """Aggregate build quality across sections (Phase 7 measurement).

    From the stored per-section grounding + critique, reports the metrics that
    actually matter: mean grounding rate, count of ungrounded claims, and how
    many gated sections passed vs need revision.
    """
    section_state.init_sections(session_id)
    states = section_state.list_sections(session_id)
    rates, ungrounded, gated_pass, gated_revise = [], 0, 0, 0
    for s in states.values():
        draft = s.get("draft")
        if not isinstance(draft, dict):
            continue
        g = draft.get("grounding") or {}
        if g.get("available") and g.get("rate") is not None:
            rates.append(g["rate"])
            ungrounded += len(g.get("ungrounded") or [])
        c = draft.get("critique") or {}
        if c.get("gated"):
            if c.get("verdict") == "revise":
                gated_revise += 1
            else:
                gated_pass += 1
    mean_grounding = round(sum(rates) / len(rates), 3) if rates else None
    return {
        "sections_with_grounding": len(rates),
        "mean_grounding_rate": mean_grounding,
        "total_ungrounded_claims": ungrounded,
        "gated_pass": gated_pass,
        "gated_revise": gated_revise,
    }


def next_actions(session_id: str) -> dict:
    """What Alex can do now: ready sections + what's blocked/in-review."""
    section_state.init_sections(session_id)
    states = section_state.list_sections(session_id)
    ready = section_state.ready_sections(session_id)
    review = [sid for sid, s in states.items() if s.get("status") == "needs_review"]
    blocked = {sid: s.get("blocked_on") for sid, s in states.items()
               if s.get("status") == "blocked_on_data"}
    return {"ready": ready, "needs_review": review, "blocked_on_data": blocked}
