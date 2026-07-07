"""
Document extraction — turns an uploaded file (PDF/DOCX/XLSX/TXT/CSV) into
plain text for the Feed pipeline.

Every step emits a trace event via emit_trace() so the frontend's Process
panel can narrate exactly what's happening to *this* file, live, instead of
showing a generic "processing..." spinner.
"""

import csv
import io
import logging
from typing import List

from tools.trace_emitter import emit_trace

logger = logging.getLogger(__name__)

SUPPORTED_EXTENSIONS = {"pdf", "docx", "xlsx", "xls", "txt", "csv"}

# Any single file above this size is rejected before extraction even starts —
# keeps a bad upload from hanging the request indefinitely.
MAX_FILE_SIZE_BYTES = 20 * 1024 * 1024  # 20 MB


class ExtractionError(Exception):
    """Raised when a file is unsupported, too large, or cannot be parsed."""


def extract_text(file_bytes: bytes, filename: str, session_id: str) -> str:
    """Extract plain text from an uploaded file, narrating progress live.

    Args:
        file_bytes: Raw file content.
        filename: Original filename (used to detect format by extension).
        session_id: Session key — trace events are broadcast to this session's
            WebSocket connections so the Process panel can show live status.

    Returns:
        Extracted plain text, ready to hand to the atomic-fact splitter.

    Raises:
        ExtractionError: If the file type is unsupported, the file is too
            large, or parsing fails for any reason.
    """
    ext = filename.rsplit(".", 1)[-1].lower() if "." in filename else ""

    emit_trace(
        session_id, "Feed", "file_received",
        f"Received {filename} ({_human_size(len(file_bytes))})",
        {"filename": filename, "size": len(file_bytes)},
    )

    if len(file_bytes) > MAX_FILE_SIZE_BYTES:
        emit_trace(
            session_id, "Feed", "extraction_failed",
            f"{filename} is too large ({_human_size(len(file_bytes))}, max 20 MB)",
            {"filename": filename},
        )
        raise ExtractionError(f"{filename} exceeds the 20 MB upload limit.")

    if ext not in SUPPORTED_EXTENSIONS:
        emit_trace(
            session_id, "Feed", "extraction_failed",
            f"Unsupported file type: .{ext}",
            {"filename": filename},
        )
        raise ExtractionError(
            f"Unsupported file type: .{ext}. Supported: PDF, DOCX, XLSX, TXT, CSV."
        )

    emit_trace(
        session_id, "Feed", "opening_file", f"Opening {filename}...",
        {"filename": filename, "format": ext},
    )

    try:
        if ext == "pdf":
            text = _extract_pdf(file_bytes, filename, session_id)
        elif ext == "docx":
            text = _extract_docx(file_bytes, filename, session_id)
        elif ext in ("xlsx", "xls"):
            text = _extract_xlsx(file_bytes, filename, session_id)
        elif ext == "csv":
            text = _extract_csv(file_bytes, filename, session_id)
        else:
            text = _extract_txt(file_bytes, filename, session_id)
    except ExtractionError:
        raise
    except Exception as e:
        logger.error("[DocumentExtractor] Failed to extract %s: %s", filename, e)
        emit_trace(
            session_id, "Feed", "extraction_failed",
            f"Couldn't read {filename}: {e}",
            {"filename": filename},
        )
        raise ExtractionError(f"Failed to extract text from {filename}: {e}")

    if not text or not text.strip():
        emit_trace(
            session_id, "Feed", "extraction_empty",
            f"No readable text found in {filename}",
            {"filename": filename},
        )
        raise ExtractionError(f"No readable text found in {filename}.")

    emit_trace(
        session_id, "Feed", "extraction_complete",
        f"Extracted {len(text):,} characters from {filename}",
        {"filename": filename, "char_count": len(text)},
    )

    return text


def _extract_pdf(file_bytes: bytes, filename: str, session_id: str) -> str:
    """Extract text from a PDF, page by page, narrating progress on long docs."""
    from pypdf import PdfReader

    reader = PdfReader(io.BytesIO(file_bytes))
    page_count = len(reader.pages)
    emit_trace(
        session_id, "Feed", "reading_pdf",
        f"Reading {page_count} page(s) from {filename}...",
        {"filename": filename, "pages": page_count},
    )

    pages_text: List[str] = []
    checkpoint = max(1, page_count // 4)
    for i, page in enumerate(reader.pages, start=1):
        pages_text.append(page.extract_text() or "")
        if page_count > 3 and i % checkpoint == 0:
            emit_trace(
                session_id, "Feed", "reading_pdf_progress",
                f"Read page {i}/{page_count} of {filename}...",
                {"filename": filename, "page": i, "total": page_count},
            )

    return "\n\n".join(pages_text)


def _extract_docx(file_bytes: bytes, filename: str, session_id: str) -> str:
    """Extract paragraph and table text from a Word document."""
    import docx

    document = docx.Document(io.BytesIO(file_bytes))
    emit_trace(
        session_id, "Feed", "reading_docx",
        f"Reading paragraphs and tables from {filename}...",
        {"filename": filename},
    )

    parts = [p.text for p in document.paragraphs if p.text.strip()]

    for table in document.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                parts.append(" | ".join(cells))

    return "\n".join(parts)


def _extract_xlsx(file_bytes: bytes, filename: str, session_id: str) -> str:
    """Extract row data from every sheet, treating the first row as headers."""
    import openpyxl

    workbook = openpyxl.load_workbook(io.BytesIO(file_bytes), data_only=True)
    emit_trace(
        session_id, "Feed", "reading_xlsx",
        f"Reading {len(workbook.sheetnames)} sheet(s) from {filename}...",
        {"filename": filename, "sheets": workbook.sheetnames},
    )

    parts: List[str] = []
    for sheet_name in workbook.sheetnames:
        sheet = workbook[sheet_name]
        emit_trace(
            session_id, "Feed", "reading_sheet",
            f"Reading sheet '{sheet_name}' from {filename}...",
            {"filename": filename, "sheet": sheet_name},
        )
        header = None
        for row in sheet.iter_rows(values_only=True):
            cells = [str(c).strip() for c in row if c is not None and str(c).strip()]
            if not cells:
                continue
            if header is None:
                header = cells
                continue
            row_text = "; ".join(
                f"{header[j]}: {c}" if j < len(header) else str(c)
                for j, c in enumerate(cells)
            )
            parts.append(f"[{sheet_name}] {row_text}")

    return "\n".join(parts)


def _extract_csv(file_bytes: bytes, filename: str, session_id: str) -> str:
    """Extract row data from a CSV, treating the first row as headers."""
    emit_trace(
        session_id, "Feed", "reading_csv", f"Reading rows from {filename}...",
        {"filename": filename},
    )

    text = file_bytes.decode("utf-8", errors="replace")
    rows = list(csv.reader(io.StringIO(text)))

    if not rows:
        return ""

    header = rows[0]
    parts = []
    for row in rows[1:]:
        row_text = "; ".join(
            f"{header[j]}: {c}" if j < len(header) else c
            for j, c in enumerate(row) if c.strip()
        )
        if row_text:
            parts.append(row_text)

    return "\n".join(parts)


def _extract_txt(file_bytes: bytes, filename: str, session_id: str) -> str:
    """Decode a plain text file."""
    emit_trace(
        session_id, "Feed", "reading_txt", f"Reading {filename}...",
        {"filename": filename},
    )
    return file_bytes.decode("utf-8", errors="replace")


def _human_size(num_bytes: int) -> str:
    """Format a byte count as a human-readable string (e.g. '1.4 MB')."""
    size = float(num_bytes)
    for unit in ("B", "KB", "MB", "GB"):
        if size < 1024:
            return f"{size:.0f} {unit}" if unit == "B" else f"{size:.1f} {unit}"
        size /= 1024
    return f"{size:.1f} TB"
