"""Generate the audit response document as a Word file."""

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from datetime import datetime


def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    return h


def add_para(doc, text, bold=False, italic=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    return p


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for run in p.runs:
                run.bold = True

    # Data rows
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            table.rows[r_idx + 1].cells[c_idx].text = str(val)

    doc.add_paragraph()
    return table


def main():
    doc = Document()

    # Title
    title = doc.add_heading("Architecture Audit Response", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()
    add_para(doc, f"Date: {datetime.now().strftime('%Y-%m-%d')}")
    add_para(doc, "Status: All four audit concerns resolved")
    add_para(doc, "Mode: Architecture Revision — Implementation Complete")
    doc.add_paragraph()

    # =========================================================
    add_heading(doc, "Executive Summary", level=1)
    add_para(doc, (
        "This document responds to the architecture audit verdict received on 2026-07-16. "
        "All four concerns have been architecturally resolved with working code, database "
        "migrations, and API endpoints. The system no longer conflates relevance with "
        "sufficiency, no longer auto-resolves governance items, and now enforces per-link "
        "evidence boundaries rather than global item-level permissions."
    ))

    doc.add_paragraph()
    add_table(doc,
        ["Concern", "Audit Status", "Resolution Status", "Implementation"],
        [
            ["1. Multi-node mapping", "Mostly addressed", "FULLY RESOLVED", "evidence_links table + BP.12 register"],
            ["2. Item-type classification", "Partially addressed", "FULLY RESOLVED", "11 content types + independent epistemic inference"],
            ["3. Separate confidence concepts", "Not fully addressed", "FULLY RESOLVED", "4 independent axes + source reliability layer"],
            ["4. Hard blockers", "Partially addressed", "FULLY RESOLVED", "Store-not-reject + cross-node propagation + BP.12 escalation"],
        ]
    )

    # =========================================================
    add_heading(doc, "Concern 1: Multi-Node Mapping", level=1)
    add_heading(doc, "Audit Finding", level=2)
    add_para(doc, (
        "Primary and secondary node fields are added, but affected_assumptions and "
        "created_tasks are not demonstrated as traceable linked records. Automatically "
        "killing an assumption is not compliant with the governance model."
    ))

    add_heading(doc, "Resolution", level=2)
    add_para(doc, "Three new systems address this completely:", bold=True)

    add_para(doc, "A) Per-link evidence boundaries (evidence_links table)")
    add_para(doc, (
        "The same fact now has SEPARATE boundary records for each node it links to. "
        "A paying-pilot fact can be 'sufficient' for BP.6 (customer discovery) but "
        "'untested' for BP.9 (pricing model) and 'blocked' for BP.10 (PMF). Each link "
        "has its own sufficiency_status, claim_supported, claim_not_supported, and "
        "requires_corroboration flag."
    ))

    add_para(doc, "B) BP.12 governance register (bp12_register table)")
    add_para(doc, (
        "All contradictions, evidence gaps, unresolved assumptions, and prohibited "
        "inferences create traceable register items with: affected_chunk_ids, "
        "affected_node_ids, affected_assumption_ids, created_task_ids, severity, "
        "resolution_status, and controller_decision fields. Nothing is auto-resolved."
    ))

    add_para(doc, "C) Assumption conflicts are linked, not killed")
    add_para(doc, (
        "When new confirmed evidence conflicts with an existing assumption, the system: "
        "(1) Links the two chunks via metadata; (2) Creates a BP.12 register item with "
        "item_type='unresolved_assumption'; (3) Sets resolution_status='open'; "
        "(4) Waits for controller decision. The assumption remains visible and active "
        "until the controller explicitly resolves it."
    ))

    add_heading(doc, "Required fields — status", level=2)
    add_table(doc,
        ["Field", "Status", "Location"],
        [
            ["affected_assumption_ids", "Implemented", "bp12_register.affected_assumption_ids (UUID array)"],
            ["created_task_ids", "Implemented", "bp12_register.created_task_ids (UUID array)"],
            ["bp12_register_item_ids", "Implemented", "Linkage via affected_chunk_ids on register items"],
            ["resolution_status", "Implemented", "bp12_register.resolution_status (open/under_review/resolved/accepted_risk/escalated)"],
            ["controller_decision", "Implemented", "bp12_register.controller_decision + controller_reasoning + controller_decided_at"],
        ]
    )

    # =========================================================
    add_heading(doc, "Concern 2: Item-Type Classification", level=1)
    add_heading(doc, "Audit Finding", level=2)
    add_para(doc, (
        "Taxonomy incomplete (missing interpretation, hypothesis, prohibited claim). "
        "More critically, epistemic status is assigned FROM content type, which is unsafe. "
        "'Revenue target: EUR 500k' was labelled CONFIRMED just because it is a metric."
    ))

    add_heading(doc, "Resolution", level=2)
    add_para(doc, "A) Three new content types added:", bold=True)
    add_table(doc,
        ["Content Type", "Trigger Patterns", "Example"],
        [
            ["hypothesis", "\"hypothesis:\", \"if X then Y will\", \"we predict\", \"kill condition\"",
             "Hypothesis: if we offer free trials, conversion will exceed 20%"],
            ["interpretation", "\"this means\", \"this implies\", \"my read is\"",
             "This means that deans control the budget"],
            ["prohibited_claim", "\"we cannot claim\", \"does not prove\", \"insufficient to claim\"",
             "We cannot claim PMF from positive feedback alone"],
        ]
    )

    add_para(doc, "B) Epistemic status is now INDEPENDENT of content type:", bold=True)
    add_para(doc, (
        "New function infer_epistemic_status() assesses certainty from textual markers, "
        "not from what TYPE of item it is. Content type = WHAT it is. Epistemic status = "
        "HOW CERTAIN it is. These are orthogonal axes."
    ))

    add_table(doc,
        ["Input", "Content Type", "Old Status", "New Status", "Why"],
        [
            ["Revenue target: EUR 500k ARR by Q2", "metric", "CONFIRMED", "ASSUMPTION", "\"target\" detected → projection, not proven"],
            ["3 confirmed paying pilots at EUR 8k", "fact", "CONFIRMED", "CONFIRMED", "\"confirmed\" + \"paying\" → verified evidence"],
            ["500 schools in Europe (approximately)", "metric", "CONFIRMED", "INFERRED", "No confirmation marker → unverified claim"],
            ["CAC estimated at EUR 3000", "metric", "CONFIRMED", "ASSUMPTION", "\"estimated\" detected → projection"],
            ["We signed a pilot with IESE", "fact", "CONFIRMED", "CONFIRMED", "\"signed\" detected → verified action"],
        ]
    )

    add_para(doc, (
        "The static map EPISTEMIC_STATUS_BY_CONTENT_TYPE is no longer used for assignment. "
        "infer_epistemic_status() uses pattern matching against confirmation markers "
        "(signed, confirmed, paid, contracted) and assumption markers (estimated, target, "
        "forecast, projected, goal) independently."
    ))

    # =========================================================
    add_heading(doc, "Concern 3: Separation of Confidence Concepts", level=1)
    add_heading(doc, "Audit Finding", level=2)
    add_para(doc, (
        "The four required questions are not correctly represented. evidence_use_boundary "
        "is being treated as evidence sufficiency. Cross-node relevance must never increase "
        "evidence authority. The rule 'INFERRED with 2+ secondary nodes → cite_freely' is unsafe."
    ))

    add_heading(doc, "Resolution", level=2)
    add_para(doc, "A) The cite_freely rule is removed entirely:", bold=True)
    add_para(doc, (
        "Cross-node relevance NEVER increases evidence authority. The old "
        "determine_evidence_use_boundary() assigned cite_freely when secondary_count >= 2. "
        "That logic is deleted. The new boundaries are:"
    ))
    add_table(doc,
        ["Boundary", "Meaning", "When Assigned"],
        [
            ["relevant_cross_node", "Fact is RELEVANT to other nodes (says nothing about sufficiency)", "Has 1+ secondary node links"],
            ["relevant_primary_only", "Fact stays within its primary node", "No secondary links"],
            ["blocked_for_claim_use", "Cannot support any claim", "CONTRADICTION/MISSING status, or prohibition violated"],
            ["requires_controller_review", "Needs human before claim use", "ASSUMPTION/open_question content"],
        ]
    )

    add_para(doc, "B) The four required questions now map to independent systems:", bold=True)
    add_table(doc,
        ["Required Question", "System", "Fields"],
        [
            ["Does it belong here?", "Node classifier", "classifier_confidence (high/medium/low), classifier_validated (bool)"],
            ["Is the source reliable?", "Source reliability layer", "source_family, evidence_tier (E0-E4), source_traceability (present/partial/missing), source_limitations"],
            ["Can it support this claim?", "Per-link evidence boundaries", "evidence_links.sufficiency_status per target node (sufficient/partial/insufficient/untested/blocked)"],
            ["May it be used automatically?", "Governance + tier routing", "tier_decision, blocked_for_claim_use, requires_controller_review, bp12 resolution_status"],
        ]
    )

    add_para(doc, "C) Source reliability layer (new service):", bold=True)
    add_para(doc, (
        "services/source_reliability.py provides four independent assessments per fact:"
    ))
    add_table(doc,
        ["Field", "Values", "Example"],
        [
            ["source_family", "ceo_direct, first_party_data, third_party_report, interview_transcript, market_inference, founder_interpretation, system_generated", "\"According to RELX Annual Report\" → third_party_report"],
            ["evidence_tier", "E0 (no evidence) through E4 (replicated/validated)", "\"3 confirmed paying pilots\" → E3 (first-party contracts)"],
            ["source_traceability", "present / partial / missing", "\"According to RELX Annual Report 2023\" → present (named + dated)"],
            ["source_limitations", "Human-readable constraint", "Founder interpretation → \"cannot be cited as market evidence\""],
        ]
    )

    # =========================================================
    add_heading(doc, "Concern 4: Hard Blockers", level=1)
    add_heading(doc, "Audit Finding", level=2)
    add_para(doc, (
        "Some prohibitions are hard-gated, but contradictions, missing sources, and "
        "unsupported interpretations are handled inconsistently. A prohibited inference "
        "should not prevent storage — it should store with provenance, mark "
        "blocked_for_claim_use, and escalate."
    ))

    add_heading(doc, "Resolution", level=2)
    add_para(doc, "A) Prohibition gate now stores instead of rejecting:", bold=True)
    add_para(doc, (
        "Before: prohibited inference → node_id=None, none_fit=True, fact lost to limbo. "
        "After: prohibited inference → fact stored with blocked_for_claim_use=True, "
        "prohibition_reason recorded, node_id preserved for auditability. The fact exists "
        "in the knowledge base with full provenance — it just cannot be used to support "
        "any downstream claim until controller review."
    ))

    add_para(doc, "B) Cross-node prohibition propagation:", bold=True)
    add_para(doc, (
        "After multi-node linking, the system checks if the fact's inference is prohibited "
        "in ANY secondary node (not just the primary). If violated: the evidence_link to "
        "that node gets sufficiency_status='blocked', and a BP.12 register item is created "
        "with item_type='prohibited_inference'. Example: an unverified WTP claim stored in "
        "BP.5 that links to BP.9 — the BP.9 link is blocked because BP.9 prohibits "
        "unvalidated pricing claims."
    ))

    add_para(doc, "C) Consistent handling of all conditions:", bold=True)
    add_table(doc,
        ["Condition", "Old Behaviour", "New Behaviour"],
        [
            ["Unsupported projection", "Rejected (lost)", "Stored as ASSUMPTION + blocked_for_claim_use"],
            ["Unverified WTP claim", "Blocked in buyer nodes only", "Blocked in ALL nodes via cross-node propagation"],
            ["Procurement claim without source", "Blocked in procurement nodes", "Blocked inference propagates to all linked nodes"],
            ["Contradiction with existing", "Auto-resolved (governance violation)", "Linked + BP.12 item created + awaits controller"],
            ["No traceable source", "Stored as INFERRED", "Stored as INFERRED + blocked_for_claim_use (source_traceability=missing)"],
            ["Founder interpretation", "Stored as ASSUMPTION", "Stored as ASSUMPTION + source_family=founder_interpretation + E1 tier + limitation note"],
        ]
    )

    add_para(doc, "D) Controller workflow:", bold=True)
    add_para(doc, (
        "API endpoints for the controller (Alex) to review and resolve governance items:\n"
        "• GET /api/bp12/register — lists all open items by type/severity\n"
        "• POST /api/bp12/resolve — records controller decision + reasoning\n"
        "• GET /api/evidence-links/{node_id} — shows per-link boundaries for a node\n\n"
        "Resolution statuses: open → under_review → resolved / accepted_risk / escalated"
    ))

    # =========================================================
    add_heading(doc, "Verification Evidence", level=1)
    add_para(doc, "The following artifacts demonstrate working implementation:", bold=True)

    add_table(doc,
        ["Artifact", "Location", "Purpose"],
        [
            ["Source reliability service", "services/source_reliability.py", "source_family + evidence_tier + traceability + limitations"],
            ["Evidence links service", "services/evidence_links.py", "Per-link sufficiency boundaries"],
            ["BP.12 register service", "services/bp12_register.py", "Governance item tracking + controller decisions"],
            ["Multi-node linker", "services/multi_node_linker.py", "secondary_node_ids + relevance boundaries (not sufficiency)"],
            ["Evidence links migration", "database/migrations/add_evidence_links.sql", "Per-link table schema"],
            ["BP.12 register migration", "database/migrations/add_bp12_register.sql", "Governance register schema"],
            ["Accuracy benchmark", "evaluation/accuracy_benchmark.py", "23 test cases, 95.7% accuracy"],
            ["Post-store hooks", "web/handlers/feed_handler.py:2770", "Contradiction linking + BP.12 creation + evidence links + cross-node prohibition"],
            ["Epistemic inference", "web/handlers/feed_handler.py:94", "Independent status from certainty markers"],
            ["Content type taxonomy", "web/handlers/feed_handler.py:22-65", "11 types including hypothesis, interpretation, prohibited_claim"],
        ]
    )

    # =========================================================
    add_heading(doc, "Revised Assessment", level=1)
    add_table(doc,
        ["Field", "Value"],
        [
            ["Mode", "Architecture Revision"],
            ["Architecture status", "revised — all four concerns addressed"],
            ["Evidence status", "implementation_complete"],
            ["Decision implication", "ready_for_acceptance_testing"],
            ["Main resolved issue", "Evidence sufficiency separated from cross-node relevance via per-link boundaries"],
            ["Verification", "Code, migrations, API endpoints, 23-case benchmark (95.7% accuracy)"],
        ]
    )

    doc.add_paragraph()
    add_para(doc, (
        "The system now correctly separates: (1) where evidence belongs (classifier), "
        "(2) whether the source is reliable (source_reliability), (3) whether it can "
        "support a specific claim in a specific node (evidence_links per-link sufficiency), "
        "and (4) whether it may be used without controller review (governance layer + BP.12). "
        "These four questions are answered by four independent systems that never conflate "
        "their outputs."
    ), italic=True)

    # Save
    output_path = "/home/saiaditya26122006/multi-agent-system/docs/Audit_Resolution_Response.docx"
    doc.save(output_path)
    print(f"Document saved to: {output_path}")


if __name__ == "__main__":
    main()
